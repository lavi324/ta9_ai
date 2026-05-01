import os
import uuid
import pathlib
import time
import traceback
from typing import List, Optional, Tuple, Dict, Any, Deque
from collections import Counter, defaultdict, deque

from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from dotenv import load_dotenv
import requests
import math
import hashlib
import fnmatch
import chromadb
import base64
from datetime import datetime
import asyncio
import re
from html import unescape
import subprocess
import tempfile
import shutil
import json
import mimetypes
import difflib
from io import BytesIO
from urllib.parse import unquote, quote, urlsplit, urlunsplit, parse_qsl, urlencode
import threading
from concurrent.futures import ThreadPoolExecutor

# Try to import optional dependencies
try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

import csv

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import fitz
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import tiktoken
    HAS_TIKTOKEN = True
except ImportError:
    HAS_TIKTOKEN = False

load_dotenv()

# ---------------------------------------------------------------------
# Env & constants
# ---------------------------------------------------------------------

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_URL = os.getenv("OPENAI_URL", "https://api.openai.com/v1").rstrip("/")
OPENAI_EMBEDDING_MODEL = os.getenv("OPENAI_EMBEDDING_MODEL", "text-embedding-3-large")
OPENAI_CHAT_MODEL = os.getenv("OPENAI_CHAT_MODEL", "gpt-4o-mini")
OPENAI_RERANK_MODEL = os.getenv("OPENAI_RERANK_MODEL", OPENAI_CHAT_MODEL)
OPENAI_VISION_MODEL = os.getenv("OPENAI_VISION_MODEL", "gpt-4o")
LOG_FULL_EMBEDDINGS = os.getenv("LOG_FULL_EMBEDDINGS", "false").lower() in ("1", "true", "yes")
EMBED_PREVIEW_COUNT = int(os.getenv("EMBED_PREVIEW_COUNT", "8"))

# SINGLE ROOT DIRECTORY: all .md files are under /app/wiki_files
WIKI_ROOT = os.getenv("WIKI_DIR", "/app/wiki_files")

# Directory exposed for "Add Knowledge" server file listing
KNOWLEDGE_DIR = os.getenv("KNOWLEDGE_DIR", WIKI_ROOT)

CHROMA_DIR = os.getenv("CHROMA_DIR", "/app/chroma_db")
SYSTEM_PROMPT_FILE = os.getenv("SYSTEM_PROMPT_FILE", "/app/system_prompt_template.txt")
SYSTEM_PROMPT_DISPLAY_PATH = os.getenv("SYSTEM_PROMPT_DISPLAY_PATH", "backend/system_prompt_template.txt")

# Chunk configuration optimized for large files (2M+ tokens)
# Larger chunks = fewer embeddings = faster processing of massive files
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "2000"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "300"))

COLLECTION_NAME = "Intsight"
MEMORY_COLLECTION_NAME = "New_Knowledge"
VECTOR_DB_PREVIEW_CHARS = 280
VECTOR_DB_DEFAULT_CHUNK_LIMIT = 25
VECTOR_DB_MAX_CHUNK_LIMIT = 100

# PDF processing: optimized for large files
# PDF_MAX_PAGES=0 means unlimited (process entire file)
PDF_MAX_PAGES = int(os.getenv("PDF_MAX_PAGES", "0"))
PDF_OCR_ENABLED = os.getenv("PDF_OCR_ENABLED", "true").lower() in ("1", "true", "yes")
# For image-heavy PDFs, analyze ALL pages with vision model (not just first 50)
PDF_OCR_MAX_PAGES = int(os.getenv("PDF_OCR_MAX_PAGES", "0"))  # 0 = unlimited = analyze all pages
PDF_TABLE_EXTRACTION_ENABLED = os.getenv("PDF_TABLE_EXTRACTION_ENABLED", "true").lower() in ("1", "true", "yes")
# Aggressive image extraction from ALL PDF pages (not just fallback)
PDF_IMAGE_EXTRACTION_ENABLED = os.getenv("PDF_IMAGE_EXTRACTION_ENABLED", "true").lower() in ("1", "true", "yes")
PDF_MAX_IMAGES_PER_PAGE = int(os.getenv("PDF_MAX_IMAGES_PER_PAGE", "4"))
PDF_MIN_IMAGE_AREA_RATIO = float(os.getenv("PDF_MIN_IMAGE_AREA_RATIO", "0.015"))

# Vision model configuration
VISION_MAX_TOKENS = int(os.getenv("VISION_MAX_TOKENS", "2500"))
VISION_TIMEOUT_SECONDS = int(os.getenv("VISION_TIMEOUT_SECONDS", "120"))
VISION_RETRY_COUNT = int(os.getenv("VISION_RETRY_COUNT", "2"))

# Large file processing support (up to 2M tokens per uploaded or fetched file)
MAX_FILE_SIZE_TOKENS = int(os.getenv("MAX_FILE_SIZE_TOKENS", "2000000"))
BLOCK_PROCESSING_ENABLED = os.getenv("BLOCK_PROCESSING_ENABLED", "true").lower() in ("1", "true", "yes")
DETAILED_INGEST_LOGS = os.getenv("DETAILED_INGEST_LOGS", "true").lower() in ("1", "true", "yes")
LLM_MAX_INPUT_TOKENS = 120000

TABLE_ROW_START_TAG = "[TABLE_ROW]"
TABLE_ROW_END_TAG = "[/TABLE_ROW]"

DEFAULT_SYSTEM_PROMPT_TEMPLATE = (
    "You are a TA9 / IntSight customer support assistant with deep technical knowledge. "
    "Your role is to provide clear, accurate, and helpful answers in a natural, conversational way.\n\n"
    "INSTRUCTIONS:\n"
    "1. When the user attaches files (images, documents), those files are your PRIMARY source - answer directly from them.\n"
    "2. When you see [Image Content from ...] sections, that means the image has been analyzed - describe what you see in the analysis naturally.\n"
    "3. For images: Provide a comprehensive explanation of what is shown in the provided analysis. Highlight key details, technical context, and likely causes when relevant.\n"
    "4. When a ticket is selected, use it to understand the user's issue and provide relevant solutions.\n"
    "4.1 For ticket responses, begin with a short explanation of the ticket before listing the solution steps.\n"
    "5. Use knowledge base context to supplement your answer or provide additional related information.\n"
    "6. Use fenced code blocks (```bash, ```python, ```sql, etc.) only for real commands, code, SQL queries, JSON, config files, or other text the user may copy and run exactly as-is.\n"
    "6.1 Do NOT use fenced code blocks for numbered steps, UI navigation sequences, button names, field labels, prose instructions, or short key-value notes such as 'Source: ...' or 'Type: ...'. Render those as normal Markdown text, numbered lists, or bullet points.\n"
    "6.2 A summary of steps (e.g., '1. Open Admin Studio and log in. 2. Locate the data model...') is NOT code — always render it as a plain numbered list, never inside a code fence.\n"
    "7. If the context contains redundant or overlapping information, synthesize it into a single coherent answer.\n"
    "8. Do NOT repeat information from different sources - intelligently merge related points.\n"
    "8.1 Use a balanced approach across both collections (Intsight and New_Knowledge) and do not assume one is always better.\n"
    "8.2 If relevant details are split across both collections, combine them into one integrated, consistent answer.\n"
    "9. Grounding is mandatory: only state facts supported by the provided context sections.\n"
    "9.1 If a requested detail is not in context, explicitly say it is not available in the current knowledge context.\n"
    "9.2 Never invent commands, configuration keys, UI paths, API names, or procedural steps.\n"
    "9.3 If the response is about a BUG ticket and the issue is not fully resolved, frame the ending as an INTERNAL escalation from support to engineering or R&D. Never tell the support engineer to 'contact support', 'reach out to support', or use wording that treats support as the customer.\n"
    "9.4 When the context provides a specific value (URL, endpoint, path, command, setting name, etc.), state it directly and confidently. "
    "NEVER hedge with words like 'usually', 'typically', 'something like', 'often', 'generally', or 'might be' when the exact value is present in the context. "
    "Replace vague phrasing with the concrete detail from the context.\n"
    "9.5 If the context contains the answer, present it as established fact — do NOT fall back to generic guidance or speculative patterns. "
    "Only use generic language when the specific detail is genuinely absent from all provided context.\n"
    "10. Keep a professional, helpful tone that encourages follow-up questions.\n"
    "11. Answer naturally and conversationally - avoid rigid structured formats unless specifically requested.\n"
    "12. For Intsight or system configuration guidance, ALWAYS provide instructions using Admin Studio (UI-based configuration) by default. Do NOT include SQL queries, INSERT/UPDATE statements, or direct database table manipulation unless the user explicitly asks.\n"
    "13. Provide database-level (DB) configuration instructions ONLY when the user explicitly mentions 'database', 'SQL', 'DB', 'table', 'query', or specifically requests backend/database-level steps. If the context contains both Admin Studio steps and DB-level steps for the same task, present ONLY the Admin Studio steps unless DB steps are explicitly requested.\n"
    "13.1 When DB steps are explicitly requested, you may include SQL queries and table-level instructions alongside the Admin Studio approach.\n"
    "14. Never hard-refuse when at least partial context exists; provide the best grounded answer possible, explicitly flag uncertainty, and ask one focused clarifying question if needed.\n"
    "14.1 For BUG tickets that remain unresolved, end with a brief internal handoff note to engineering or R&D and the information support should gather before escalation. Do NOT say 'contact support'.\n\n"
    "SOURCE SEPARATION AND ACCURACY:\n"
    "15. Each context block (labeled PRIMARY MATCH or SUPPLEMENTAL MATCH) comes from a DIFFERENT document. Treat each document as a separate, self-contained source.\n"
    "15.1 NEVER combine or merge procedural steps from different source documents into a single procedure. Each document describes its own workflow — mixing steps from Document A with steps from Document B creates incorrect instructions.\n"
    "15.2 If two source documents describe different procedures for related but distinct tasks, present them separately with clear labels. Do NOT interleave their steps.\n"
    "15.3 When multiple sources cover the same topic, use the one that best matches the user's specific question. Use supplemental sources only for additional context that the primary source does not cover.\n\n"
    "QUESTION ANALYSIS:\n"
    "16. Before answering, carefully analyze the user's question to understand what they ALREADY HAVE vs what they NEED:\n"
    "16.1 If the user says 'in a data model' or 'in my data model', the data model already exists — do NOT provide steps to create it. Focus on the specific operation they are asking about.\n"
    "16.2 If the user says 'configure X' or 'define X', they want to set up that specific feature — do NOT provide steps for creating the parent object that already exists.\n"
    "16.3 Pay attention to prepositions: 'in', 'on', 'for', 'within' usually indicate an existing context. 'Create', 'new', 'set up from scratch' indicate they need to build something new.\n"
    "16.4 Match the scope of your answer to the scope of the question. If the user asks about one specific feature within a larger system, focus your answer on that feature — do not explain the entire system setup.\n"
)

app = FastAPI(title="Wiki RAG API")

# Track if a structured "first ticket response" was already sent per ticket key (id or URL fallback)
ticket_first_reply_done: Dict[str, bool] = {}

# Lightweight in-memory conversation memory (per conversation key)
MAX_CONVERSATION_MESSAGES = 12
conversation_store: Dict[str, Deque[Dict[str, str]]] = defaultdict(
    lambda: deque(maxlen=MAX_CONVERSATION_MESSAGES)
)
conversation_state_store: Dict[str, Dict[str, Any]] = {}

# In-memory cancellation flags for long-running knowledge ingestion requests.
knowledge_cancel_lock = threading.Lock()
knowledge_cancel_flags: Dict[str, bool] = {}

# In-memory job store so knowledge_add returns immediately and the UI polls for the result.
knowledge_job_store: Dict[str, dict] = {}
knowledge_job_store_lock = threading.Lock()

# In-memory cache of merged vector document content for large document viewing.
VECTOR_DOCUMENT_CACHE_MAX_ITEMS = int(os.getenv("VECTOR_DOCUMENT_CACHE_MAX_ITEMS", "24"))
vector_document_cache: Dict[str, dict] = {}
vector_document_cache_lock = threading.Lock()

# In-memory cache for collection metadata to avoid repeated full-collection dumps.
# Keys: collection name, Values: {"metadatas": [...], "documents": [...], "timestamp": float}
_collection_cache: Dict[str, dict] = {}
_collection_cache_lock = threading.Lock()
_COLLECTION_CACHE_TTL = 120  # seconds


def _get_cached_collection_data(col_obj, collection_name: str, include_documents: bool = False) -> dict:
    """Return cached collection data (metadatas + optionally documents). Refreshes every TTL seconds."""
    cache_key = f"{collection_name}:{'docs' if include_documents else 'meta'}"
    with _collection_cache_lock:
        cached = _collection_cache.get(cache_key)
        if cached and (time.time() - cached["timestamp"]) < _COLLECTION_CACHE_TTL:
            return cached["data"]
    includes = ["metadatas"]
    if include_documents:
        includes.append("documents")
    data = col_obj.get(include=includes)
    with _collection_cache_lock:
        _collection_cache[cache_key] = {"data": data, "timestamp": time.time()}
    return data


def invalidate_collection_cache(collection_name: Optional[str] = None):
    """Clear collection cache after modifications (add/delete/update)."""
    with _collection_cache_lock:
        if collection_name:
            keys_to_remove = [k for k in _collection_cache if k.startswith(f"{collection_name}:")]
            for k in keys_to_remove:
                del _collection_cache[k]
        else:
            _collection_cache.clear()


def _request_log_prefix(request_id: Optional[str]) -> str:
    return f"[rid={request_id}]" if request_id else "[rid=-]"


def _ingest_log(message: str, request_id: Optional[str] = None, force: bool = False) -> None:
    if DETAILED_INGEST_LOGS or force:
        print(f"[INGEST]{_request_log_prefix(request_id)} {message}")


def _rag_log_step(step_number: int, total_steps: int, title: str, detail: Optional[str] = None) -> None:
    prefix = f"[RAG][Step {step_number}/{total_steps}] {title}"
    if detail:
        print(f"{prefix}: {detail}")
    else:
        print(prefix)


def _mark_knowledge_request_started(request_id: Optional[str]) -> None:
    if not request_id:
        return
    with knowledge_cancel_lock:
        # If cancel arrived first, do not overwrite it.
        if knowledge_cancel_flags.get(request_id) is True:
            _ingest_log("request start observed after cancel signal; keeping canceled=true", request_id, force=True)
            return
        knowledge_cancel_flags[request_id] = False
    _ingest_log("request marked active", request_id)


def _cancel_knowledge_request(request_id: Optional[str]) -> bool:
    if not request_id:
        return False
    with knowledge_cancel_lock:
        existed = request_id in knowledge_cancel_flags
        knowledge_cancel_flags[request_id] = True
    _ingest_log(f"cancel signal stored (was_active={existed})", request_id, force=True)
    return existed


def _is_knowledge_request_cancelled(request_id: Optional[str]) -> bool:
    if not request_id:
        return False
    with knowledge_cancel_lock:
        return bool(knowledge_cancel_flags.get(request_id, False))


def _clear_knowledge_request(request_id: Optional[str]) -> None:
    if not request_id:
        return
    with knowledge_cancel_lock:
        knowledge_cancel_flags.pop(request_id, None)
    _ingest_log("request state cleared", request_id)


def _store_job_result(request_id: Optional[str], result: dict) -> None:
    if not request_id:
        return
    with knowledge_job_store_lock:
        # Evict very old completed jobs (older than 24 h) to keep memory bounded.
        cutoff = datetime.utcnow().isoformat()
        keys_to_remove = [
            k for k, v in knowledge_job_store.items()
            if v.get("status") == "done" and v.get("stored_at", cutoff) < cutoff
        ]
        # Keep at most 100 completed jobs regardless of age.
        done_jobs = [k for k, v in knowledge_job_store.items() if v.get("status") == "done"]
        if len(done_jobs) > 100:
            keys_to_remove += done_jobs[:-100]
        for k in keys_to_remove:
            knowledge_job_store.pop(k, None)
        result["stored_at"] = datetime.utcnow().isoformat() + "Z"
        knowledge_job_store[request_id] = result


def _get_job_result(request_id: Optional[str]) -> Optional[dict]:
    if not request_id:
        return None
    with knowledge_job_store_lock:
        return knowledge_job_store.get(request_id)


def _vector_document_cache_key(collection_name: str, source: str) -> str:
    return f"{collection_name}::{source}"


def _get_cached_vector_document_payload(collection_name: str, source: str) -> Optional[dict]:
    key = _vector_document_cache_key(collection_name, source)
    with vector_document_cache_lock:
        payload = vector_document_cache.get(key)
        if payload is None:
            return None
        # Refresh insertion order for simple LRU behavior.
        vector_document_cache.pop(key, None)
        vector_document_cache[key] = payload
        return dict(payload)


def _store_cached_vector_document_payload(
    collection_name: str,
    source: str,
    full_content: str,
    token_count: int,
) -> None:
    key = _vector_document_cache_key(collection_name, source)
    payload = {
        "full_content": full_content,
        "token_count": int(token_count or 0),
        "cached_at": datetime.utcnow().isoformat() + "Z",
    }
    with vector_document_cache_lock:
        vector_document_cache.pop(key, None)
        vector_document_cache[key] = payload
        while len(vector_document_cache) > VECTOR_DOCUMENT_CACHE_MAX_ITEMS:
            oldest_key = next(iter(vector_document_cache))
            vector_document_cache.pop(oldest_key, None)


def _invalidate_cached_vector_document_payload(collection_name: str, source: str) -> None:
    key = _vector_document_cache_key(collection_name, source)
    with vector_document_cache_lock:
        vector_document_cache.pop(key, None)


def _validate_system_prompt_template(template: str) -> str:
    candidate = str(template or "")
    if not candidate.strip():
        raise ValueError("System prompt template cannot be empty.")

    return candidate


def _read_system_prompt_template() -> str:
    try:
        with open(SYSTEM_PROMPT_FILE, "r", encoding="utf-8") as handle:
            template = handle.read()
        return _validate_system_prompt_template(template)
    except FileNotFoundError:
        return DEFAULT_SYSTEM_PROMPT_TEMPLATE
    except Exception as exc:
        print(f"[PROMPT][WARN] Failed to read system prompt template from {SYSTEM_PROMPT_FILE}: {exc}")
        return DEFAULT_SYSTEM_PROMPT_TEMPLATE


def _write_system_prompt_template(template: str) -> str:
    validated = _validate_system_prompt_template(template)
    parent = os.path.dirname(SYSTEM_PROMPT_FILE)
    if parent:
        os.makedirs(parent, exist_ok=True)
    with open(SYSTEM_PROMPT_FILE, "w", encoding="utf-8") as handle:
        handle.write(validated)
    return validated


def _build_system_prompt(
    question: str,
    combined_context: str,
    foundational_instruction: str = "",
    ta9_instruction: str = "",
) -> Tuple[str, str]:
    """Return (system_message, user_message) for proper role separation."""
    template = _read_system_prompt_template().strip()

    # --- System message: persona + rules ---
    system_parts: List[str] = [template]
    if foundational_instruction and foundational_instruction.strip():
        system_parts.append(foundational_instruction.strip())
    if ta9_instruction and ta9_instruction.strip():
        system_parts.append(ta9_instruction.strip())
    system_msg = "\n\n".join(part for part in system_parts if part)

    # --- User message: retrieved context + question ---
    user_parts: List[str] = []
    if combined_context and combined_context.strip():
        user_parts.append(
            "Below are the retrieved context documents. READ THEM CAREFULLY — your answer MUST be based on their content.\n\n"
            + combined_context.strip()
        )
        user_parts.append(
            "IMPORTANT REMINDER: The documents above contain the information you need. "
            "Extract and use the relevant details from these documents to answer the question below. "
            "Do NOT say the information is unavailable if it appears anywhere in the documents above."
        )
    user_parts.append(f"Question: {question}")
    user_parts.append("Answer (based on the retrieved documents above):")
    user_msg = "\n\n".join(part for part in user_parts if part)

    return system_msg, user_msg


# ---------------------------------------------------------------------
# Embedding + LLM
# ---------------------------------------------------------------------

def embed_text(text: str) -> List[float]:
    snippet = text[:120].replace("\n", " ")
    print(f"[EMBED] Calling OpenAI embeddings model={OPENAI_EMBEDDING_MODEL} len={len(text)} snippet='{snippet}...'")
    start = time.time()
    headers = {"Authorization": f"Bearer {OPENAI_API_KEY}", "Content-Type": "application/json"}
    body = {"model": OPENAI_EMBEDDING_MODEL, "input": text}
    try:
        resp = requests.post(f"{OPENAI_URL}/embeddings", json=body, headers=headers, timeout=120)
    except Exception as e:
        print(f"[EMBED][ERROR] Exception while calling OpenAI embeddings: {e}")
        traceback.print_exc()
        raise RuntimeError(f"Failed to call OpenAI embeddings: {e}")
    duration = time.time() - start
    print(f"[EMBED] OpenAI embeddings status={resp.status_code} took={duration:.2f}s")

    if resp.status_code != 200:
        print(f"[EMBED][ERROR] Non-200 from OpenAI embeddings: {resp.text[:400]}")
        raise RuntimeError(f"OpenAI embeddings error: {resp.text}")

    data = resp.json()
    try:
        emb = data["data"][0]["embedding"]
    except Exception as e:
        print(f"[EMBED][ERROR] Unexpected OpenAI response shape: {data}")
        raise RuntimeError(f"OpenAI embeddings response missing embedding: {e}")

    preview = ",".join([f"{x:.6f}" for x in emb[:EMBED_PREVIEW_COUNT]])
    norm = math.sqrt(sum([x * x for x in emb]))
    h = hashlib.sha256(
        ",".join([f"{x:.6f}" for x in emb[:16]]).encode("utf-8")
    ).hexdigest()[:12]
    print(f"[EMBED] Got OpenAI embedding length={len(emb)} preview=[{preview}] norm={norm:.6f} hash={h}")
    if LOG_FULL_EMBEDDINGS:
        print(f"[EMBED][FULL] {emb}")
    return emb


def _embed_texts_single_batch(texts: List[str]) -> List[List[float]]:
    """Call OpenAI embeddings API for a single batch of texts."""
    headers = {"Authorization": f"Bearer {OPENAI_API_KEY}", "Content-Type": "application/json"}
    body = {"model": OPENAI_EMBEDDING_MODEL, "input": texts}
    try:
        resp = requests.post(f"{OPENAI_URL}/embeddings", json=body, headers=headers, timeout=120)
    except Exception as e:
        print(f"[EMBED_BATCH][ERROR] Exception: {e}")
        traceback.print_exc()
        raise RuntimeError(f"Failed to call OpenAI embeddings batch: {e}")
    if resp.status_code != 200:
        print(f"[EMBED_BATCH][ERROR] Non-200: {resp.text[:400]}")
        raise RuntimeError(f"OpenAI embeddings batch error: {resp.text}")
    data = resp.json()
    try:
        sorted_items = sorted(data["data"], key=lambda x: x["index"])
        return [item["embedding"] for item in sorted_items]
    except Exception as e:
        print(f"[EMBED_BATCH][ERROR] Unexpected response shape: {e}")
        raise RuntimeError(f"OpenAI embeddings batch parse error: {e}")


def embed_texts_batch(texts: List[str]) -> List[List[float]]:
    """Batch embed multiple texts, splitting into sub-batches of 100 to avoid API limits."""
    if not texts:
        return []
    if len(texts) == 1:
        return [embed_text(texts[0])]
    BATCH_SIZE = 100
    total = len(texts)
    print(f"[EMBED_BATCH] Calling OpenAI embeddings model={OPENAI_EMBEDDING_MODEL} total={total} sub_batches={-(-total // BATCH_SIZE)}")
    start = time.time()
    all_embeddings: List[List[float]] = []
    for i in range(0, total, BATCH_SIZE):
        batch = texts[i:i + BATCH_SIZE]
        batch_start = time.time()
        print(f"[EMBED_BATCH] Sub-batch {i // BATCH_SIZE + 1}: texts {i+1}-{min(i+len(batch), total)} of {total}")
        batch_embeddings = _embed_texts_single_batch(batch)
        all_embeddings.extend(batch_embeddings)
        print(f"[EMBED_BATCH] Sub-batch done in {time.time() - batch_start:.2f}s")
    duration = time.time() - start
    print(f"[EMBED_BATCH] Got {len(all_embeddings)} embeddings in {duration:.2f}s")
    return all_embeddings


def call_llm(prompt: str, temperature: float = 0.2, model: Optional[str] = None, max_tokens: Optional[int] = None, system_message: Optional[str] = None) -> str:
    """Call OpenAI chat for final answer."""
    selected_model = (model or OPENAI_CHAT_MODEL).strip() or OPENAI_CHAT_MODEL
    print(f"[LLM] Calling OpenAI chat model={selected_model} prompt_len={len(prompt)} temp={temperature} max_tokens={max_tokens} has_system={bool(system_message)}")
    start = time.time()
    headers = {"Authorization": f"Bearer {OPENAI_API_KEY}", "Content-Type": "application/json"}
    if system_message:
        messages = [
            {"role": "system", "content": system_message},
            {"role": "user", "content": prompt},
        ]
    else:
        messages = [{"role": "user", "content": prompt}]
    body = {
        "model": selected_model,
        "messages": messages,
        "temperature": temperature,
    }
    if max_tokens is not None:
        body["max_tokens"] = max_tokens
    try:
        resp = requests.post(f"{OPENAI_URL}/chat/completions", json=body, headers=headers, timeout=120)
    except Exception as e:
        print(f"[LLM][ERROR] Exception while calling OpenAI chat: {e}")
        traceback.print_exc()
        raise RuntimeError(f"Failed to call OpenAI chat: {e}")
    duration = time.time() - start
    print(f"[LLM] OpenAI chat status={resp.status_code} took={duration:.2f}s")

    if resp.status_code != 200:
        print(f"[LLM][ERROR] Non-200 from OpenAI chat: {resp.text[:400]}")
        raise RuntimeError(f"OpenAI chat error: {resp.text}")

    try:
        data = resp.json()
        msg = data.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
    except Exception as e:
        print(f"[LLM][ERROR] Unexpected OpenAI chat response: {resp.text[:400]}")
        traceback.print_exc()
        raise RuntimeError(f"OpenAI chat parse error: {e}")

    print(f"[LLM] Got OpenAI answer length={len(msg)}")
    return msg


def call_llm_stream(prompt: str, temperature: float = 0.2, model: str = None, max_tokens: int = None, system_message: str = None):
    """Call OpenAI chat with streaming. Yields text chunks as they arrive."""
    selected_model = (model or OPENAI_CHAT_MODEL).strip() or OPENAI_CHAT_MODEL
    print(f"[LLM-STREAM] Calling OpenAI chat model={selected_model} prompt_len={len(prompt)} temp={temperature} max_tokens={max_tokens} has_system={bool(system_message)}")
    headers = {"Authorization": f"Bearer {OPENAI_API_KEY}", "Content-Type": "application/json"}
    if system_message:
        messages = [
            {"role": "system", "content": system_message},
            {"role": "user", "content": prompt},
        ]
    else:
        messages = [{"role": "user", "content": prompt}]
    body = {
        "model": selected_model,
        "messages": messages,
        "temperature": temperature,
        "stream": True,
    }
    if max_tokens is not None:
        body["max_tokens"] = max_tokens
    try:
        resp = requests.post(f"{OPENAI_URL}/chat/completions", json=body, headers=headers, timeout=120, stream=True)
    except Exception as e:
        print(f"[LLM-STREAM][ERROR] Exception: {e}")
        raise RuntimeError(f"Failed to call OpenAI chat: {e}")
    if resp.status_code != 200:
        print(f"[LLM-STREAM][ERROR] Non-200: {resp.text[:400]}")
        raise RuntimeError(f"OpenAI chat error: {resp.text}")
    for line in resp.iter_lines(decode_unicode=True):
        if not line or not line.startswith("data: "):
            continue
        payload = line[6:]
        if payload.strip() == "[DONE]":
            break
        try:
            chunk = json.loads(payload)
            delta = chunk.get("choices", [{}])[0].get("delta", {})
            text = delta.get("content", "")
            if text:
                yield text
        except Exception:
            pass
    print("[LLM-STREAM] Stream complete")


# ---------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------

def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    tokens_estimate = len(text) // 4
    print(f"[CHUNK] Splitting text len={len(text)} chars (~{tokens_estimate} tokens) size={size} overlap={overlap}")
    chunks: List[str] = []
    start_idx = 0
    while start_idx < len(text):
        end = min(start_idx + size, len(text))
        chunk = text[start_idx:end].strip()
        if chunk:
            chunks.append(chunk)
        start_idx += size - overlap
    total_tokens = sum(len(c) // 4 for c in chunks)
    print(f"[CHUNK] Produced {len(chunks)} chunks (~{total_tokens} tokens total)")
    return chunks


def chunk_text_preserve_table_rows(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    if not text:
        return []

    pattern = re.compile(r"\[TABLE_ROW\][\s\S]*?\[/TABLE_ROW\]")
    chunks: List[str] = []
    cursor = 0

    for match in pattern.finditer(text):
        regular_part = text[cursor:match.start()].strip()
        if regular_part:
            chunks.extend(chunk_text(regular_part, size=size, overlap=overlap))

        table_row_block = match.group(0).strip()
        if table_row_block:
            if len(table_row_block) <= size:
                chunks.append(table_row_block)
            else:
                chunks.extend(chunk_text(table_row_block, size=size, overlap=0))

        cursor = match.end()

    tail = text[cursor:].strip()
    if tail:
        chunks.extend(chunk_text(tail, size=size, overlap=overlap))

    print(f"[CHUNK] Produced {len(chunks)} chunks with TABLE_ROW preservation")
    return chunks


def extract_table_row_metadata(chunk: str) -> Dict[str, int]:
    if not chunk:
        return {}
    match = re.search(r"\[TABLE_ROW\]\s*page=(\d+)\s+table=(\d+)\s+row=(\d+)", chunk)
    if not match:
        return {}
    return {
        "table_page": int(match.group(1)),
        "table_index": int(match.group(2)),
        "table_row_index": int(match.group(3)),
    }


# ---------------------------------------------------------------------
# Iterate ALL markdown files under WIKI_ROOT
# ---------------------------------------------------------------------

def iter_markdown_files():
    root = pathlib.Path(WIKI_ROOT)
    print(f"[FILES] Scanning for .md files under: {WIKI_ROOT}")
    if not root.exists():
        print(f"[FILES][WARN] WIKI_ROOT path does not exist: {WIKI_ROOT}")
        return
    for path in root.rglob("*.md"):
        if path.is_file():
            yield path


# ---------------------------------------------------------------------
# ChromaDB Setup
# ---------------------------------------------------------------------

print(f"[CHROMA] Initializing PersistentClient path={CHROMA_DIR}")
client = chromadb.PersistentClient(path=CHROMA_DIR)
print(f"[CHROMA] Getting/creating collection '{COLLECTION_NAME}'")
collection = client.get_or_create_collection(name=COLLECTION_NAME)
print(f"[CHROMA] Getting/creating memory collection '{MEMORY_COLLECTION_NAME}'")
memory_collection = client.get_or_create_collection(name=MEMORY_COLLECTION_NAME)

# ---------------------------------------------------------------------
# Azure DevOps config & helpers
# ---------------------------------------------------------------------

ADO_ORG = os.getenv("ADO_ORG")
ADO_PROJECT = os.getenv("ADO_PROJECT")
ADO_PROJECT_TICKET_SUMMARY = os.getenv("ADO_PROJECT_TICKET_SUMMARY", "").strip('"')
ADO_PAT = os.getenv("ADO_PAT")
ADO_SUPPORT_TICKET_QUERY_TARGETS = os.getenv(
    "ADO_SUPPORT_TICKET_QUERY_TARGETS",
    "TA9 Support::My Queries/external tickets|STE Support Board::My Queries/external tickets",
)
ADO_SUPPORT_TICKET_ASSIGNED_TO = os.getenv(
    "ADO_SUPPORT_TICKET_ASSIGNED_TO",
    "Support <support@ta-9.com>",
)
ADO_KEYWORD_TICKET_QUERY_TARGETS = os.getenv(
    "ADO_KEYWORD_TICKET_QUERY_TARGETS",
    f"{ADO_PROJECT}::My Queries/New Query" if ADO_PROJECT else "",
)
# Direct WIQL override for keyword ticket search – bypasses saved query fetch entirely.
# If set, this WIQL is executed directly to get the pool of ticket IDs to search.
ADO_KEYWORD_WIQL = os.getenv("ADO_KEYWORD_WIQL", "").strip()

def _ado_headers() -> dict:
    if not (ADO_ORG and ADO_PROJECT and ADO_PAT):
        raise RuntimeError("Azure DevOps is not configured. Set ADO_ORG, ADO_PROJECT, ADO_PAT env vars.")
    token = f":{ADO_PAT}"  # PAT used as password with empty username
    b64 = base64.b64encode(token.encode("utf-8")).decode("utf-8")
    return {
        "Authorization": f"Basic {b64}",
        "Content-Type": "application/json",
    }

def _ado_base() -> str:
    if ADO_PROJECT:
        return f"https://dev.azure.com/{ADO_ORG}/{ADO_PROJECT}"
    return f"https://dev.azure.com/{ADO_ORG}"


def _ado_project_base(project: Optional[str] = None) -> str:
    project_name = (project or ADO_PROJECT or "").strip()
    if project_name:
        return f"https://dev.azure.com/{ADO_ORG}/{project_name}"
    return f"https://dev.azure.com/{ADO_ORG}"


def _ado_query_api_path(query_path: str) -> str:
    parts = [quote(part.strip(), safe="") for part in str(query_path or "").split("/") if part.strip()]
    if not parts:
        raise RuntimeError("Azure DevOps query path is empty.")
    return "/".join(parts)


def _ado_support_ticket_query_definitions() -> List[Tuple[str, str]]:
    targets: List[Tuple[str, str]] = []
    for raw_target in str(ADO_SUPPORT_TICKET_QUERY_TARGETS or "").split("|"):
        target = raw_target.strip()
        if not target:
            continue
        if "::" not in target:
            print(f"[ADO][TICKETS][WARN] Ignoring malformed query target '{target}'. Expected Project::Query Path.")
            continue
        project, query_path = target.split("::", 1)
        project = project.strip()
        query_path = query_path.strip()
        if project and query_path:
            targets.append((project, query_path))
    if not targets:
        default_project = (ADO_PROJECT_TICKET_SUMMARY or ADO_PROJECT or "").strip()
        if default_project:
            targets.append((default_project, "My Queries/external tickets"))
    return targets


def _ado_keyword_ticket_query_definitions() -> List[Tuple[str, str]]:
    targets: List[Tuple[str, str]] = []
    for raw_target in str(ADO_KEYWORD_TICKET_QUERY_TARGETS or "").split("|"):
        target = raw_target.strip()
        if not target:
            continue
        if "::" not in target:
            print(f"[ADO][TICKETS][WARN] Ignoring malformed keyword query target '{target}'. Expected Project::Query Path.")
            continue
        project, query_path = target.split("::", 1)
        project = project.strip()
        query_path = query_path.strip()
        if project and query_path:
            targets.append((project, query_path))
    return targets


def _ado_support_ticket_fallback_wiql(project: str) -> str:
    project_name = str(project or "").replace("'", "''")
    assigned_to = str(ADO_SUPPORT_TICKET_ASSIGNED_TO or "").replace("'", "''")
    query = (
        "SELECT [System.Id] FROM WorkItems "
        f"WHERE [System.TeamProject] = '{project_name}' "
        "AND [System.State] <> 'Closed' "
    )
    if assigned_to:
        query += f"AND [System.AssignedTo] = '{assigned_to}' "
    query += "ORDER BY [System.ChangedDate] DESC"
    return query


def _ado_fetch_saved_query_wiql(project: str, query_path: str) -> str:
    url = (
        f"{_ado_project_base(project)}/_apis/wit/queries/{_ado_query_api_path(query_path)}"
        "?$expand=wiql&api-version=7.1"
    )
    resp = requests.get(url, headers=_ado_headers(), timeout=60)
    if resp.status_code != 200:
        raise RuntimeError(f"saved query fetch failed: {resp.status_code} {resp.text[:300]}")
    payload = resp.json()
    wiql = str(payload.get("wiql") or "").strip()
    if not wiql:
        raise RuntimeError(f"saved query '{query_path}' in project '{project}' did not return WIQL text")
    return wiql


def _ado_execute_wiql(project: str, wiql: str, label: str) -> List[int]:
    url = f"{_ado_project_base(project)}/_apis/wit/wiql?api-version=7.1-preview.2&$top=20000"
    try:
        resp = requests.post(url, json={"query": wiql}, headers=_ado_headers(), timeout=8)
    except Exception as e:
        print(f"[ADO][TICKETS][ERROR] Query '{label}' failed during execution: {e}")
        traceback.print_exc()
        raise RuntimeError(f"ADO WIQL failed: {e}")
    if resp.status_code != 200:
        print(f"[ADO][TICKETS][ERROR] Query '{label}' returned {resp.status_code}: {resp.text[:300]}")
        raise RuntimeError(f"ADO WIQL error: {resp.text}")
    items = resp.json().get("workItems", [])
    ids = [int(it.get("id")) for it in items if it.get("id")]
    print(f"[ADO][TICKETS] Query '{label}' returned {len(ids)} work items")
    return ids


def _ado_fetch_work_items_by_ids(ids: List[int]) -> List[dict]:
    if not ids:
        return []

    results: List[dict] = []
    org_base = _ado_project_base(None)
    for start in range(0, len(ids), 200):
        batch_ids = ids[start:start + 200]
        det_url = (
            f"{org_base}/_apis/wit/workitems?ids={','.join(str(item_id) for item_id in batch_ids)}"
            "&$expand=all&api-version=7.1"
        )
        det = requests.get(det_url, headers=_ado_headers(), timeout=60)
        if det.status_code != 200:
            print(f"[ADO][TICKETS][ERROR] Work item detail fetch returned {det.status_code}: {det.text[:300]}")
            raise RuntimeError(f"ADO workitems details error: {det.text}")
        results.extend(det.json().get("value", []))
    return results


def _extract_images_from_html(html: str) -> List[str]:
    """
    Extract image URLs/references from HTML.
    Returns list of image src attributes.
    """
    if not html:
        return []
    try:
        img_pattern = r'<img[^>]+src=["\']?([^"\'>\s]+)["\']?'
        matches = re.findall(img_pattern, html or "")
        return [m for m in matches if m.strip()]
    except Exception:
        return []


def _normalize_ado_resource_url(resource_url: str, project: Optional[str] = None) -> str:
    raw_url = str(resource_url or "").strip()
    if not raw_url:
        return ""
    if raw_url.startswith("data:"):
        return raw_url
    if raw_url.startswith("//"):
        return "https:" + raw_url
    if raw_url.startswith("http://") or raw_url.startswith("https://"):
        return raw_url
    if raw_url.startswith("/_apis/") or raw_url.startswith("/_attachments/"):
        return f"{_ado_project_base(project)}{raw_url}"
    if raw_url.startswith("_apis/") or raw_url.startswith("_attachments/"):
        return f"{_ado_project_base(project)}/{raw_url}"
    if raw_url.startswith("/"):
        return f"{_ado_project_base(project)}{raw_url}"
    return raw_url


def _ado_guess_filename(resource_url: str, fallback_name: str = "attachment") -> str:
    try:
        parsed = urlsplit(resource_url)
        query = dict(parse_qsl(parsed.query, keep_blank_values=True))
        for key in ("fileName", "filename", "name"):
            value = (query.get(key) or "").strip()
            if value:
                return unquote(value)
        tail = pathlib.Path(unquote(parsed.path)).name
        if tail:
            return tail
    except Exception:
        pass
    return fallback_name


def _ado_with_download_flag(resource_url: str) -> str:
    parsed = urlsplit(resource_url)
    query = dict(parse_qsl(parsed.query, keep_blank_values=True))
    if "download" not in query:
        query["download"] = "true"
    return urlunsplit((parsed.scheme, parsed.netloc, parsed.path, urlencode(query), parsed.fragment))


def _download_ado_binary_resource(
    resource_url: str,
    project: Optional[str] = None,
    file_name: Optional[str] = None,
) -> Optional[dict]:
    normalized_url = _normalize_ado_resource_url(resource_url, project=project)
    if not normalized_url:
        return None

    if normalized_url.startswith("data:"):
        normalized_b64, payload_mime = _normalize_image_base64_payload(normalized_url)
        if not normalized_b64:
            return None
        try:
            decoded_bytes = base64.b64decode(normalized_b64)
        except Exception:
            return None
        resolved_name = file_name or "embedded-image"
        return {
            "file_name": resolved_name,
            "content_type": payload_mime or mimetypes.guess_type(resolved_name)[0] or "application/octet-stream",
            "data": decoded_bytes,
            "url": normalized_url,
        }

    urls_to_try = [normalized_url]
    download_url = _ado_with_download_flag(normalized_url)
    if download_url != normalized_url:
        urls_to_try.append(download_url)

    last_error: Optional[str] = None
    while urls_to_try:
        candidate_url = urls_to_try.pop(0)
        try:
            response = requests.get(candidate_url, headers=_ado_headers(), timeout=120, allow_redirects=True)
        except Exception as exc:
            last_error = str(exc)
            continue

        if response.status_code != 200:
            last_error = f"{response.status_code} {response.text[:200]}"
            continue

        content_type = (response.headers.get("content-type") or "").split(";")[0].strip().lower()
        if content_type == "application/json":
            try:
                payload = response.json()
            except Exception:
                payload = {}
            follow_url = str(payload.get("downloadUrl") or payload.get("url") or "").strip()
            if follow_url and follow_url not in urls_to_try:
                urls_to_try.append(follow_url)
                continue

        resolved_name = file_name or _ado_guess_filename(response.url or candidate_url)
        return {
            "file_name": resolved_name,
            "content_type": content_type or mimetypes.guess_type(resolved_name)[0] or "application/octet-stream",
            "data": response.content,
            "url": response.url or candidate_url,
        }

    print(f"[ADO][TICKET][WARN] Failed to download Azure DevOps resource '{normalized_url}': {last_error}")
    return None


def _describe_ticket_image_reference(image_ref: str, source_label: str, project: Optional[str] = None) -> str:
    normalized_ref = _normalize_ado_resource_url(image_ref, project=project)
    if not normalized_ref:
        return f"[Image Content from {source_label}]\n[Image source was empty]"

    downloaded = _download_ado_binary_resource(normalized_ref, project=project, file_name=source_label)
    if downloaded:
        resolved_name = downloaded.get("file_name") or source_label
        content_type = str(downloaded.get("content_type") or "")
        if content_type.startswith("image/") or resolved_name.split(".")[-1].lower() in POPULAR_IMAGE_EXTS:
            file_b64 = base64.b64encode(downloaded["data"]).decode("utf-8")
            return _process_image_file(
                resolved_name,
                file_b64,
                content_type,
                describe_func=_describe_image_via_vision,
            )

    try:
        description = _describe_image_via_vision(normalized_ref)
    except Exception as exc:
        description = f"[Image analysis failed: {exc}]"
    return f"[Image Content from {source_label}]\n{description}"


def _render_ado_html_with_inline_images(html: str, section_label: str, project: Optional[str] = None, skip_image_analysis: bool = False) -> str:
    raw_html = str(html or "")
    if not raw_html.strip():
        return ""

    parts: List[str] = []
    cursor = 0
    image_index = 0
    for match in re.finditer(r'<img[^>]+src=["\']?([^"\'>\s]+)["\']?[^>]*>', raw_html, flags=re.IGNORECASE):
        text_segment = _strip_html(raw_html[cursor:match.start()]).strip()
        if text_segment:
            parts.append(text_segment)

        image_index += 1
        if skip_image_analysis:
            parts.append(f"[Image {image_index} in {section_label}]")
        else:
            parts.append(
                _describe_ticket_image_reference(
                    match.group(1),
                    source_label=f"{section_label} image {image_index}",
                    project=project,
                )
            )
        cursor = match.end()

    tail_text = _strip_html(raw_html[cursor:]).strip()
    if tail_text:
        parts.append(tail_text)

    if not parts:
        return _strip_html(raw_html).strip()
    return "\n\n".join(part for part in parts if part)


def _build_ticket_attachment_context(
    work_item_id: int,
    attachments: List[dict],
    project: Optional[str] = None,
    skip_image_analysis: bool = False,
) -> str:
    if not attachments:
        return ""

    ordered_attachments = sorted(
        attachments,
        key=lambda item: (item.get("createdDate") or "", item.get("name") or ""),
    )
    blocks: List[str] = []
    total = len(ordered_attachments)

    for index, attachment in enumerate(ordered_attachments, start=1):
        attachment_name = str(attachment.get("name") or f"attachment-{index}")
        print(f"[ADO][TICKET] Downloading attachment {index}/{total} for work item {work_item_id}: {attachment_name}")
        downloaded = _download_ado_binary_resource(
            str(attachment.get("url") or ""),
            project=project,
            file_name=attachment_name,
        )

        header_lines = [f"Attachment {index}: {attachment_name}"]
        created_date = str(attachment.get("createdDate") or "").strip()
        comment = str(attachment.get("comment") or "").strip()
        if created_date:
            header_lines.append(f"Added: {created_date}")
        if comment:
            header_lines.append(f"Note: {comment}")

        if not downloaded:
            header_lines.append("[Unable to download attachment content from Azure DevOps]")
            blocks.append("\n".join(header_lines))
            continue

        upload_payload = {
            "name": str(downloaded.get("file_name") or attachment_name),
            "type": str(downloaded.get("content_type") or "application/octet-stream"),
            "data": base64.b64encode(downloaded["data"]).decode("utf-8"),
        }
        img_func = None if skip_image_analysis else _describe_image_via_vision
        extracted_text = process_uploaded_file(
            upload_payload,
            describe_image_func=img_func,
        ).strip()
        if extracted_text:
            header_lines.append(extracted_text)
        else:
            header_lines.append("[Attachment was downloaded but produced no extractable content]")
        blocks.append("\n".join(header_lines))

    return "\n\n---\n\n".join(blocks)


def _describe_image_via_vision(image_url: str) -> str:
    """
    Use OpenAI vision (gpt-4o) to describe an image in detail.
    Returns a comprehensive text description focusing on technical content, UI, data, and context.
    Supports both HTTP(S) URLs and data URIs (base64 encoded images).
    """
    try:
        # Accept both URLs and data URIs
        if not (image_url.startswith("http://") or image_url.startswith("https://") or image_url.startswith("data:image/")):
            print(f"[VISION][WARN] Invalid image URL format: {image_url[:60]}...")
            return f"[Local image: {image_url}]"
        
        headers = {"Authorization": f"Bearer {OPENAI_API_KEY}", "Content-Type": "application/json"}
        body = {
            "model": OPENAI_VISION_MODEL,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                "You are a detailed technical image analyzer. Provide an exhaustive description of this image. Include:\n"
                                "1. What type of document/UI/diagram is this (screenshot, chart, diagram, form, etc.)\n"
                                "2. Main content, purpose, and overall layout\n"
                                "3. ALL visible text — every label, header, field name, button, menu item, tooltip, breadcrumb, tab name, and status message. Transcribe them exactly as shown.\n"
                                "4. Technical details: configuration settings, error messages, status indicators, selected values, toggled options\n"
                                "5. ALL tables, data grids, metrics, numerical values — reproduce table content row by row\n"
                                "6. UI elements: panels, sections, checkboxes (checked/unchecked), dropdowns (selected value), input fields (current value), radio buttons, toggles\n"
                                "7. Navigation context: which page/screen/tab is active, breadcrumb path, sidebar selections\n"
                                "8. Any highlighted, selected, or focused elements\n"
                                "9. Color-coded indicators: red/green/yellow status, error highlights, warnings\n\n"
                                "Be exhaustive. Extract EVERY piece of visible information. Do not summarize — list everything you see. "
                                "Do not refuse to analyze. Focus on technical relevance."
                            ),
                        },
                        {"type": "image_url", "image_url": {"url": image_url, "detail": "high"}},
                    ],
                }
            ],
            "max_tokens": 2000,
        }
        resp = requests.post(f"{OPENAI_URL}/chat/completions", json=body, headers=headers, timeout=60)
        if resp.status_code == 200:
            data = resp.json()
            description = data.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
            print(f"[VISION] Detailed image analysis ({len(description)} chars) for {image_url[:60]}...")
            return description
        else:
            print(f"[VISION][WARN] Vision call failed for {image_url}: {resp.status_code} {resp.text[:200]}")
            return f"[Image analysis unavailable]"
    except Exception as e:
        print(f"[VISION][WARN] Failed to describe image {image_url[:60]}...: {e}")
        return f"[Image analysis unavailable]"


def _extract_structured_text_from_image(
    image_ref: str,
    source_label: str = "image",
    cancel_check=None,
) -> str:
    """
    Extract OCR-style text from an image using gpt-4o vision and preserve table content as markdown.
    Uses best OpenAI vision model for superior accuracy.
    """
    if not OPENAI_API_KEY:
        return ""

    if cancel_check and cancel_check():
        print(f"[VISION] Skipping OCR extraction for {source_label}: canceled")
        return ""

    try:
        headers = {"Authorization": f"Bearer {OPENAI_API_KEY}", "Content-Type": "application/json"}
        body = {
            "model": OPENAI_VISION_MODEL,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                "You are an advanced OCR and visual content extractor. Extract all useful information from this image comprehensively. "
                                "Follow this exact output structure:\n"
                                "[IMAGE_SUMMARY]\n"
                                "2-4 detailed sentences describing what this screenshot/image shows, its purpose, and context.\n\n"
                                "[VISIBLE_TEXT]\n"
                                "Extract ALL readable text: labels, values, menu names, buttons, error text, headings, configuration keys, field names, etc. "
                                "Preserve the exact text and format. Use one item per line. If none, write: <none>.\n\n"
                                "[TABLES]\n"
                                "If a table/grid/matrix exists, output it in markdown format preserving all row/column values and structure exactly. If none, write: <none>.\n\n"
                                "[DATA_VALUES]\n"
                                "Extract any numbers, metrics, timestamps, thresholds, or measurable data shown. Include context (e.g., 'CPU: 85%').\n\n"
                                "[UI_ELEMENTS]\n"
                                "List all important UI components: buttons, panels, tabs, sections, checkboxes, dropdowns, input fields, dialogs, and their states.\n\n"
                                "[TECHNICAL_SIGNALS]\n"
                                "Identify technical details: screen/page type, selected items, active filters, visible warnings, error codes, status indicators, configuration mode.\n\n"
                                "Rules: Extract everything visible. Do not refuse. Do not invent unseen details. Mark unclear or illegible content as [unclear]. "
                                "This is for technical documentation and troubleshooting support."
                            ),
                        },
                        {"type": "image_url", "image_url": {"url": image_ref}},
                    ],
                }
            ],
            "max_tokens": min(VISION_MAX_TOKENS, 2500),
            "temperature": 0,
        }
        for attempt in range(VISION_RETRY_COUNT + 1):
            if cancel_check and cancel_check():
                print(f"[VISION] Stopping OCR extraction for {source_label}: canceled")
                return ""
            try:
                resp = requests.post(
                    f"{OPENAI_URL}/chat/completions",
                    json=body,
                    headers=headers,
                    timeout=VISION_TIMEOUT_SECONDS,
                )
            except requests.exceptions.ReadTimeout as e:
                if attempt >= VISION_RETRY_COUNT:
                    print(f"[VISION][WARN] OCR extraction timeout for {source_label} after retries: {e}")
                    return ""
                wait_sec = (attempt + 1) * 2
                print(f"[VISION][WARN] OCR timeout for {source_label}, retry {attempt + 1}/{VISION_RETRY_COUNT} in {wait_sec}s")
                time.sleep(wait_sec)
                continue
            except Exception as e:
                print(f"[VISION][WARN] OCR extraction error for {source_label}: {e}")
                return ""

            if resp.status_code != 200:
                print(f"[VISION][WARN] OCR extraction failed for {source_label}: {resp.status_code} {resp.text[:200]}")
                return ""

            data = resp.json()
            text = data.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
            if text:
                print(f"[VISION] OCR extraction success for {source_label} using {OPENAI_VISION_MODEL}, len={len(text)}")
            return text
        return ""
    except Exception as e:
        print(f"[VISION][WARN] OCR extraction error for {source_label}: {e}")
        return ""


def _normalize_image_base64_payload(payload: str) -> Tuple[Optional[str], Optional[str]]:
    """Return (normalized_b64, mime_type_if_present) from either raw base64 or data URI."""
    raw = (payload or "").strip()
    if not raw:
        return None, None

    mime_type = None
    if raw.startswith("data:") and "," in raw:
        header, raw = raw.split(",", 1)
        match = re.match(r"^data:([^;]+);base64$", header, re.IGNORECASE)
        if match:
            mime_type = (match.group(1) or "").lower()

    # Remove whitespace/newlines and normalize padding.
    raw = re.sub(r"\s+", "", raw)
    if not raw:
        return None, mime_type
    padding = len(raw) % 4
    if padding:
        raw += "=" * (4 - padding)

    return raw, mime_type


def _is_low_signal_image_context(text: str) -> bool:
    candidate = (text or "").strip().lower()
    if not candidate:
        return True

    low_signal_markers = [
        "unable to extract text",
        "cannot extract text",
        "can't extract text",
        "no readable text",
        "image analysis unavailable",
        "i cannot",
        "i can't",
        "<none>",
    ]
    if any(marker in candidate for marker in low_signal_markers):
        return True

    # Very short outputs are typically not enough for troubleshooting context.
    return len(candidate) < 80


def _normalize_table_rows(rows: List[List[Any]]) -> List[List[str]]:
    normalized: List[List[str]] = []
    for row in rows:
        row_values: List[str] = []
        for cell in (row or []):
            cell_text = "" if cell is None else str(cell)
            cell_text = " ".join(cell_text.replace("\r", " ").split())
            cell_text = cell_text.replace("|", "\\|")
            row_values.append(cell_text)
        normalized.append(row_values)
    return normalized


def _format_table_rows_as_markdown(rows: List[List[Any]]) -> str:
    if not rows:
        return ""

    normalized = _normalize_table_rows(rows)

    if not normalized:
        return ""

    width = max(len(r) for r in normalized)
    if width == 0:
        return ""

    padded = [r + [""] * (width - len(r)) for r in normalized]
    header = padded[0]
    separator = ["---"] * width

    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]

    for row in padded[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def _format_table_rows_as_blocks(rows: List[List[Any]], page_number: int, table_number: int) -> List[str]:
    normalized = _normalize_table_rows(rows)
    if len(normalized) <= 1:
        return []

    width = max(len(r) for r in normalized)
    if width == 0:
        return []

    padded = [r + [""] * (width - len(r)) for r in normalized]
    header = padded[0]
    row_blocks: List[str] = []

    for row_idx, row_values in enumerate(padded[1:], start=1):
        row_dict = {}
        for col_idx, cell_val in enumerate(row_values, start=1):
            key = header[col_idx - 1].strip() or f"col_{col_idx}"
            row_dict[key] = cell_val

        row_payload = {
            "page": page_number,
            "table": table_number,
            "row": row_idx,
            "cells": row_values,
            "by_header": row_dict,
        }
        row_blocks.append(
            f"{TABLE_ROW_START_TAG} page={page_number} table={table_number} row={row_idx}\n"
            + json.dumps(row_payload, ensure_ascii=False)
            + f"\n{TABLE_ROW_END_TAG}"
        )

    return row_blocks


def _extract_text_from_pypdf2_page(page: Any) -> str:
    try:
        text = page.extract_text(extraction_mode="layout")
        if text:
            return text
    except TypeError:
        pass
    except Exception:
        pass

    try:
        return page.extract_text() or ""
    except Exception:
        return ""


def _looks_like_command_or_code_line(line: str) -> bool:
    candidate = str(line or "").strip()
    if not candidate:
        return False

    upper_candidate = candidate.upper()
    if upper_candidate in {"SQL", "BASH", "SHELL", "CMD", "POWERSHELL", "PYTHON"}:
        return True

    sql_markers = (
        "CREATE TABLE", "ALTER TABLE", "DROP TABLE", "INSERT INTO", "UPDATE ", "DELETE FROM",
        "SELECT ", "FROM ", "WHERE ", "LEFT JOIN", "RIGHT JOIN", "INNER JOIN", "PRIMARY KEY",
        "FOREIGN KEY", "UNIQUE KEY", "ENGINE=", "COLLATE=", "DEFAULT CHARSET=", "NOT NULL",
    )
    if any(marker in upper_candidate for marker in sql_markers):
        return True

    shell_markers = ("sudo ", "docker ", "kubectl ", "python ", "pip ", "npm ", "yarn ", "git ", "curl ", "wget ")
    if any(candidate.startswith(marker) for marker in shell_markers):
        return True

    if re.search(r'[`{}();=<>"]', candidate):
        return True
    # Lines starting with # that contain UI/navigation language are NOT code
    if re.match(r"^\s*#", candidate):
        _comment_body = re.sub(r"^\s*#\s*", "", candidate)
        if re.search(
            r"\b(click|select|open|navigate|button|tab|screen|ui|portainer|admin studio|"
            r"field|dropdown|checkbox|page|section|restart|restarting|service|clearing|"
            r"browser|cache|setting|menu|dialog|typically|usually|involves|done through|via)"
            r"\b", _comment_body, re.IGNORECASE
        ):
            return False
        return True
    if re.match(r"^\s*(--|/\*|\*)", candidate):
        return True
    if re.match(r"^[A-Za-z0-9_.-]+\s*=\s*.+$", candidate):
        return True
    return False


def _normalize_pdf_extracted_text(text: str) -> str:
    raw_lines = [line.rstrip() for line in str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    normalized_lines: List[str] = []
    paragraph_parts: List[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_parts
        if paragraph_parts:
            paragraph = " ".join(part for part in paragraph_parts if part).strip()
            if paragraph:
                normalized_lines.append(paragraph)
            paragraph_parts = []

    for raw_line in raw_lines:
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            if normalized_lines and normalized_lines[-1] != "":
                normalized_lines.append("")
            continue

        heading_like = bool(re.match(r"^(?:---\s*Page\s+\d+\s*---|\[Table\s+\d+\]|Table:\s|\d+\.\s+|#+\s+)", line))
        code_like = _looks_like_command_or_code_line(line)

        if heading_like or code_like:
            flush_paragraph()
            normalized_lines.append(line)
            continue

        if paragraph_parts and paragraph_parts[-1].endswith("-"):
            paragraph_parts[-1] = paragraph_parts[-1][:-1] + line
        else:
            paragraph_parts.append(line)

    flush_paragraph()

    collapsed: List[str] = []
    previous_blank = False
    for line in normalized_lines:
        if line == "":
            if not previous_blank:
                collapsed.append(line)
            previous_blank = True
        else:
            collapsed.append(line)
            previous_blank = False

    return "\n".join(collapsed).strip()


def _extract_text_from_pymupdf_page(page: Any) -> str:
    try:
        blocks = page.get_text("blocks", sort=True) or []
    except Exception:
        try:
            return _normalize_pdf_extracted_text(page.get_text("text", sort=True) or "")
        except Exception:
            return ""

    block_texts: List[str] = []
    for block in blocks:
        text = str(block[4] or "").strip() if len(block) > 4 else ""
        if not text:
            continue
        normalized = _normalize_pdf_extracted_text(text)
        if normalized:
            block_texts.append(normalized)

    return "\n\n".join(block_texts).strip()


def _extract_layout_ordered_segments_from_pymupdf_page(
    page: Any,
    file_name: str,
    page_number: int,
    allow_image_analysis: bool,
    cancel_check=None,
    request_id: Optional[str] = None,
) -> List[str]:
    """Extract page content in visual reading order (text and images interleaved)."""
    try:
        data = page.get_text("dict", sort=True) or {}
        blocks = data.get("blocks", []) or []
    except Exception as e:
        print(f"[FILE][PDF][WARN] Layout extraction failed page={page_number}: {e}")
        return []

    segments: List[str] = []
    image_count = 0
    page_rect = getattr(page, "rect", None)
    page_area = 0.0
    if page_rect is not None:
        page_area = max(1.0, float(page_rect.width) * float(page_rect.height))

    for block_idx, block in enumerate(blocks):
        if cancel_check and cancel_check():
            _ingest_log(f"pdf page={page_number}: layout extraction canceled", request_id, force=True)
            break
        block_type = int(block.get("type", 0))
        bbox = block.get("bbox") or []
        if DETAILED_INGEST_LOGS:
            kind = "text" if block_type == 0 else ("image" if block_type == 1 else f"type-{block_type}")
            _ingest_log(
                f"pdf page={page_number}: scanning block={block_idx + 1}/{len(blocks)} kind={kind} bbox={bbox if len(bbox) == 4 else 'n/a'}",
                request_id,
            )

        if block_type == 0:
            text_parts: List[str] = []
            for line in block.get("lines", []) or []:
                spans = line.get("spans", []) or []
                line_text = "".join(str(span.get("text", "")) for span in spans).strip()
                if line_text:
                    text_parts.append(line_text)
            text = "\n".join(text_parts).strip()
            if text:
                normalized = _normalize_pdf_extracted_text(text)
                if normalized:
                    segments.append(normalized)
                    _ingest_log(
                        f"pdf page={page_number}: text block accepted len={len(normalized)} -> queued for embedding",
                        request_id,
                    )
            continue

        if block_type != 1 or not allow_image_analysis:
            continue

        if PDF_MAX_IMAGES_PER_PAGE > 0 and image_count >= PDF_MAX_IMAGES_PER_PAGE:
            continue

        bbox = block.get("bbox") or []
        if len(bbox) != 4:
            continue

        x0, y0, x1, y1 = [float(v) for v in bbox]
        width = max(0.0, x1 - x0)
        height = max(0.0, y1 - y0)
        block_area = width * height
        area_ratio = block_area / page_area if page_area else 0.0

        if area_ratio < PDF_MIN_IMAGE_AREA_RATIO:
            _ingest_log(
                f"pdf page={page_number}: image block skipped ratio={area_ratio:.4f} (< {PDF_MIN_IMAGE_AREA_RATIO:.4f})",
                request_id,
            )
            continue

        try:
            clip = fitz.Rect(x0, y0, x1, y1)
            pix = page.get_pixmap(clip=clip, dpi=200, alpha=False)
            image_b64 = base64.b64encode(pix.tobytes("png")).decode("utf-8")
            data_uri = f"data:image/png;base64,{image_b64}"
            image_text = _extract_structured_text_from_image(
                data_uri,
                source_label=f"{file_name}:page-{page_number}:image-{block_idx + 1}",
                cancel_check=cancel_check,
            )
            if image_text:
                image_count += 1
                segments.append("[IMAGE_BLOCK_ANALYSIS]\n" + image_text)
                _ingest_log(
                    f"pdf page={page_number}: image block analyzed with vision block={block_idx + 1} ratio={area_ratio:.3f} desc_len={len(image_text)} -> queued for embedding",
                    request_id,
                )
        except Exception as e:
            print(
                f"[FILE][PDF][WARN] Inline image block extraction failed page={page_number} block={block_idx + 1}: {e}"
            )

    return segments


def _strip_html(text: str) -> str:
    try:
        txt = re.sub(r"<[^>]+>", " ", text or "")
        txt = unescape(txt)
        return " ".join(txt.split())
    except Exception:
        return text or ""


# ---------------------------------------------------------------------
# File Processing Functions (for uploaded files in chat)
# ---------------------------------------------------------------------

POPULAR_IMAGE_EXTS = {
    "png", "jpg", "jpeg", "gif", "webp", "bmp", "tif", "tiff", "svg", "ico", "heic", "heif", "avif", "jfif"
}
POPULAR_EXCEL_EXTS = {"xlsx", "xls", "xlsm", "xltx", "xltm", "xlsb", "ods"}
POPULAR_TEXT_EXTS = {
    "txt", "md", "markdown", "rst", "csv", "tsv", "log", "ini", "cfg", "conf", "properties", "env",
    "json", "jsonl", "xml", "yaml", "yml", "toml", "sql",
    "sh", "bash", "zsh", "ps1", "bat", "cmd",
    "py", "js", "ts", "jsx", "tsx", "java", "cs", "go", "rb", "php", "c", "cpp", "h", "hpp",
    "html", "htm", "css", "scss", "less",
}
POPULAR_OTHER_EXTS = {
    "ppt", "pptx", "odp", "odt", "rtf", "eml", "msg",
    "zip", "rar", "7z", "tar", "gz",
    "mp3", "wav", "m4a", "aac", "ogg", "flac",
    "mp4", "mkv", "mov", "avi", "wmv", "webm",
}


def _process_unstructured_popular_file(file_name: str, file_ext: str) -> str:
    """Graceful fallback for popular file types that are uploadable but not text-extractable yet."""
    return (
        f"**File: {file_name}**\n\n"
        f"[NOTICE] .{file_ext} is accepted for upload, but automatic text extraction is not available yet for this type."
    )

def process_uploaded_file(
    file_data: Dict[str, str],
    describe_image_func=None,
    cancel_check=None,
    request_id: Optional[str] = None,
) -> str:
    """
    Process an uploaded file (base64 encoded) and extract its content.
    Supports: images (via OCR/vision), Excel, CSV, PDF, text, Word docs.
    
    Args:
        file_data: Dict with keys 'name', 'type', 'data' (base64 encoded)
        describe_image_func: Optional function to describe images via Vision API
    
    Returns:
        Extracted text content
    """
    try:
        if cancel_check and cancel_check():
            _ingest_log("file processing skipped due to cancel before start", request_id, force=True)
            return "[CANCELED] Upload processing stopped by user."

        file_name = file_data.get("name", "unknown")
        file_type = file_data.get("type", "")
        file_b64 = file_data.get("data", "")
        
        if not file_b64:
            return f"[ERROR] Empty file: {file_name}"
        
        file_ext = file_name.split(".")[-1].lower()
        
        # Decode base64
        try:
            file_bytes = base64.b64decode(file_b64)
        except Exception as e:
            print(f"[FILE] Failed to decode base64 for {file_name}: {e}")
            return f"[ERROR] Failed to decode file: {file_name}"
        
        print(f"[FILE] Processing {file_name} (ext={file_ext}, size={len(file_bytes)} bytes)")
        _ingest_log(f"file begin name='{file_name}' ext={file_ext} size_bytes={len(file_bytes)}", request_id, force=True)
        
        # Route to appropriate handler
        if file_ext in POPULAR_IMAGE_EXTS or (file_type and file_type.startswith("image/")):
            return _process_image_file(file_name, file_b64, file_type, describe_image_func)
        elif file_ext in POPULAR_EXCEL_EXTS:
            return _process_excel_file(file_name, file_bytes)
        elif file_ext == "csv":
            return _process_csv_file(file_name, file_bytes)
        elif file_ext == "pdf":
            return _process_pdf_file(file_name, file_bytes, cancel_check=cancel_check, request_id=request_id)
        elif file_ext in POPULAR_TEXT_EXTS:
            return _process_text_file(file_name, file_bytes)
        elif file_ext in ["doc", "docx"]:
            return _process_word_file(file_name, file_bytes)
        elif file_ext == "zip":
            return _process_archive_file(file_name, file_bytes, describe_image_func=describe_image_func, cancel_check=cancel_check, request_id=request_id)
        elif any(file_name.lower().endswith(ext) for ext in (".tar.gz", ".tgz", ".tar.bz2", ".tar")) or file_ext == "tar":
            return _process_archive_file(file_name, file_bytes, describe_image_func=describe_image_func, cancel_check=cancel_check, request_id=request_id)
        elif file_ext in POPULAR_OTHER_EXTS:
            return _process_unstructured_popular_file(file_name, file_ext)
        else:
            return f"[WARNING] Unsupported file type: {file_name} (.{file_ext})"
    
    except Exception as e:
        print(f"[FILE] Error processing file: {e}")
        return f"[ERROR] Failed to process file: {str(e)}"


def _process_image_file(file_name: str, file_b64: str, file_type: Optional[str] = None, describe_func=None) -> str:
    """Process image file using OpenAI Vision API or fallback."""
    try:
        if not OPENAI_API_KEY:
            return f"[Image file {file_name} - OPENAI_API_KEY is not configured]"

        normalized_b64, payload_mime = _normalize_image_base64_payload(file_b64)
        if not normalized_b64:
            return f"[Image file {file_name} - invalid base64 data]"

        # Decode base64 to bytes for potential conversion
        try:
            file_bytes = base64.b64decode(normalized_b64)
        except Exception as e:
            print(f"[FILE] Failed to decode base64 for image {file_name}: {e}")
            return f"[Image file {file_name} - invalid base64 data]"

        # Determine mime type
        mime_type = None
        if file_type and file_type.startswith("image/"):
            mime_type = file_type
        if not mime_type and payload_mime and payload_mime.startswith("image/"):
            mime_type = payload_mime
        if not mime_type:
            mime_type = mimetypes.guess_type(file_name)[0]
        if not mime_type:
            mime_type = "image/jpeg"

        # If unsupported by vision, try convert to JPEG via PIL
        vision_supported = {"image/jpeg", "image/png", "image/gif", "image/webp"}
        if mime_type not in vision_supported:
            if HAS_PIL:
                try:
                    img = Image.open(BytesIO(file_bytes))
                    if img.mode not in ("RGB", "L"):
                        img = img.convert("RGB")
                    buf = BytesIO()
                    img.save(buf, format="JPEG")
                    jpeg_b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
                    data_uri = f"data:image/jpeg;base64,{jpeg_b64}"
                except Exception as e:
                    print(f"[FILE] Failed to convert image {file_name} to JPEG: {e}")
                    return f"[Image file {file_name} - unsupported format for analysis]"
            else:
                return f"[Image file {file_name} - unsupported format for analysis]"
        else:
            data_uri = f"data:{mime_type};base64,{normalized_b64}"

        # Run both OCR extraction and visual description in parallel for complete coverage
        _ocr_result = [""]
        _visual_result = [""]

        def _run_ocr():
            _ocr_result[0] = _extract_structured_text_from_image(data_uri, source_label=file_name)

        def _run_visual():
            if describe_func:
                _visual_result[0] = (describe_func(data_uri) or "").strip()

        with ThreadPoolExecutor(max_workers=2) as img_pool:
            img_pool.submit(_run_ocr)
            img_pool.submit(_run_visual)
            img_pool.shutdown(wait=True)

        extracted = _ocr_result[0]
        visual_desc = _visual_result[0]

        merged_parts: List[str] = []
        if extracted and extracted.strip():
            merged_parts.append(extracted.strip())
        if visual_desc:
            # Always include visual description — it captures layout/context OCR may miss
            merged_parts.append("[VISUAL_DESCRIPTION]\n" + visual_desc)

        image_context = "\n\n".join(part for part in merged_parts if part).strip()
        if not image_context:
            image_context = "[Image analysis unavailable]"

        return f"[Image Content from {file_name}]\n{image_context}"
    except Exception as e:
        print(f"[FILE] Vision API failed for {file_name}: {e}")
        return f"[Image file {file_name} - analysis failed: {e}]"


def _process_text_file(file_name: str, file_bytes: bytes) -> str:
    """Extract text from text file."""
    try:
        content = file_bytes.decode("utf-8", errors="ignore")
        return f"**File: {file_name}**\n\n{content}"
    except Exception as e:
        print(f"[FILE] Failed to read text file {file_name}: {e}")
        return f"[ERROR] Failed to read text file: {file_name}"


def _process_excel_file(file_name: str, file_bytes: bytes) -> str:
    """Extract text from Excel file."""
    try:
        if not HAS_OPENPYXL:
            return f"**File: {file_name}** [Excel support not installed - openpyxl required]"
        
        excel_file = BytesIO(file_bytes)
        wb = openpyxl.load_workbook(excel_file)
        
        content_lines = [f"**File: {file_name}**\n"]
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            content_lines.append(f"\n**Sheet: {sheet_name}**\n")
            
            row_count = 0
            for row in ws.iter_rows(max_row=1000):
                row_data = []
                for cell in row:
                    val = cell.value
                    row_data.append(str(val) if val is not None else "")
                
                if any(row_data):
                    content_lines.append(" | ".join(row_data))
                    row_count += 1
            
            if row_count == 1000:
                content_lines.append("\n... (file too large, showing first 1000 rows)")
        
        return "\n".join(content_lines)
    
    except Exception as e:
        print(f"[FILE] Failed to process Excel file {file_name}: {e}")
        return f"**File: {file_name}** [Error reading Excel: {str(e)}]"


def _process_csv_file(file_name: str, file_bytes: bytes) -> str:
    """Extract text from CSV file."""
    try:
        csv_text = file_bytes.decode("utf-8", errors="ignore")
        reader = csv.reader(csv_text.splitlines())
        
        content_lines = [f"**File: {file_name}**\n"]
        
        row_count = 0
        for row in reader:
            if row_count >= 500:
                content_lines.append("\n... (file too large, showing first 500 rows)")
                break
            
            if any(row):
                content_lines.append(" | ".join(row))
                row_count += 1
        
        return "\n".join(content_lines)
    
    except Exception as e:
        print(f"[FILE] Failed to process CSV file {file_name}: {e}")
        return f"**File: {file_name}** [Error reading CSV: {str(e)}]"


def _process_pdf_file(file_name: str, file_bytes: bytes, cancel_check=None, request_id: Optional[str] = None) -> str:
    """Extract text from PDF file."""
    try:
        import PyPDF2
        pdf_file = BytesIO(file_bytes)
        reader = PyPDF2.PdfReader(pdf_file)

        total_pages = len(reader.pages)
        pages_to_process = total_pages
        if PDF_MAX_PAGES > 0:
            pages_to_process = min(total_pages, PDF_MAX_PAGES)

        _ingest_log(
            f"pdf begin file='{file_name}' pages_total={total_pages} pages_to_process={pages_to_process}",
            request_id,
            force=True,
        )

        content_lines = [f"**File: {file_name}** ({total_pages} pages)\n"]

        plumber_doc = None
        if HAS_PDFPLUMBER and PDF_TABLE_EXTRACTION_ENABLED:
            try:
                plumber_doc = pdfplumber.open(BytesIO(file_bytes))
            except Exception as e:
                print(f"[FILE][PDF][WARN] pdfplumber open failed for {file_name}: {e}")

        pymupdf_doc = None
        if HAS_PYMUPDF:
            try:
                pymupdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
            except Exception as e:
                print(f"[FILE][PDF][WARN] PyMuPDF open failed for {file_name}: {e}")

        for page_idx in range(pages_to_process):
            if cancel_check and cancel_check():
                _ingest_log(
                    f"cancel acknowledged during PDF parse at page={page_idx + 1}; stopping parse",
                    request_id,
                    force=True,
                )
                content_lines.append("\n[CANCELED] PDF processing stopped by user.")
                break
            page = reader.pages[page_idx]
            page_lines: List[str] = [f"\n--- Page {page_idx + 1} ---"]
            _ingest_log(f"pdf page={page_idx + 1}: start top-to-bottom scan", request_id)
            allow_image_analysis = (
                PDF_IMAGE_EXTRACTION_ENABLED
                and PDF_OCR_ENABLED
                and OPENAI_API_KEY
                and (PDF_OCR_MAX_PAGES == 0 or page_idx < PDF_OCR_MAX_PAGES)
            )

            text = ""
            used_layout_segments = False
            if pymupdf_doc and page_idx < len(pymupdf_doc):
                try:
                    pymupdf_page = pymupdf_doc.load_page(page_idx)
                    ordered_segments = _extract_layout_ordered_segments_from_pymupdf_page(
                        pymupdf_page,
                        file_name=file_name,
                        page_number=page_idx + 1,
                        allow_image_analysis=allow_image_analysis,
                        cancel_check=cancel_check,
                        request_id=request_id,
                    )
                    if ordered_segments:
                        used_layout_segments = True
                        page_lines.append("\n\n".join(ordered_segments))
                        _ingest_log(
                            f"pdf page={page_idx + 1}: ordered segments ready count={len(ordered_segments)}",
                            request_id,
                        )
                    else:
                        text = _extract_text_from_pymupdf_page(pymupdf_page)
                except Exception as e:
                    print(f"[FILE][PDF][WARN] PyMuPDF text extraction failed page={page_idx + 1}: {e}")

            if not used_layout_segments and not text.strip():
                text = _normalize_pdf_extracted_text(_extract_text_from_pypdf2_page(page))

            if text and text.strip():
                page_lines.append(text)
                _ingest_log(
                    f"pdf page={page_idx + 1}: fallback text extracted len={len(text)} -> queued for embedding",
                    request_id,
                )

            table_count = 0
            if plumber_doc and page_idx < len(plumber_doc.pages):
                try:
                    tables = plumber_doc.pages[page_idx].extract_tables() or []
                    for table_idx, table_rows in enumerate(tables, start=1):
                        table_md = _format_table_rows_as_markdown(table_rows)
                        if table_md:
                            table_count += 1
                            page_lines.append(f"\n[Table {table_idx}]\n{table_md}")
                            _ingest_log(
                                f"pdf page={page_idx + 1}: table extracted idx={table_idx} len={len(table_md)} -> queued for embedding",
                                request_id,
                            )
                            row_blocks = _format_table_rows_as_blocks(
                                table_rows,
                                page_number=page_idx + 1,
                                table_number=table_idx,
                            )
                            if row_blocks:
                                page_lines.append("\n" + "\n".join(row_blocks))
                except Exception as e:
                    print(f"[FILE][PDF][WARN] Table extraction failed page={page_idx + 1}: {e}")

            should_try_ocr = (
                pymupdf_doc is not None
                and allow_image_analysis
                and len(page_lines) == 1
            )

            if should_try_ocr:
                try:
                    p = pymupdf_doc.load_page(page_idx)
                    pix = p.get_pixmap(dpi=200, alpha=False)
                    image_bytes = pix.tobytes("png")
                    image_b64 = base64.b64encode(image_bytes).decode("utf-8")
                    data_uri = f"data:image/png;base64,{image_b64}"
                    ocr_text = _extract_structured_text_from_image(
                        data_uri,
                        source_label=f"{file_name}:page-{page_idx + 1}",
                        cancel_check=cancel_check,
                    )
                    if ocr_text:
                        page_lines.append("\n[PAGE_IMAGE_ANALYSIS_FALLBACK]\n" + ocr_text)
                        print(f"[FILE][PDF] Fallback page image analysis extracted from page {page_idx + 1}")
                        _ingest_log(
                            f"pdf page={page_idx + 1}: fallback full-page vision len={len(ocr_text)} -> queued for embedding",
                            request_id,
                        )
                except Exception as e:
                    print(f"[FILE][PDF][WARN] Image extraction failed page={page_idx + 1}: {e}")

            if len(page_lines) == 1:
                page_lines.append("[No extractable text found on this page]")

            content_lines.append("\n".join(page_lines))
            _ingest_log(f"pdf page={page_idx + 1}: complete sections={max(0, len(page_lines)-1)}", request_id)

        if pages_to_process < total_pages:
            content_lines.append(
                f"\n... (showing first {pages_to_process} pages of {total_pages} total; set PDF_MAX_PAGES=0 for all pages)"
            )

        if plumber_doc:
            try:
                plumber_doc.close()
            except Exception:
                pass
        if pymupdf_doc:
            try:
                pymupdf_doc.close()
            except Exception:
                pass
        
        return "\n".join(content_lines)
    
    except ImportError:
        print(f"[FILE] PyPDF2 not installed, cannot process PDF: {file_name}")
        return f"**File: {file_name}** [PDF support not installed - PyPDF2 required]"
    except Exception as e:
        print(f"[FILE] Failed to process PDF file {file_name}: {e}")
        return f"**File: {file_name}** [Error reading PDF: {str(e)}]"


def _process_word_file(file_name: str, file_bytes: bytes) -> str:
    """Extract text from Word document."""
    try:
        from docx import Document
        docx_file = BytesIO(file_bytes)
        doc = Document(docx_file)
        
        content_lines = [f"**File: {file_name}**\n"]
        
        for para in doc.paragraphs:
            if para.text.strip():
                content_lines.append(para.text)
        
        for table_idx, table in enumerate(doc.tables):
            content_lines.append(f"\n**Table {table_idx + 1}:**")
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                content_lines.append(" | ".join(row_data))
        
        return "\n".join(content_lines)
    
    except ImportError:
        print(f"[FILE] python-docx not installed, cannot process Word doc: {file_name}")
        return f"**File: {file_name}** [Word support not installed - python-docx required]"
    except Exception as e:
        print(f"[FILE] Failed to process Word file {file_name}: {e}")
        return f"**File: {file_name}** [Error reading Word doc: {str(e)}]"


def _process_archive_file(
    file_name: str,
    file_bytes: bytes,
    describe_image_func=None,
    cancel_check=None,
    request_id: Optional[str] = None,
) -> str:
    """Extract and process all files from a zip or tar archive.
    Images inside the archive are passed through the vision model.
    """
    import io
    import zipfile
    import tarfile as tarfile_mod

    MAX_EXTRACTED_BYTES = 200 * 1024 * 1024  # 200 MB total extracted size guard
    MAX_FILES = 50  # max number of entries to process

    file_name_lower = file_name.lower()
    file_ext = file_name.split(".")[-1].lower()
    results: List[str] = []
    total_extracted_bytes = 0

    def _process_entry(entry_name: str, entry_bytes: bytes) -> None:
        nonlocal total_extracted_bytes
        total_extracted_bytes += len(entry_bytes)
        if total_extracted_bytes > MAX_EXTRACTED_BYTES:
            results.append(
                f"[TRUNCATED] Archive size limit of {MAX_EXTRACTED_BYTES // (1024 * 1024)} MB reached. "
                "Remaining files not processed."
            )
            return
        display_name = os.path.basename(entry_name) or entry_name
        entry_b64 = base64.b64encode(entry_bytes).decode("utf-8")
        mime_type, _ = mimetypes.guess_type(display_name)
        entry_data = {"name": display_name, "type": mime_type or "", "data": entry_b64}
        print(f"[ARCHIVE] Processing entry '{display_name}' ({len(entry_bytes)} bytes) from {file_name}")
        content = process_uploaded_file(
            entry_data,
            describe_image_func=describe_image_func,
            cancel_check=cancel_check,
            request_id=request_id,
        )
        results.append(f"[From archive: {file_name} → {display_name}]\n{content}")

    try:
        if file_ext == "zip":
            try:
                with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                    entries = [e for e in zf.infolist() if not e.is_dir()]
                    if not entries:
                        return f"**Archive: {file_name}**\n\n[Archive is empty]"
                    for entry in entries[:MAX_FILES]:
                        if cancel_check and cancel_check():
                            break
                        if total_extracted_bytes > MAX_EXTRACTED_BYTES:
                            break
                        try:
                            entry_bytes = zf.read(entry.filename)
                        except Exception as e:
                            results.append(f"[From {file_name} → {entry.filename}]\n[Error reading entry: {e}]")
                            continue
                        _process_entry(entry.filename, entry_bytes)
            except zipfile.BadZipFile:
                return f"**Archive: {file_name}**\n\n[Invalid or corrupted zip file]"
        else:
            # tar, tar.gz, tgz, tar.bz2 — tarfile auto-detects compression
            try:
                with tarfile_mod.open(fileobj=io.BytesIO(file_bytes)) as tf:
                    members = [m for m in tf.getmembers() if m.isfile()]
                    if not members:
                        return f"**Archive: {file_name}**\n\n[Archive is empty]"
                    for member in members[:MAX_FILES]:
                        if cancel_check and cancel_check():
                            break
                        if total_extracted_bytes > MAX_EXTRACTED_BYTES:
                            break
                        try:
                            f = tf.extractfile(member)
                            if f is None:
                                continue
                            entry_bytes = f.read()
                        except Exception as e:
                            results.append(f"[From {file_name} → {member.name}]\n[Error reading entry: {e}]")
                            continue
                        _process_entry(member.name, entry_bytes)
            except tarfile_mod.TarError as e:
                return f"**Archive: {file_name}**\n\n[Invalid or corrupted archive: {e}]"

        if not results:
            return f"**Archive: {file_name}**\n\n[No processable files found in archive]"

        header = f"**Archive: {file_name}** ({len(results)} file(s) extracted)\n\n"
        return header + "\n\n---ARCHIVE ENTRY SEPARATOR---\n\n".join(results)

    except Exception as e:
        print(f"[FILE] Error processing archive {file_name}: {e}")
        return f"**Archive: {file_name}**\n\n[Error extracting archive: {e}]"


def build_file_context(
    files: Optional[List[Dict[str, str]]],
    describe_image_func=None,
    cancel_check=None,
    request_id: Optional[str] = None,
) -> str:
    """
    Process all uploaded files and build a context string to include in LLM prompt.
    
    Args:
        files: List of file dicts with 'name', 'type', 'data' keys
        describe_image_func: Optional function to describe images
    
    Returns:
        Combined text from all files, or empty string if no files.
    """
    if not files:
        return ""
    
    print(f"[FILE] Processing {len(files)} uploaded files")
    _ingest_log(f"file batch start count={len(files)}", request_id, force=True)
    
    file_contents = []
    for idx, file_data in enumerate(files, start=1):
        if cancel_check and cancel_check():
            _ingest_log(f"cancel acknowledged before file index={idx}; stopping batch", request_id, force=True)
            file_contents.append("[CANCELED] File processing stopped by user.")
            break
        file_name = str((file_data or {}).get("name", "unknown"))
        _ingest_log(f"file batch item {idx}/{len(files)} begin name='{file_name}'", request_id)
        content = process_uploaded_file(
            file_data,
            describe_image_func,
            cancel_check=cancel_check,
            request_id=request_id,
        )
        file_contents.append(content)
        _ingest_log(f"file batch item {idx}/{len(files)} done name='{file_name}' extracted_len={len(content)}", request_id)
    
    combined = "\n\n---FILE SEPARATOR---\n\n".join(file_contents)
    print(f"[FILE] Extracted content from {len(files)} files, total length={len(combined)}")
    _ingest_log(f"file batch complete total_combined_len={len(combined)}", request_id, force=True)
    return combined


# ---------------------------------------------------------------------
# ADO Helpers (continued)
# ---------------------------------------------------------------------



def _first_field(fields: dict, keys: List[str]) -> str:
    for k in keys:
        if k in fields and fields.get(k) not in (None, ""):
            v = fields.get(k)
            if isinstance(v, dict):
                return v.get("displayName") or v.get("uniqueName") or str(v)
            return str(v)
    return ""


def ado_fetch_ticket_full(work_item_id: int) -> dict:
    """Fetch main fields + comments for a work item."""
    base = _ado_project_base(None)
    hdrs = _ado_headers()
    main = requests.get(
        f"{base}/_apis/wit/workitems/{work_item_id}?$expand=all&api-version=7.1",
        headers=hdrs,
        timeout=60,
    )
    if main.status_code != 200:
        raise RuntimeError(f"ADO workitem fetch error: {main.text}")
    data = main.json()
    fields = data.get("fields", {})
    relations = data.get("relations", [])
    project_name = fields.get("System.TeamProject") or ADO_PROJECT

    comments: List[dict] = []
    try:
        c = requests.get(
            f"{_ado_project_base(project_name)}/_apis/wit/workItems/{work_item_id}/comments?api-version=7.1-preview.3",
            headers=hdrs,
            timeout=60,
        )
        if c.status_code == 200:
            for it in c.json().get("comments", []):
                raw_html = it.get("text", "") or ""
                comments.append(
                    {
                        "text": _strip_html(raw_html),
                        "raw_html": raw_html,
                        "createdBy": (it.get("createdBy") or {}).get("displayName")
                        or (it.get("createdBy") or {}).get("uniqueName")
                        or "",
                        "createdDate": it.get("createdDate") or "",
                    }
                )
    except Exception as e:
        print(f"[ADO][WARN] comments fetch failed: {e}")

    attachments: List[dict] = []
    for relation in relations or []:
        if str(relation.get("rel") or "") != "AttachedFile":
            continue
        attributes = relation.get("attributes") or {}
        relation_url = relation.get("url") or ""
        attachment_name = attributes.get("name") or _ado_guess_filename(relation_url)
        attachments.append(
            {
                "name": attachment_name,
                "url": relation_url,
                "comment": attributes.get("comment") or "",
                "createdDate": attributes.get("resourceCreatedDate") or attributes.get("authorizedDate") or "",
            }
        )

    return {
        "id": work_item_id,
        "project": project_name,
        "fields": fields,
        "comments": comments,
        "attachments": attachments,
    }


def ado_list_tickets(tag_contains: str = "CC") -> List[dict]:
    """Load open support tickets from the configured Azure DevOps saved queries."""
    del tag_contains

    query_targets = _ado_support_ticket_query_definitions()
    if not query_targets:
        return []

    merged_ids: List[int] = []
    seen_ids = set()
    for project, query_path in query_targets:
        label = f"{project}::{query_path}"
        try:
            wiql = _ado_fetch_saved_query_wiql(project, query_path)
            ids = _ado_execute_wiql(project, wiql, label)
        except Exception as exc:
            print(
                f"[ADO][TICKETS][WARN] Saved query '{label}' could not be used ({exc}). "
                "Falling back to the same filter shown in the screenshot."
            )
            ids = _ado_execute_wiql(project, _ado_support_ticket_fallback_wiql(project), f"{label} [fallback]")
        for item_id in ids:
            if item_id in seen_ids:
                continue
            seen_ids.add(item_id)
            merged_ids.append(item_id)

    if not merged_ids:
        return []

    results = []
    for wi in _ado_fetch_work_items_by_ids(merged_ids):
        fid = wi.get("id")
        flds = wi.get("fields", {})
        title = flds.get("System.Title", "")
        state = flds.get("System.State", "")
        tags = flds.get("System.Tags", "")
        project_name = flds.get("System.TeamProject", "")
        changed_date = flds.get("System.ChangedDate", "")
        web_url = f"https://dev.azure.com/{ADO_ORG}/{project_name}/_workitems/edit/{fid}"
        results.append(
            {
                "id": fid,
                "title": title,
                "state": state,
                "tags": tags,
                "project": project_name,
                "changedDate": changed_date,
                "url": web_url,
            }
        )

    results.sort(key=lambda item: (item.get("changedDate") or "", item.get("id") or 0), reverse=True)
    return results


def ado_get_ticket_picker_item(work_item_id: int) -> dict:
    """Fetch a single work item and normalize it for the ticket dropdown picker."""
    base = _ado_project_base(None)
    hdrs = _ado_headers()
    resp = requests.get(
        f"{base}/_apis/wit/workitems/{work_item_id}?$expand=fields&api-version=7.1",
        headers=hdrs,
        timeout=60,
    )

    if resp.status_code == 404:
        raise HTTPException(status_code=404, detail=f"Ticket #{work_item_id} was not found in Azure DevOps.")
    if resp.status_code != 200:
        raise RuntimeError(f"ADO workitem lookup error ({resp.status_code}): {resp.text}")

    wi = resp.json() or {}
    fields = wi.get("fields", {}) or {}
    project_name = str(fields.get("System.TeamProject") or ADO_PROJECT or "").strip()
    expected_project = str(ADO_PROJECT or "").strip()

    # Allow tickets from the configured project AND from "TA9 Support".
    allowed_projects = {expected_project.lower()} if expected_project else set()
    allowed_projects.add("ta9 support")

    if allowed_projects and project_name and project_name.lower() not in allowed_projects:
        raise HTTPException(
            status_code=404,
            detail=(
                f"Ticket #{work_item_id} belongs to project '{project_name}'. "
                f"Please provide a ticket from one of: {', '.join(sorted(p.title() for p in allowed_projects if p))}."
            ),
        )

    # For non-primary projects, enforce a 2-year age limit.
    if project_name and expected_project and project_name.lower() != expected_project.lower():
        created_str = str(fields.get("System.CreatedDate") or "").strip()
        if created_str:
            try:
                from datetime import timezone, timedelta
                created_dt = datetime.fromisoformat(created_str.replace("Z", "+00:00"))
                cutoff = datetime.now(timezone.utc) - timedelta(days=730)
                if created_dt < cutoff:
                    raise HTTPException(
                        status_code=404,
                        detail=(
                            f"Ticket #{work_item_id} from '{project_name}' is older than 2 years."
                        ),
                    )
            except HTTPException:
                raise
            except Exception:
                pass  # If date parsing fails, allow it through

    title = str(fields.get("System.Title") or "").strip()
    state = str(fields.get("System.State") or "").strip()
    tags = str(fields.get("System.Tags") or "").strip()
    changed_date = str(fields.get("System.ChangedDate") or "").strip()
    web_url = f"https://dev.azure.com/{ADO_ORG}/{project_name}/_workitems/edit/{work_item_id}"

    return {
        "id": int(wi.get("id") or work_item_id),
        "title": title or f"Work item {work_item_id}",
        "state": state,
        "tags": tags,
        "project": project_name,
        "changedDate": changed_date,
        "url": web_url,
    }


def _normalize_ticket_search_text(value: str) -> str:
    text = unescape(_strip_html(str(value or "")))
    text = text.lower()
    text = re.sub(r"[^a-z0-9\s]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _ticket_search_terms(query: str) -> List[str]:
    normalized = _normalize_ticket_search_text(query)
    if not normalized:
        return []
    stopwords = {
        "a", "an", "and", "are", "as", "at", "be", "by", "do", "for", "from",
        "had", "has", "have", "he", "her", "him", "his", "how", "if", "in",
        "into", "is", "it", "its", "just", "like", "look", "me", "my", "no",
        "nor", "not", "of", "on", "or", "our", "out", "own", "so", "some",
        "than", "that", "the", "them", "then", "there", "these", "they",
        "this", "to", "too", "up", "us", "very", "was", "we", "were", "what",
        "when", "where", "which", "who", "whom", "why", "will", "with", "you",
        "your", "also", "but", "can", "did", "does", "get", "got", "may",
        "much", "must", "need", "see", "should", "such", "would", "could",
        "about", "after", "all", "been", "before", "between", "both",
        "each", "few", "more", "most", "other", "over", "same",
        "through", "under", "until",
    }
    parts = [p for p in normalized.split(" ") if len(p) >= 2 and p not in stopwords]
    seen = set()
    terms: List[str] = []
    for part in parts:
        if part in seen:
            continue
        seen.add(part)
        terms.append(part)
    # Cap at 15 most significant terms — since we combine them into a single
    # Contains Words clause, more terms don't slow down the WIQL query.
    # Prefer longer terms (more specific) when there are too many.
    if len(terms) > 15:
        terms.sort(key=lambda t: -len(t))
        terms = terms[:15]
    return terms


def _ticket_stem(term: str) -> str:
    token = str(term or "")
    for suffix in ("ingly", "edly", "ing", "ed", "es", "s"):
        if len(token) > 4 and token.endswith(suffix):
            return token[: -len(suffix)]
    return token


def _ticket_soft_phrase_similarity(query_phrase: str, haystack: str) -> float:
    """Return best soft similarity between query phrase and local token windows in text."""
    phrase = _normalize_ticket_search_text(query_phrase)
    text = _normalize_ticket_search_text(haystack)
    if not phrase or not text:
        return 0.0

    phrase_tokens = phrase.split(" ")
    text_tokens = text.split(" ")
    if not phrase_tokens or not text_tokens:
        return 0.0

    window_size = max(2, min(len(phrase_tokens) + 2, 10))
    max_windows = 500
    best = 0.0

    if len(text_tokens) <= window_size:
        return difflib.SequenceMatcher(None, phrase, " ".join(text_tokens)).ratio()

    checked = 0
    for index in range(0, len(text_tokens) - window_size + 1):
        if checked >= max_windows:
            break
        window = " ".join(text_tokens[index:index + window_size])
        ratio = difflib.SequenceMatcher(None, phrase, window).ratio()
        if ratio > best:
            best = ratio
            if best >= 0.95:
                break
        checked += 1

    return best


def _ticket_keyword_match_score(query: str, title: str, description: str, discussion: str) -> Tuple[float, int, int]:
    """Return (score, exact_term_hits, fuzzy_term_hits) for title+description+discussion search."""
    phrase = _normalize_ticket_search_text(query)
    if not phrase:
        return 0.0, 0, 0

    terms = _ticket_search_terms(query)
    
    title_norm = _normalize_ticket_search_text(title)
    description_norm = _normalize_ticket_search_text(description)
    discussion_norm = _normalize_ticket_search_text(discussion)
    body_combined = f"{description_norm} {discussion_norm}".strip()
    all_combined = f"{title_norm} {body_combined}".strip()
    
    if not all_combined:
        return 0.0, 0, 0

    # Soft similarity first - this is the primary matching strategy for "data model" type queries
    soft_title_similarity = _ticket_soft_phrase_similarity(phrase, title_norm) if title_norm else 0.0
    soft_body_similarity = _ticket_soft_phrase_similarity(phrase, body_combined) if body_combined else 0.0
    soft_similarity = max(soft_title_similarity, soft_body_similarity)
    
    # If we have decent soft similarity, that's a match
    if soft_similarity >= 0.50:
        return 2.0 + (soft_similarity * 3.0), 0, 0

    # Otherwise try exact/partial/fuzzy matching as secondary strategy
    if not terms:
        # If no terms extracted but phrase is present exactly, score it
        if phrase in all_combined:
            return 3.0, 1, 0
        return 0.0, 0, 0

    phrase_boost = 0.0
    if phrase in body_combined:
        phrase_boost += 2.5
    if phrase in title_norm:
        phrase_boost += 1.2

    words = set(all_combined.split(" "))
    title_words = set(title_norm.split(" ")) if title_norm else set()
    stems = {_ticket_stem(w) for w in words if w}
    exact_hits = 0
    fuzzy_hits = 0
    partial_hits = 0
    total_occurrences = 0

    for term in terms:
        term_stem = _ticket_stem(term)
        body_occurrences = body_combined.count(term)
        title_occurrences = title_norm.count(term)
        occurrences = body_occurrences + title_occurrences
        total_occurrences += occurrences
        # Slightly prefer title matches while keeping body text as primary signal.
        if term in words or occurrences > 0:
            exact_hits += 1
            continue
        if term_stem and (
            term_stem in stems
            or any(w.startswith(term_stem) or term_stem.startswith(w) for w in words if w)
        ):
            partial_hits += 1
            continue
        close = difflib.get_close_matches(term, list(words), n=1, cutoff=0.72)
        if close:
            fuzzy_hits += 1

    # Very permissive threshold - include anything with at least soft similarity or ANY matching signal
    if exact_hits == 0 and fuzzy_hits == 0 and partial_hits == 0 and phrase_boost == 0 and soft_similarity < 0.40:
        return 0.0, 0, 0

    exact_ratio = exact_hits / len(terms) if terms else 0
    fuzzy_ratio = fuzzy_hits / len(terms) if terms else 0
    partial_ratio = partial_hits / len(terms) if terms else 0
    occurrence_boost = min(2.0, total_occurrences * 0.15)
    title_exact_boost = 0.4 * sum(1 for term in terms if term in title_words)
    soft_similarity_boost = max(0.0, soft_similarity - 0.30) * 4.0
    
    score = (
        phrase_boost
        + (exact_ratio * 4.0)
        + (partial_ratio * 1.7)
        + (fuzzy_ratio * 1.2)
        + occurrence_boost
        + title_exact_boost
        + soft_similarity_boost
    )
    
    return score, exact_hits, fuzzy_hits


def _ado_inject_keyword_conditions(base_wiql: str, query: str, use_or: bool = False) -> str:
    """Inject keyword filtering into an existing WIQL.

    Strategy: use individual `Contains Words` clauses per word, OR'd across
    Title, Description, and ReproSteps.  Each word is a separate clause
    because ADO `Contains Words 'w1 w2 w3'` fails when words include
    possessives (e.g., "Client's" doesn't match "client").

    We use up to 3 distinctive words, each requiring at least one field to
    contain it.  The clauses are AND'd together so all 3 words must appear
    somewhere.  Strict all-terms filtering happens in Python.
    """
    raw = re.sub(r"[^\w\s']", " ", str(query or "")).lower()
    raw_words = [w.strip("' ") for w in raw.split()]
    raw_words = [w for w in raw_words if len(w) >= 3]

    # Handle apostrophes: "client's" → "client"
    clean_words: List[str] = []
    for w in raw_words:
        if "'" in w:
            base = w.split("'")[0]
            if len(base) >= 3:
                clean_words.append(base)
        else:
            clean_words.append(w)

    seen: set = set()
    unique_words: List[str] = []
    for w in clean_words:
        if w not in seen:
            seen.add(w)
            unique_words.append(w)

    noise = {
        "the", "and", "for", "are", "but", "not", "you", "all", "can", "had",
        "her", "was", "one", "our", "out", "has", "his", "how", "its", "may",
        "new", "now", "old", "see", "way", "who", "did", "get", "got", "let",
        "say", "she", "too", "use", "into", "that", "this", "with", "from",
        "will", "won", "don", "isn", "does", "been", "have",
        "just", "like", "also", "need", "were", "some", "them",
        "than", "then", "what", "when", "where", "which", "would", "could",
        "should", "about", "there", "these", "those",
    }
    content_words = [w for w in unique_words if w not in noise]
    if not content_words:
        content_words = unique_words[:3]
    if not content_words:
        return base_wiql

    # Pick the single most distinctive (longest) word for the WIQL.
    # Using multiple AND'd Contains Words clauses causes ADO query timeouts
    # (30s) on large projects.  One word keeps WIQL fast; strict all-terms
    # filtering happens in Python, and the Search API (Strategy C) catches
    # anything the single-word WIQL misses.
    distinctive = sorted(content_words, key=lambda t: -len(t))[:1]
    print(f"[ADO][WIQL] clean_words={content_words} | distinctive={distinctive}")

    fields = [
        "[System.Title]",
        "[System.Description]",
        "[Microsoft.VSTS.TCM.ReproSteps]",
    ]

    # Single word across all fields with OR
    safe_word = distinctive[0].replace("'", "''")
    field_parts = " OR ".join(f"{f} Contains Words '{safe_word}'" for f in fields)
    keyword_clause = f"({field_parts})"

    wiql_upper = base_wiql.upper()
    order_pos = wiql_upper.rfind("ORDER BY")

    if order_pos > 0:
        before_order = base_wiql[:order_pos].rstrip()
        order_part = base_wiql[order_pos:]
        return f"{before_order} AND {keyword_clause} {order_part}"
    else:
        return f"{base_wiql.rstrip()} AND {keyword_clause}"


def _ado_search_api(query: str, top: int = 200) -> List[int]:
    """Use the ADO Search API to find work items by full-text search.

    This searches ALL fields including discussion/comments, which the WIQL
    Contains Words operator cannot do efficiently.  Returns work item IDs.
    """
    search_url = (
        f"https://almsearch.dev.azure.com/{ADO_ORG}"
        f"/_apis/search/workitemsearchresults?api-version=7.1-preview.1"
    )

    # Build search text: use the longest content words joined by AND so all
    # must appear somewhere in the work item (any field).
    raw = re.sub(r"[^\w\s']", " ", str(query or "")).lower()
    raw_words = [w.strip("' ") for w in raw.split()]
    raw_words = [w for w in raw_words if len(w) >= 3]
    # Strip possessives
    clean: List[str] = []
    for w in raw_words:
        if "'" in w:
            base = w.split("'")[0]
            if len(base) >= 3:
                clean.append(base)
        else:
            clean.append(w)
    # Deduplicate
    seen: set = set()
    unique: List[str] = []
    for w in clean:
        if w not in seen:
            seen.add(w)
            unique.append(w)
    noise = {
        "the", "and", "for", "are", "but", "not", "you", "all", "can", "had",
        "her", "was", "one", "our", "out", "has", "his", "how", "its", "may",
        "new", "now", "old", "see", "way", "who", "did", "get", "got", "let",
        "say", "she", "too", "use", "into", "that", "this", "with", "from",
        "will", "won", "don", "isn", "does", "been", "have",
        "just", "like", "also", "need", "were", "some", "them",
        "than", "then", "what", "when", "where", "which", "would", "could",
        "should", "about", "there", "these", "those",
    }
    content = [w for w in unique if w not in noise]
    if not content:
        content = unique[:5]
    if not content:
        return []

    # Pick up to 8 most distinctive words for search query
    distinctive = sorted(content, key=lambda t: -len(t))[:8]
    search_text = " AND ".join(distinctive)
    print(f"[ADO][SEARCH-API] searchText='{search_text}'")

    body: dict = {
        "searchText": search_text,
        "$top": min(top, 1000),
        "$skip": 0,
    }

    try:
        resp = requests.post(search_url, json=body, headers=_ado_headers(), timeout=15)
    except Exception as exc:
        print(f"[ADO][SEARCH-API][ERROR] Request failed: {exc}")
        return []

    if resp.status_code != 200:
        print(f"[ADO][SEARCH-API][ERROR] {resp.status_code}: {resp.text[:300]}")
        return []

    results = resp.json().get("results", [])
    ids: List[int] = []
    for r in results:
        fields = r.get("fields", {})
        wid = fields.get("system.id")
        if wid:
            try:
                ids.append(int(wid))
            except (ValueError, TypeError):
                pass
    print(f"[ADO][SEARCH-API] Returned {len(ids)} work item IDs")
    return ids


def _ado_fetch_work_items_lightweight(ids: List[int]) -> List[dict]:
    """Fetch work items requesting only the fields needed for keyword search (fast)."""
    if not ids:
        return []

    fields = [
        "System.Id", "System.Title", "System.Description",
        "Microsoft.VSTS.TCM.ReproSteps", "System.State",
        "System.Tags", "System.TeamProject", "System.ChangedDate",
    ]
    fields_param = ",".join(fields)
    results: List[dict] = []
    org_base = _ado_project_base(None)
    for start in range(0, len(ids), 200):
        batch_ids = ids[start:start + 200]
        det_url = (
            f"{org_base}/_apis/wit/workitems?ids={','.join(str(item_id) for item_id in batch_ids)}"
            f"&fields={fields_param}&api-version=7.1"
        )
        det = requests.get(det_url, headers=_ado_headers(), timeout=30)
        if det.status_code != 200:
            print(f"[ADO][TICKETS][ERROR] Lightweight work item fetch returned {det.status_code}: {det.text[:300]}")
            raise RuntimeError(f"ADO workitems details error: {det.text}")
        results.extend(det.json().get("value", []))
    return results


def _ado_fetch_comments_for_search(work_item_ids: List[int], max_workers: int = 10) -> Dict[int, str]:
    """Fetch discussion comments for a batch of work items concurrently.

    Returns a dict mapping work-item ID → plain-text discussion (HTML stripped).
    """
    if not work_item_ids:
        return {}

    hdrs = _ado_headers()

    def _fetch_one(wid: int) -> Tuple[int, str]:
        try:
            url = (
                f"{_ado_project_base(None)}/_apis/wit/workItems/{wid}"
                f"/comments?api-version=7.1-preview.3"
            )
            resp = requests.get(url, headers=hdrs, timeout=15)
            if resp.status_code == 200:
                texts = [
                    _strip_html(c.get("text", "") or "")
                    for c in resp.json().get("comments", [])
                ]
                return wid, " ".join(t for t in texts if t)
        except Exception as exc:
            print(f"[ADO][SEARCH][WARN] Comment fetch for #{wid} failed: {exc}")
        return wid, ""

    results: Dict[int, str] = {}
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(_fetch_one, wid): wid for wid in work_item_ids}
        for future in futures:
            wid, text = future.result()
            results[wid] = text

    print(f"[ADO][SEARCH] Fetched comments for {len(results)} work items "
          f"({sum(1 for v in results.values() if v)} have discussion)")
    return results


def ado_search_tickets_by_keywords(query: str, limit: int = 20, max_scan: int = 500) -> List[dict]:
    """Search work items by keywords.  Only tickets that contain ALL search
    terms (in their normalized title + description) are returned."""
    normalized_query = _normalize_ticket_search_text(query)
    terms = _ticket_search_terms(query)

    print(f"[ADO][SEARCH][INIT] query='{query}' | normalized='{normalized_query}' | terms={terms}")

    if not normalized_query or len(normalized_query) < 1:
        print(f"[ADO][SEARCH][REJECT] normalized_query too short or empty")
        return []

    safe_max_scan = max(1, min(int(max_scan or 500), 2000))
    safe_limit = max(1, min(int(limit or 100), 100))

    # --- Phase 1: Broad WIQL to get candidate IDs ----------------------------
    # The WIQL uses OR with a few distinctive terms to cast a wide net.
    # Strict ALL-terms filtering happens in Python on normalized text (Phase 3).
    merged_ids: List[int] = []
    seen_ids: set = set()

    def _add_ids(ids: List[int]):
        for item_id in ids:
            if item_id not in seen_ids:
                seen_ids.add(item_id)
                merged_ids.append(item_id)

    # --- Run WIQL (Strategy A) and Search API (Strategy C) in parallel ---
    # Both hit ADO over HTTP; running them concurrently halves the wait time.
    wiql_ids: List[int] = []
    search_api_ids: List[int] = []

    def _run_wiql():
        nonlocal wiql_ids
        if not ADO_KEYWORD_WIQL:
            return
        filtered_wiql = _ado_inject_keyword_conditions(ADO_KEYWORD_WIQL, query)
        print(f"[ADO][SEARCH] Using keyword-filtered WIQL: {filtered_wiql[:400]}")
        try:
            project = (ADO_PROJECT or "").strip()
            ids = _ado_execute_wiql(project, filtered_wiql, "ADO_KEYWORD_WIQL_FILTERED")
            print(f"[ADO][SEARCH] Keyword-filtered WIQL returned {len(ids)} IDs")
            wiql_ids = ids
        except Exception as exc:
            print(f"[ADO][SEARCH][WARN] Keyword-filtered WIQL failed: {exc}")

    def _run_search_api():
        nonlocal search_api_ids
        try:
            search_api_ids = _ado_search_api(query, top=200)
        except Exception as exc:
            print(f"[ADO][SEARCH][WARN] ADO Search API failed: {exc}")

    with ThreadPoolExecutor(max_workers=2) as executor:
        executor.submit(_run_wiql)
        executor.submit(_run_search_api)
        executor.shutdown(wait=True)

    _add_ids(wiql_ids)

    # Strategy B: Saved query targets with keyword injection (fallback)
    if not merged_ids:
        query_targets = _ado_keyword_ticket_query_definitions()
        print(f"[ADO][SEARCH] query_targets={query_targets}")
        for project, query_path in query_targets:
            label = f"{project}::{query_path}"
            print(f"[ADO][SEARCH] trying saved query target: {label}")
            try:
                wiql = _ado_fetch_saved_query_wiql(project, query_path)
                filtered_wiql = _ado_inject_keyword_conditions(wiql, query)
                print(f"[ADO][SEARCH] fetched & filtered WIQL: {filtered_wiql[:400]}")
                ids = _ado_execute_wiql(project, filtered_wiql, f"{label}_filtered")
                print(f"[ADO][SEARCH] got {len(ids)} IDs")
                _add_ids(ids)
            except Exception as exc:
                print(f"[ADO][SEARCH][ERROR] saved query '{label}' failed: {exc}")
                continue

    # Merge Search API results
    search_api_id_set: set = set()
    if search_api_ids:
        search_api_id_set = set(search_api_ids)
        before = len(merged_ids)
        _add_ids(search_api_ids)
        print(f"[ADO][SEARCH] Search API added {len(merged_ids) - before} new IDs (total now {len(merged_ids)})")

    print(f"[ADO][SEARCH] Total candidate IDs from WIQL: {len(merged_ids)}")

    if not merged_ids:
        print(f"[ADO][SEARCH] No ticket IDs found from any strategy.")
        return []

    # --- Phase 2: Fetch work item details (lightweight, specific fields) -----
    candidate_ids = merged_ids[:safe_max_scan]
    print(f"[ADO][SEARCH] Fetching details for {len(candidate_ids)} candidates...")
    candidates: List[dict] = []
    try:
        for wi in _ado_fetch_work_items_lightweight(candidate_ids):
            fid = wi.get("id")
            flds = wi.get("fields", {})
            title = flds.get("System.Title", "")
            desc_raw = flds.get("System.Description", "") or ""
            repro_raw = flds.get("Microsoft.VSTS.TCM.ReproSteps", "") or ""
            description = f"{desc_raw} {repro_raw}".strip()
            state = flds.get("System.State", "")
            tags = flds.get("System.Tags", "")
            project_name = flds.get("System.TeamProject", "")
            changed_date = flds.get("System.ChangedDate", "")
            web_url = f"https://dev.azure.com/{ADO_ORG}/{project_name}/_workitems/edit/{fid}"
            candidates.append(
                {
                    "id": fid,
                    "title": title,
                    "state": state,
                    "tags": tags,
                    "project": project_name,
                    "changedDate": changed_date,
                    "url": web_url,
                    "_description": description,
                }
            )
    except Exception as exc:
        print(f"[ADO][SEARCH][ERROR] Work item batch fetch failed: {exc}")

    # --- Phase 2.5: Two-pass filtering to avoid fetching comments for all candidates.
    # Pass A: check title + description only (already fetched). Tickets where
    # ALL terms match go straight to results. Others are queued for discussion
    # check, sorted by how many terms already match (best partial matches first).
    print(f"[ADO][SEARCH] Pass A: title+description filter for {len(candidates)} candidates ({len(terms)} terms)...")

    matched: List[dict] = []
    needs_discussion: List[dict] = []

    for item in candidates:
        work_item_id = item.get("id")
        if not work_item_id:
            continue
        title = item.get("title", "") or ""
        description = item.get("_description", "") or ""
        td_norm = _normalize_ticket_search_text(f"{title} {description}")

        missing = [t for t in terms if t not in td_norm]
        if not missing:
            # All terms found in title+description — immediate match
            title_norm = _normalize_ticket_search_text(title)
            hits_in_title = sum(1 for t in terms if t in title_norm)
            total_occurrences = sum(td_norm.count(t) for t in terms)
            matched.append({**item, "_hits_in_title": hits_in_title, "_total_occurrences": total_occurrences, "_combined_norm": td_norm})
        elif terms:
            hit_count = len(terms) - len(missing)
            # Queue for discussion check — sort by hit_count later to prioritize best partial matches
            needs_discussion.append({**item, "_missing": missing, "_td_norm": td_norm, "_hit_count": hit_count})

    print(f"[ADO][SEARCH] Pass A done: {len(matched)} immediate matches, {len(needs_discussion)} need discussion check")

    # Pass B: fetch discussion ONLY for top candidates sorted by partial match quality.
    # If Pass A already found matches, limit discussion fetch to save time.
    # If Pass A found 0 matches, invest more in discussion search.
    # Candidates found by the Search API are always included (they matched
    # the full query server-side, so they're very likely to pass strict filtering).
    discussion_cap = 10 if matched else 50
    if needs_discussion:
        # Partition: Search-API hits first (always checked), then the rest sorted by hit_count.
        from_search_api = [x for x in needs_discussion if x.get("id") in search_api_id_set]
        from_wiql = [x for x in needs_discussion if x.get("id") not in search_api_id_set]
        from_wiql.sort(key=lambda x: -x.get("_hit_count", 0))
        remaining_cap = max(0, discussion_cap - len(from_search_api))
        needs_discussion = from_search_api + from_wiql[:remaining_cap]
        discussion_ids = [item["id"] for item in needs_discussion]
        print(f"[ADO][SEARCH] Pass B: fetching discussion for {len(discussion_ids)} tickets ({len(from_search_api)} from Search API)...")
        try:
            discussion_map = _ado_fetch_comments_for_search(discussion_ids, max_workers=20)
        except Exception as exc:
            print(f"[ADO][SEARCH][WARN] Discussion fetch failed: {exc}")
            discussion_map = {}

        for item in needs_discussion:
            work_item_id = item["id"]
            discussion = discussion_map.get(work_item_id, "")
            if not discussion:
                continue
            discussion_norm = _normalize_ticket_search_text(discussion)
            # Check if the missing terms are in the discussion
            still_missing = [t for t in item["_missing"] if t not in discussion_norm]
            if still_missing:
                continue
            # All terms accounted for across title+desc+discussion
            title = item.get("title", "") or ""
            title_norm = _normalize_ticket_search_text(title)
            combined_norm = f"{item['_td_norm']} {discussion_norm}"
            hits_in_title = sum(1 for t in terms if t in title_norm)
            total_occurrences = sum(combined_norm.count(t) for t in terms)
            matched.append({
                k: v for k, v in item.items() if not k.startswith("_")
            } | {"_hits_in_title": hits_in_title, "_total_occurrences": total_occurrences, "_combined_norm": combined_norm})

    # Sort: more title hits first, then more total occurrences, then newest
    matched.sort(
        key=lambda x: (
            -int(x.get("_hits_in_title") or 0),
            -int(x.get("_total_occurrences") or 0),
            str(x.get("changedDate") or ""),
        )
    )

    # --- Phrase proximity filter ---
    # When multiple tickets match all keyword terms, prefer those that contain
    # the original query as a contiguous phrase.  This eliminates false positives
    # where keywords happen to appear scattered across unrelated sentences.
    query_words = normalized_query.split()
    if len(matched) > 1 and len(query_words) >= 3:
        phrase_matches = [item for item in matched if normalized_query in (item.get("_combined_norm") or "")]
        if phrase_matches and len(phrase_matches) < len(matched):
            print(f"[ADO][SEARCH] Phrase filter: {len(phrase_matches)} of {len(matched)} tickets contain the full query phrase")
            matched = phrase_matches

    results = [
        {k: v for k, v in item.items() if not k.startswith("_")}
        for item in matched[:safe_limit]
    ]
    print(f"[ADO][SEARCH] Done. {len(matched)} tickets have ALL terms, returning top {len(results)}")
    return results

def ado_parse_id_from_url(url: str) -> Optional[int]:
    try:
        # expected e.g. https://dev.azure.com/{org}/{project}/_workitems/edit/{id}
        # Be resilient to trailing params or alternative shapes; grab the last digit run.
        digits = re.findall(r"(\d+)", url or "")
        if not digits:
            return None
        return int(digits[-1])
    except Exception:
        return None

def ado_fetch_ticket_text(work_item_id: int, skip_image_analysis: bool = False) -> str:
    """Fetch full ticket context as chronological plain text, including attachments."""
    ticket = ado_fetch_ticket_full(work_item_id)
    fields = ticket.get("fields", {})
    comments = sorted(ticket.get("comments", []), key=lambda item: item.get("createdDate") or "")
    attachments = ticket.get("attachments", [])
    project_name = ticket.get("project") or fields.get("System.TeamProject") or ADO_PROJECT

    title = fields.get("System.Title", "")
    state = fields.get("System.State", "")
    tags = fields.get("System.Tags", "")
    work_item_type = fields.get("System.WorkItemType", "")
    assigned_to = _first_field(fields, ["System.AssignedTo"])
    raw_desc = fields.get("System.Description", "") or fields.get("Microsoft.VSTS.TCM.ReproSteps", "") or ""
    description_text = _render_ado_html_with_inline_images(
        raw_desc,
        section_label=f"ticket {work_item_id} description",
        project=project_name,
        skip_image_analysis=skip_image_analysis,
    ).strip() or "Not provided."

    discussion_blocks: List[str] = []
    total_comments = len(comments)
    for index, comment in enumerate(comments, start=1):
        author = comment.get("createdBy") or "Unknown"
        created_date = comment.get("createdDate") or ""
        body = _render_ado_html_with_inline_images(
            comment.get("raw_html") or comment.get("text") or "",
            section_label=f"ticket {work_item_id} discussion comment {index}",
            project=project_name,
            skip_image_analysis=skip_image_analysis,
        ).strip() or "[Empty discussion entry]"
        header = f"[{created_date}] {author}"
        discussion_blocks.append(f"{header}\n{body}")
        print(f"[ADO][TICKET] Processed discussion entry {index}/{total_comments} for work item {work_item_id}")

    attachment_text = _build_ticket_attachment_context(
        work_item_id,
        attachments,
        project=project_name,
        skip_image_analysis=skip_image_analysis,
    ).strip()

    sections = [
        f"Ticket ID: {work_item_id}",
        f"Project: {project_name}",
        f"Type: {work_item_type}",
        f"Title: {title}",
        f"State: {state}",
        f"Assigned To: {assigned_to}",
        f"Tags: {tags}",
        "",
        "Description (full, chronological):",
        description_text,
        "",
        "Discussion (full, chronological):",
        "\n\n---\n\n".join(discussion_blocks) if discussion_blocks else "Not provided.",
    ]

    if attachment_text:
        sections.extend(
            [
                "",
                "Attachments (full content, chronological when Azure DevOps provides dates):",
                attachment_text,
            ]
        )

    blob = "\n".join(section for section in sections if section is not None).strip()
    if not blob:
        blob = f"(Empty work item {work_item_id})"
    return blob


def collection_empty() -> bool:
    try:
        c = collection.count()
        print(f"[CHROMA] collection.count() = {c}")
        return c == 0
    except Exception as e:
        print(f"[CHROMA][ERROR] Failed to get collection count: {e}")
        traceback.print_exc()
        return True


def get_ingested_sources() -> List[str]:
    """
    Return the list of all 'source' values from collection metadatas.
    Used for 'only_missing' ingest mode.
    """
    try:
        print("[CHROMA] Fetching existing metadatas for ingested sources list")
        data = collection.get(include=["metadatas"])
        ingested = sorted({m["source"] for m in data["metadatas"] if m})
        print(f"[CHROMA] Found {len(ingested)} distinct sources already ingested")
        return ingested
    except Exception as e:
        print(f"[CHROMA][ERROR] Failed to get ingested sources: {e}")
        traceback.print_exc()
        return []


# ---------------------------------------------------------------------
# Ingest wiki files
# ---------------------------------------------------------------------

def ingest_wiki_files(
    force: bool = False,
    only_missing: bool = False,
    include_empty: bool = False,
    selected_paths: Optional[List[str]] = None,
    ignore_patterns: Optional[List[str]] = None,
) -> int:
    """
    Ingest .md files under WIKI_ROOT into the Chroma collection.

    Args:
        force:
            - If True:
                * Delete the existing collection and recreate it.
                * Ingest from scratch (ignores only_missing).
        only_missing:
            - If True and force is False:
                * Do NOT skip when collection is non-empty.
                * Only ingest files whose 'source' is NOT already in Chroma.
        include_empty:
            - If True, 0-byte / empty .md files are still ingested
              as a single placeholder chunk.
        selected_paths:
            - Optional list of specific relative paths (from WIKI_ROOT)
              to ingest. If provided, only those files will be processed.
    """
    print(
        f"[INGEST] ingest_wiki_files(force={force}, only_missing={only_missing}, "
        f"include_empty={include_empty}, selected_paths_count={0 if not selected_paths else len(selected_paths)})"
    )
    print(f"[INGEST] WIKI_ROOT={WIKI_ROOT} CHROMA_DIR={CHROMA_DIR} COLLECTION_NAME={COLLECTION_NAME}")

    if not os.path.isdir(WIKI_ROOT):
        print(f"[INGEST][ERROR] WIKI_ROOT not found: {WIKI_ROOT}")
        raise RuntimeError(f"WIKI_ROOT not found: {WIKI_ROOT}")

    global collection

    # ----- FORCE mode: drop and recreate collection -----
    if force:
        print(f"[INGEST] FORCE mode: deleting collection '{COLLECTION_NAME}'")
        start_del = time.time()
        client.delete_collection(COLLECTION_NAME)
        print(f"[INGEST] Collection '{COLLECTION_NAME}' deleted in {time.time() - start_del:.2f}s")
        collection = client.get_or_create_collection(name=COLLECTION_NAME)
        print(f"[INGEST] Collection '{COLLECTION_NAME}' recreated")

    # ----- Build base file list -----
    if selected_paths:
        # Use only the explicitly specified relative paths
        all_files: List[pathlib.Path] = []
        for rel in selected_paths:
            abs_path = os.path.join(WIKI_ROOT, rel)
            if os.path.isfile(abs_path):
                all_files.append(pathlib.Path(abs_path))
                print(f"[INGEST] Selected file found: {rel}")
            else:
                print(f"[INGEST][WARN] Selected file not found on disk: {rel}")
    else:
        # Use all markdown files under WIKI_ROOT
        all_files = list(iter_markdown_files())

    # ----- Apply .ragignore or provided ignore_patterns -----
    ignore_file = os.path.join(WIKI_ROOT, ".ragignore")
    # Also allow a repo-level .ragignore (one level up from WIKI_ROOT)
    repo_ignore = os.path.join(os.path.dirname(WIKI_ROOT), ".ragignore")
    compiled_ignores: List[str] = []
    if os.path.exists(ignore_file):
        try:
            with open(ignore_file, "r", encoding="utf-8") as f:
                for ln in f:
                    ln = ln.strip()
                    if not ln or ln.startswith("#"):
                        continue
                    compiled_ignores.append(ln)
            print(f"[INGEST] Loaded {len(compiled_ignores)} ignore patterns from .ragignore")
        except Exception as e:
            print(f"[INGEST][WARN] Failed to read .ragignore: {e}")

    if ignore_patterns:
        compiled_ignores.extend(ignore_patterns)

    # load repo-level .ragignore if present
    if os.path.exists(repo_ignore):
        try:
            with open(repo_ignore, "r", encoding="utf-8") as f:
                for ln in f:
                    ln = ln.strip()
                    if not ln or ln.startswith("#"):
                        continue
                    compiled_ignores.append(ln)
            print(f"[INGEST] Loaded {len(compiled_ignores)} ignore patterns (including repo .ragignore)")
        except Exception as e:
            print(f"[INGEST][WARN] Failed to read repo .ragignore: {e}")

    if compiled_ignores:
        filtered_files = []
        for p in all_files:
            rel = os.path.relpath(str(p), WIKI_ROOT)
            skip = False
            for pat in compiled_ignores:
                # treat patterns as globs relative to WIKI_ROOT
                if fnmatch.fnmatch(rel, pat):
                    print(f"[INGEST] Skipping file due to ignore pattern: {rel} (pattern: {pat})")
                    skip = True
                    break
            if not skip:
                filtered_files.append(p)
        all_files = filtered_files

    total_files = len(all_files)
    print(f"[INGEST] Base file set size (before only_missing filter) = {total_files}")

    # ----- only_missing mode (skip already ingested sources) -----
    if only_missing and not force:
        existing_sources = set(get_ingested_sources())
        print(f"[INGEST] Filtering files by only_missing=True (existing_sources={len(existing_sources)})")
        filtered_files: List[pathlib.Path] = []
        for p in all_files:
            rel_path = os.path.relpath(str(p), WIKI_ROOT)
            if rel_path in existing_sources:
                print(f"[INGEST] Skipping already ingested file: {rel_path}")
                continue
            filtered_files.append(p)
        all_files = filtered_files
        total_files = len(all_files)
        print(f"[INGEST] After only_missing filter, files to ingest={total_files}")
    else:
        if not force and not selected_paths and not only_missing:
            # Original behavior: if collection not empty and no special mode → skip
            empty = collection_empty()
            print(f"[INGEST] collection_empty()={empty}")
            if not empty:
                print("[INGEST] Collection not empty and force=False, only_missing=False → skipping ingest entirely")
                return 0

    print(f"[INGEST] FINAL file count to process = {total_files}")

    added_chunks = 0
    started_at = time.time()

    for file_idx, md_path in enumerate(all_files, start=1):
        rel_path = os.path.relpath(str(md_path), WIKI_ROOT)
        print(f"[INGEST] [{file_idx}/{total_files}] Processing file: {rel_path}")

        try:
            with open(md_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        except Exception as e:
            print(f"[INGEST][ERROR] Failed to read file {rel_path}: {e}")
            traceback.print_exc()
            continue

        if not text.strip():
            if include_empty:
                print(f"[INGEST] File {rel_path} is empty → ingesting placeholder chunk")
                text = f"(Empty TA9 wiki page: {rel_path})"
            else:
                print(f"[INGEST] File {rel_path} is empty → skipping (include_empty=False)")
                continue

        chunks = chunk_text(text)
        if not chunks:
            print(f"[INGEST] File {rel_path} produced 0 chunks → skipping")
            continue

        print(f"[INGEST] File {rel_path} produced {len(chunks)} chunks")

        ids: List[str] = []
        docs: List[str] = []
        metas: List[dict] = []
        embeds: List[List[float]] = []

        for i, chunk in enumerate(chunks):
            try:
                print(f"[INGEST] Embedding chunk {i+1}/{len(chunks)} of file {rel_path}")
                emb = embed_text(chunk)
            except Exception as e:
                print(f"[INGEST][ERROR] Embedding failed for file={rel_path} chunk={i}: {e}")
                traceback.print_exc()
                # Skip this chunk but continue with others
                continue

            ids.append(str(uuid.uuid4()))
            docs.append(chunk)
            metas.append({"source": rel_path, "chunk": i})
            embeds.append(emb)

        if not ids:
            print(f"[INGEST][WARN] No successful chunks for file {rel_path} → skipping add()")
            continue

        try:
            print(f"[INGEST] Adding {len(ids)} chunks for file {rel_path} to Chroma")
            collection.add(
                ids=ids,
                documents=docs,
                metadatas=metas,
                embeddings=embeds,
            )
            added_chunks += len(ids)
            print(f"[INGEST] Successfully added {len(ids)} chunks. Total so far={added_chunks}")
        except Exception as e:
            print(f"[INGEST][ERROR] Failed to add chunks for file {rel_path}: {e}")
            traceback.print_exc()

    total_time = time.time() - started_at
    print(f"[INGEST] Completed ingest. Total chunks added={added_chunks} in {total_time:.2f}s")
    try:
        new_count = collection.count()
        print(f"[INGEST] collection.count() after ingest = {new_count}")
    except Exception as e:
        print(f"[INGEST][WARN] Could not get collection count after ingest: {e}")
    return added_chunks


# ---------------------------------------------------------------------
# RAG “brain” – augmentation + reranking
# ---------------------------------------------------------------------

def _normalize_question(question: str) -> str:
    """Normalize question to a canonical form to improve recall."""
    if not question:
        return ""
    q = question.lower()
    q = re.sub(r"[^a-z0-9\s\-]", " ", q)
    q = re.sub(r"\s+", " ", q).strip()
    return q


def _is_ta9_question(question: str) -> bool:
    if not question:
        return False
    q = question.lower()
    ta9_terms = [
        "ta9", "intsight", "int-sight", "intsight", "platform", "system",
        "features", "capabilities", "modules", "use cases", "overview",
        "dashboard", "admin studio", "data model", "dm",
        "situational awareness", "incident", "incidents", "map", "layer", "layers",
        "camera", "cameras", "protocol", "federated search", "ontology","kyc"
    ]
    return any(t in q for t in ta9_terms)


def _is_foundational_question(question: str) -> bool:
    """Detect if question is about foundational product concepts that should always be answered."""
    if not question:
        return False
    q = question.lower()
    foundational_terms = [
        "what is", "what are", "explain", "define", "how does", "how do",
        "data model", "entity", "detection", "dashboard", "admin studio",
        "module", "feature", "capability", "system", "platform",
        "link analysis", "timeline", "geospatial", "visualization",
        "criteria", "field", "dm", "configuration"
    ]
    return any(t in q for t in foundational_terms)


def _has_ta9_context(text: str) -> bool:
    """Domain gate: detects clear TA9/IntSight product context."""
    if not text:
        return False
    t = text.lower()
    term_groups = {
        "brand": [
            "ta9", "t-a9", "t a9", "intsight", "int-sight", "int sight",
        ],
        "core_modules": [
            "admin studio", "kyc", "federated search", "data model", "data models", "entities",
            "relations", "ontology manager", "main graph", "link analysis", "cases",
            "tasks", "situational awareness", "dashboard", "insight", "annotations", "autoloader",
        ],
        "platform_terms": [
            "identifier", "taxonomy", "query builder", "cluster query", "field role",
            "is free text", "is id", "sequence", "permission mode", "case profile",
            "indexing service", "index to federated", "admin tools", "lookup manager",
        ],
    }

    group_hits = 0
    for _, terms in term_groups.items():
        if any(term in t for term in terms):
            group_hits += 1

    return group_hits >= 1


def _has_ta9_support_signal(text: str) -> bool:
    """
    Detects meaningful technical/procedural support content aligned with TA9 guide material.
    Accepts operational documentation, configuration steps, integration guidance, and troubleshooting.
    """
    if not text:
        return False

    t = text.lower()

    workflow_terms = [
        "step", "click", "open", "select", "save", "configure", "set", "define", "navigate",
        "upload", "download", "reset", "create", "edit", "delete", "assign", "grant",
        "permission", "role", "profile", "workflow", "prerequisite", "note:",
    ]
    technical_terms = [
        "api", "rest", "endpoint", "json", "regex", "query", "sql", "mysql", "mariadb",
        "solr", "orient", "indexing", "schema", "field", "mapping", "lookup", "parser",
        "config", "dll", "path", "service", "cache", "cron", "batch", "token", "2fa",
        "sso", "active directory", "ldap", "authentication", "authorization", "audit",
    ]
    product_artifacts = [
        "admin", "analyst", "developer", "entity", "relation", "case", "incident", "protocol",
        "map", "graph", "gantt", "timeline", "facet", "widget", "dashboard", "federated",
        "data loader", "load file manager", "autoloader", "document viewer", "speech to text",
        "translation", "system config", "localization", "situational awareness", "situational picture", "sit pit",
    ]
    integration_terms = [
        "kerberos", "krb5", "krb5.conf", "krb5.keytab", "keytab", "spn", "realm", "kdc",
        "mssql", "sql server", "odbc", "odbc.ini", "trusted_connection", "integratedsecurity",
        "authmech", "krbservicename", "krbhostfqdn", "krbauthrealm", "active directory",
        "domain controller", "dc", "ad user",
    ]
    application_reference_terms = [
        "integration", "integrations", "integrate", "integrated with", "connector", "connectors",
        "application", "applications", "module", "modules", "plugin", "plugins", "supported app",
        "supported apps", "external system", "external systems", "third-party", "third party",
        "known as", "also known as", "alias", "aliases", "alternative name", "alternative names",
        "naming", "mapped to", "referred to as",
    ]

    workflow_hits = sum(1 for term in workflow_terms if term in t)
    technical_hits = sum(1 for term in technical_terms if term in t)
    artifact_hits = sum(1 for term in product_artifacts if term in t)
    integration_hits = sum(1 for term in integration_terms if term in t)
    application_reference_hits = sum(1 for term in application_reference_terms if term in t)

    has_structured_config_pattern = bool(
        re.search(r"\b[a-z0-9_.-]+\s*=\s*[^\s].+", t)
        or re.search(r"<add\s+key=", t)
        or re.search(r"\{\s*\"[a-z0-9_\-]+\"\s*:", t)
    )
    has_stepwise_pattern = bool(re.search(r"\bstep\s*\d+\b", t))
    has_version_or_section_pattern = bool(re.search(r"\b(v|version)\s*\d+(\.\d+)*\b", t))

    # Broad but meaningful acceptance for guide-like knowledge
    if workflow_hits >= 3 and artifact_hits >= 1:
        return True
    if technical_hits >= 3 and artifact_hits >= 1:
        return True
    if artifact_hits >= 2 and (has_structured_config_pattern or has_stepwise_pattern):
        return True
    if artifact_hits >= 2 and technical_hits >= 2:
        return True
    if has_version_or_section_pattern and artifact_hits >= 2 and (workflow_hits >= 2 or technical_hits >= 2):
        return True
    # Explicit acceptance for enterprise integration/config guides (e.g., Kerberos + MSSQL setup)
    if integration_hits >= 3 and has_structured_config_pattern:
        return True
    if integration_hits >= 4 and (workflow_hits >= 2 or technical_hits >= 2):
        return True
    if "prerequisite" in t and integration_hits >= 3:
        return True
    # Accept IntSight application/integration reference knowledge even when it is descriptive rather than procedural.
    if application_reference_hits >= 2 and artifact_hits >= 1:
        return True
    if application_reference_hits >= 1 and artifact_hits >= 2:
        return True

    return False


def _looks_clearly_out_of_scope_knowledge(text: str) -> bool:
    """Reject only obvious junk or clearly unrelated content."""
    if not text:
        return True

    if _has_ta9_context(text) or _has_ta9_support_signal(text):
        return False

    candidate = str(text or "").strip()
    lower_candidate = candidate.lower()
    word_count = len(re.findall(r"\b\w+\b", candidate))

    if word_count < 4:
        return True

    if re.fullmatch(r"[\W_]+", candidate):
        return True

    repeated_tokens = re.findall(r"\b([a-z0-9]{2,})\b", lower_candidate)
    if repeated_tokens and len(set(repeated_tokens)) == 1 and len(repeated_tokens) >= 4:
        return True

    off_topic_terms = {
        "food": ["banana", "recipe", "cook", "cooking", "bake", "oven", "ingredient", "meal", "diet", "nutrition"],
        "travel": ["flight", "hotel", "vacation", "beach", "tourism", "itinerary", "passport", "resort"],
        "entertainment": ["movie", "series", "celebrity", "song", "lyrics", "album", "netflix"],
        "lifestyle": ["dating", "horoscope", "astrology", "workout", "gym", "fashion", "makeup"],
        "creative": ["joke", "poem", "haiku", "story prompt", "fan fiction"],
    }
    off_topic_hits = 0
    for terms in off_topic_terms.values():
        if any(term in lower_candidate for term in terms):
            off_topic_hits += 1

    conversational_irrelevance = any(
        phrase in lower_candidate
        for phrase in [
            "tell me a joke",
            "write a poem",
            "how to cook",
            "how do i eat",
            "what should i eat",
            "travel plan",
            "write me a joke",
            "write me a poem",
        ]
    )

    return off_topic_hits >= 1 or conversational_irrelevance


def _strip_non_knowledge_markers(text: str) -> str:
    candidate = str(text or "")
    candidate = re.sub(r"\*\*File:\s*.+?\*\*", " ", candidate)
    candidate = candidate.replace("---FILE SEPARATOR---", " ")
    candidate = re.sub(r"(?mi)^\[(error|warning|notice|canceled)[^\n]*\]", " ", candidate)
    candidate = re.sub(r"(?mi)^\[image content from [^\n]+\]", " ", candidate)
    candidate = re.sub(r"(?mi)^\[visual_fallback\]", " ", candidate)
    candidate = re.sub(r"(?mi)^\[image analysis unavailable\]", " ", candidate)
    candidate = re.sub(r"\s+", " ", candidate)
    return candidate.strip()


def _looks_like_failed_or_empty_extraction(text: str) -> bool:
    candidate = str(text or "").strip()
    if not candidate:
        return True

    has_failure_marker = bool(re.search(r"(?mi)^\[(error|warning|notice|canceled)[^\n]*\]", candidate))
    normalized = _strip_non_knowledge_markers(candidate)
    meaningful_word_count = len(re.findall(r"\b\w+\b", normalized))

    if has_failure_marker and meaningful_word_count < 8:
        return True

    if meaningful_word_count < 4:
        return True

    return False


def _validate_ta9_knowledge_content(content_text: str) -> Tuple[bool, str]:
    """Approve by default and reject only obvious junk or irrelevant content."""
    text = (content_text or "").strip()
    if not text:
        return False, "The content is empty. Please provide TA9-related details before adding knowledge."

    if _looks_like_failed_or_empty_extraction(text):
        return False, (
            "The new knowledge could not be extracted into usable content. "
            "Please provide meaningful text or upload a file that contains readable, relevant information."
        )

    if _has_ta9_context(text) or _has_ta9_support_signal(text):
        return True, "Approved"

    text_word_count = len(re.findall(r"\b\w+\b", text))
    if text_word_count < 4:
        return False, (
            "The content is too short to classify. "
            "Please add a few more words so it can be stored as useful knowledge."
        )

    if not _looks_clearly_out_of_scope_knowledge(text):
        return True, "Approved"

    return False, (
        "The content looks clearly out of scope for this knowledge base. "
        "Only obvious junk or unrelated topics are blocked."
    )


_COMPOUND_SPLIT_PATTERN = re.compile(
    # Splits BEFORE the next question/imperative word using a lookahead so the
    # sub-question retains its leading "how", "what", etc. and stays a strong
    # embedding query.
    r"\s*(?:,\s*)?\band\s+then\s+(?=(?:how|what|where|when|why|which|who|create|configure|set\s*up|setup|add|edit|update|delete|remove|enable|disable|fix|resolve|install|deploy|search|find|view|open)\b)"
    r"|\s*(?:,\s*)?\bthen\s+(?=(?:how|what|where|when|why|which|who|create|configure|set\s*up|setup|add|edit|update|delete|remove|enable|disable|fix|resolve|install|deploy|search|find|view|open)\b)"
    r"|\s*(?:,\s*)?\band\s+(?=(?:how|what|where|when|why|which|who)\b)"
    r"|\s*(?:,\s*)?\band\s+also\s+(?=(?:how|what|where|when|why|which|who|create|configure|set\s*up|setup|add|edit|update|delete|remove|enable|disable|fix|resolve|install|deploy|search|find|view|open)\b)"
    r"|\s*(?:,\s*)?\bafter\s+that\s*,?\s*(?=(?:how|what|where|when|why|which|who|create|configure|set\s*up|setup|add|edit|update|delete|remove|enable|disable|fix|resolve|install|deploy|search|find|view|open)\b)"
    r"|\s*(?:,\s*)?\bas\s+well\s+as\s+(?=(?:how|what|where|when|why|which|who|create|configure|set\s*up|setup|add|edit|update|delete|remove|enable|disable|fix|resolve|install|deploy|search|find|view|open)\b)"
    r"|\s*(?:,\s*)?\bin\s+addition\s+to\s+(?=(?:how|what|where|when|why|which|who|create|configure|set\s*up|setup|add|edit|update|delete|remove|enable|disable|fix|resolve|install|deploy|search|find|view|open)\b)"
    r"|\s*(?:,\s*)?\bplus\s+(?=(?:how|what|where|when|why|which|who|create|configure|set\s*up|setup|add|edit|update|delete|remove|enable|disable|fix|resolve|install|deploy|search|find|view|open)\b)"
    r"|\s*;\s*"
    r"|\s*\?\s+(?=(?:how|what|where|when|why|which|who|create|configure|set\s*up|setup|add|edit|update|delete|remove|enable|disable|fix|resolve|install|deploy|search|find|view|open)\b)",
    flags=re.IGNORECASE,
)

_QUESTION_LEAD_RE = re.compile(
    r"^\s*(?:how(?:\s+(?:to|do|can|would|should))?|what|where|when|why|which|who|can\s+i|could\s+i|please|create|configure|set\s*up|setup|add|edit|update|delete|remove|enable|disable|fix|resolve|install|deploy|search|find|view|open)\b",
    flags=re.IGNORECASE,
)


# Hard cap to keep retrieval+rerank bounded even for extreme inputs.
_MAX_SUB_QUESTIONS = 15


def _decompose_compound_question(question: str) -> List[str]:
    """
    Heuristically split a compound question into atomic sub-questions.

    Used to drive multi-query retrieval so each sub-question gets its own
    embedding and its own top-k pool, instead of competing for slots inside a
    single averaged-embedding query. Capped at _MAX_SUB_QUESTIONS to keep
    retrieval bounded for any input size.
    """
    if not question:
        return []
    text = question.strip()
    if not text:
        return []

    # First split on strong compound markers ("and then", "and how", ";", "?", etc.)
    raw_parts = [p.strip(" ,.-\t") for p in _COMPOUND_SPLIT_PATTERN.split(text) if p and p.strip(" ,.-\t")]

    sub_questions: List[str] = []
    for part in raw_parts:
        # Drop trailing punctuation, keep contents
        cleaned = part.strip(" ,.-\t?")
        if not cleaned:
            continue
        # Only treat as a sub-question if it looks like a question/imperative
        if _QUESTION_LEAD_RE.match(cleaned) or len(cleaned.split()) >= 4:
            # Re-append '?' so the embedding matches the standalone-question form
            # ("how to create a relation?" embeds slightly differently from
            # "how to create a relation" and standalone form is the known-good one).
            normalized = cleaned + "?"
            if normalized not in sub_questions:
                sub_questions.append(normalized)

    # Require at least 2 sub-questions to be considered compound
    if len(sub_questions) < 2:
        return []

    return sub_questions[:_MAX_SUB_QUESTIONS]


def _build_query_variants(question: str, ta9_mode: bool) -> List[str]:
    """Generate a small set of query variants for multi-retrieval (no LLM calls)."""
    variants = []
    base = question.strip()
    if base:
        variants.append(base)
    normalized = _normalize_question(question)
    if normalized and normalized not in variants:
        variants.append(normalized)

    # Extract keyword-only variant for broader recall
    keywords = _extract_query_keywords(base)
    if keywords:
        keyword_query = " ".join(keywords)
        if keyword_query and keyword_query not in variants:
            variants.append(keyword_query)

    head_variants = variants[:3]

    # Compound-question decomposition: append each atomic sub-question as its own
    # retrieval variant. This restores per-sub-question recall when the user
    # combines multiple distinct asks in one prompt.
    sub_questions = _decompose_compound_question(base)
    for sq in sub_questions:
        if sq and sq not in head_variants:
            head_variants.append(sq)

    return head_variants

def augment_question(question: str) -> str:
    """
    Minimal augmentation: only expand very short questions with common synonyms.
    Keep general-purpose for any domain.
    """
    if not question or len(question.split()) > 15:
        # Only augment genuinely short queries
        return question

    if _question_has_entity_fact_intent(question):
        return question
    
    q_low = question.lower()
    synonyms: List[str] = []
    
    # Minimal, general augmentation for short queries
    if (("how" in q_low and "how many" not in q_low and "how much" not in q_low) or "add" in q_low or "create" in q_low):
        synonyms.extend(["how to", "instructions", "steps", "procedure", "create", "add", "setup", "configure"])
    
    if "what" in q_low:
        synonyms.extend(["definition", "explanation", "description", "overview"])
    
    if "help" in q_low or "issue" in q_low or "problem" in q_low:
        synonyms.extend(["troubleshooting", "solution", "fix", "resolve", "error"])
    
    if not synonyms:
        return question
    
    augmented = question + "\n" + " ".join(set(synonyms))
    return augmented


def _question_has_procedural_intent(question: str) -> bool:
    q = (question or "").lower()
    if not q:
        return False
    procedural_terms = [
        "how to", "step by step", "step-by-step", "steps", "procedure", "configure",
        "configuration", "setup", "set up", "create", "add", "edit", "change",
        "update", "install", "enable", "disable", "fix", "resolve",
    ]
    return any(term in q for term in procedural_terms)


def _question_requests_broad_coverage(question: str) -> bool:
    q = (question or "").lower().strip()
    if not q:
        return False

    explicit_patterns = [
        r"\bis there (any )?more (ways|options|methods|approaches)\b",
        r"\bmore (ways|options|methods|approaches)\b",
        r"\bother (ways|options|methods|approaches)\b",
        r"\banother (way|option|method|approach)\b",
        r"\bdifferent (ways|options|methods|approaches)\b",
        r"\balternative(s)?\b",
        r"\bwhat else\b",
        r"\bany other\b",
        r"\blist (all )?(ways|options|methods|approaches|types|variants)\b",
        r"\bwhat are the (ways|options|methods|approaches|types)\b",
        r"\bwhich (ways|options|methods|approaches|types)\b",
        r"\bsupported (ways|options|methods|approaches)\b",
        r"\bavailable (ways|options|methods|approaches)\b",
    ]
    if any(re.search(pattern, q) for pattern in explicit_patterns):
        return True

    breadth_terms = ["way", "ways", "option", "options", "method", "methods", "approach", "approaches"]
    breadth_modifiers = ["more", "other", "another", "different", "alternative", "alternatives", "available", "supported", "all"]
    return any(term in q for term in breadth_terms) and any(modifier in q for modifier in breadth_modifiers)


def _question_has_entity_fact_intent(question: str) -> bool:
    q = (question or "").lower()
    if not q:
        return False

    explicit_patterns = [
        r"\bhow many\s+(product|products|module|modules|service|services|platform|platforms)\b",
        r"\bwhich\s+(product|products|module|modules|service|services|platform|platforms)\b",
        r"\bwhat\s+(product|products|module|modules|service|services|platform|platforms)\b",
        r"\b(company|vendor|organization|brand)\s+name\b",
        r"\bproduct\s+name(s)?\b",
        r"\bnaming\s+(policy|convention|rule|rules)\b",
        r"\breferred\s+to\s+as\b",
        r"\bcalled\b",
        r"\balias(es)?\b",
        r"\bmain\s+platform\b",
    ]
    if any(re.search(pattern, q) for pattern in explicit_patterns):
        return True

    fact_terms = [
        "company", "companies", "product", "products", "platform", "platforms",
        "module", "modules", "vendor", "brand", "name", "names", "named",
        "naming", "called", "alias", "aliases", "provide", "provides", "offered",
        "offer", "offers", "created", "create",
    ]
    term_hits = sum(1 for term in fact_terms if term in q)
    return term_hits >= 2


_LEXICAL_STOPWORDS = {
    "the", "is", "are", "a", "an", "and", "or", "to", "of", "in", "on", "for", "with", "by",
    "from", "as", "at", "be", "this", "that", "these", "those", "it", "its", "if", "then",
    "how", "what", "when", "where", "why", "who", "can", "could", "should", "would", "do",
    "does", "did", "i", "we", "you", "they", "he", "she", "my", "our", "your", "their",
    "about", "into", "through", "using", "use", "want", "need", "like", "please",
}


def _normalize_term(term: str) -> str:
    t = (term or "").strip().lower()
    t = re.sub(r"[^a-z0-9_\-]", "", t)
    if not t:
        return ""
    # Light stemming to match camera/cameras, layer/layers, etc.
    if len(t) > 5 and t.endswith("ies"):
        t = t[:-3] + "y"
    elif len(t) > 4 and t.endswith("es"):
        t = t[:-2]
    elif len(t) > 3 and t.endswith("s"):
        t = t[:-1]
    return t


def _tokenize_normalized(text: str) -> List[str]:
    raw_tokens = re.findall(r"[a-zA-Z0-9_\-]+", (text or "").lower())
    out: List[str] = []
    for token in raw_tokens:
        norm = _normalize_term(token)
        if len(norm) <= 2 or norm in _LEXICAL_STOPWORDS:
            continue
        out.append(norm)
    return out


def _extract_key_phrases(question: str, max_phrases: int = 6) -> List[str]:
    q = (question or "").strip().lower()
    if not q:
        return []

    phrases: List[str] = []
    # Quoted phrases should be treated as high-signal hints.
    for quoted in re.findall(r'"([^"]{3,80})"', q):
        cleaned = " ".join(quoted.split()).strip()
        if cleaned:
            phrases.append(cleaned)

    tokens = _tokenize_normalized(q)
    # Add bigrams/trigrams for intent like "situational awareness", "camera request".
    for n in (3, 2):
        for i in range(0, max(0, len(tokens) - n + 1)):
            candidate = " ".join(tokens[i:i + n])
            if len(candidate) >= 8:
                phrases.append(candidate)
            if len(phrases) >= max_phrases:
                break
        if len(phrases) >= max_phrases:
            break

    deduped: List[str] = []
    seen = set()
    for item in phrases:
        if item in seen:
            continue
        seen.add(item)
        deduped.append(item)
        if len(deduped) >= max_phrases:
            break
    return deduped


def _normalize_compact_text(text: str) -> str:
    normalized = _normalize_question(text or "")
    return re.sub(r"\s+", " ", normalized).strip()


def lexical_boost_score(question: str, doc: str, meta: dict) -> float:
    """
    Simple, general-purpose lexical score.
    Avoid domain-specific or collection-specific biases; reward token overlap only.
    """
    if doc is None:
        doc = ""
    
    q = (question or "").lower()
    d = (doc or "").lower()
    source = str(meta.get("source", "")).lower()
    score = 0.0

    # Normalized overlap reward with light stemming for singular/plural variants.
    q_tokens = set(_tokenize_normalized(q))
    d_tokens = set(_tokenize_normalized(d))
    matched = len(q_tokens & d_tokens)
    if matched > 0:
        score += min(4.0, matched * 0.65)

    # Exact key phrase match is a strong relevance signal for long manuals.
    key_phrases = _extract_key_phrases(q)
    phrase_hits = 0
    for phrase in key_phrases:
        if phrase in d:
            phrase_hits += 1
    if phrase_hits > 0:
        score += min(3.0, phrase_hits * 1.2)

    normalized_question = _normalize_compact_text(q)
    normalized_doc = _normalize_compact_text(d)
    if normalized_question and normalized_doc:
        if normalized_question in normalized_doc:
            score += 6.0
        elif len(normalized_question) >= 24:
            question_terms = normalized_question.split()
            for size in range(min(6, len(question_terms)), 2, -1):
                spans = [" ".join(question_terms[i:i + size]) for i in range(0, len(question_terms) - size + 1)]
                if any(span and span in normalized_doc for span in spans):
                    score += min(4.0, size * 0.7)
                    break

    # Prioritize chunks where heading lines contain query terms.
    heading_hits = 0
    heading_lines = re.findall(r"(?m)^\s*(?:\d+(?:\.\d+)*\.?\s+)?[^\n]{3,120}$", d)
    if heading_lines and q_tokens:
        for line in heading_lines[:10]:
            line_tokens = set(_tokenize_normalized(line))
            if line_tokens and len(q_tokens & line_tokens) >= 2:
                heading_hits += 1
    if heading_hits > 0:
        score += min(2.0, heading_hits * 0.8)

    # Source name match
    for tok in q_tokens:
        if tok in source:
            score += 1.0
    
    return score


def _content_similarity(text1: str, text2: str) -> float:
    """
    Calculate similarity between two texts based on overlapping tokens.
    Returns a value between 0 and 1.
    """
    if not text1 or not text2:
        return 0.0
    
    tokens1 = set([t.lower() for t in text1.split() if len(t) > 3])
    tokens2 = set([t.lower() for t in text2.split() if len(t) > 3])
    
    if not tokens1 or not tokens2:
        return 0.0
    
    intersection = len(tokens1 & tokens2)
    union = len(tokens1 | tokens2)
    
    return intersection / union if union > 0 else 0.0


def _smart_deduplicate_and_diversify(
    docs: List[str],
    metas: List[dict],
    distances: Optional[List[float]] = None,
    ids: Optional[List[str]] = None,
    similarity_threshold: float = 0.65,
    max_chunks_per_source: int = 5,
) -> Tuple[List[str], List[dict], List[float], List[str]]:
    """
    Intelligent deduplication:
    1. Removes near-duplicate documents (content similarity > threshold)
    2. Limits chunks per source while allowing more for comprehensive topics
    3. Prefers diverse sources for better context
    4. Adaptive thresholds based on document characteristics
    
    Returns deduplicated and diversified lists.
    """
    if not docs:
        return docs, metas, distances or [], ids or []
    
    print(f"[DEDUP] Starting deduplication: {len(docs)} docs, similarity_threshold={similarity_threshold}")
    
    selected_items = []
    seen_contents = []  # Track content we've already selected
    source_chunk_count = {}  # Track how many chunks from each source
    
    for idx, (doc, meta, dist, id_val) in enumerate(zip(
        docs, metas, distances or [None]*len(docs), ids or [None]*len(docs)
    )):
        if not doc:
            continue
        
        source = meta.get("source", "unknown") if meta else "unknown"
        source_key = source  # Use full source as key
        max_allowed = max_chunks_per_source
        
        # Check if we've already included max chunks from this source
        if source_key in source_chunk_count and source_chunk_count[source_key] >= max_allowed:
            print(f"[DEDUP] Skipping doc {idx} (source={source_key} reached max={max_allowed})")
            continue
        
        # Check for content similarity with already selected docs
        is_duplicate = False
        doc_len = len(doc.split())
        
        for seen_doc in seen_contents:
            seen_len = len(seen_doc.split())
            # For very short docs (<50 words), use stricter threshold; for long docs, more lenient
            adaptive_threshold = similarity_threshold
            if doc_len < 50 or seen_len < 50:
                adaptive_threshold = 0.75  # Stricter for short snippets
            elif doc_len > 500 or seen_len > 500:
                adaptive_threshold = 0.60  # More lenient for comprehensive docs
            
            sim = _content_similarity(doc, seen_doc)
            if sim > adaptive_threshold:
                print(f"[DEDUP] Skipping doc {idx} (similarity={sim:.3f} > {adaptive_threshold:.3f})")
                is_duplicate = True
                break
        
        if is_duplicate:
            continue
        
        # This doc is acceptable
        selected_items.append({
            "doc": doc,
            "meta": meta or {},
            "dist": dist,
            "id": id_val,
            "source": source,
        })
        seen_contents.append(doc)
        source_chunk_count[source_key] = source_chunk_count.get(source_key, 0) + 1
    
    dedup_docs = [item["doc"] for item in selected_items]
    dedup_metas = [item["meta"] for item in selected_items]
    dedup_distances = [item["dist"] for item in selected_items]
    dedup_ids = [item["id"] for item in selected_items]
    
    print(f"[DEDUP] After deduplication: {len(dedup_docs)} docs (removed {len(docs) - len(dedup_docs)} duplicates/chunks)")
    return dedup_docs, dedup_metas, dedup_distances, dedup_ids


def rerank_results(
    question: str,
    docs: List[str],
    metas: List[dict],
    distances: Optional[List[float]] = None,
    ids: Optional[List[str]] = None,
    max_chunks_per_source: int = 5,
) -> Tuple[List[str], List[dict], List[float], List[str]]:
    if not docs:
        print("[RERANK] No docs to rerank")
        return docs, metas, distances or [], ids or []

    print(f"[RERANK] Reranking {len(docs)} docs for question='{question[:120]}...'")
    
    items = []
    for idx, (d, m) in enumerate(zip(docs, metas)):
        # Safety: skip None documents
        if d is None:
            print(f"[RERANK] Skipping None doc at idx={idx}")
            continue
        meta = m or {}
        s = lexical_boost_score(question, d, m)
        dist = None if distances is None or idx >= len(distances) else distances[idx]
        idv = None if ids is None or idx >= len(ids) else ids[idx]
        items.append({
            "doc": d,
            "meta": meta,
            "score": s,
            "idx": idx,
            "dist": dist,
            "id": idv,
        })

    if all(it["score"] == 0 for it in items):
        print("[RERANK] All scores=0 → sorting by vector distance (balanced across collections)")
        # Still print a summary of results
        for it in items[:10]:
            print(f"[RERANK] src={it['meta'].get('source')} chunk={it['meta'].get('chunk')} id={it['id']} dist={it['dist']}")

        items.sort(key=lambda x: ((x["dist"] if x["dist"] is not None else 999.0), x["idx"]))
        
        # Even with 0 score, apply deduplication
        reranked_docs = [it["doc"] for it in items]
        reranked_metas = [it["meta"] for it in items]
        reranked_distances = [it["dist"] for it in items]
        reranked_ids = [it["id"] for it in items]
        return _smart_deduplicate_and_diversify(
            reranked_docs, reranked_metas, reranked_distances, reranked_ids,
            max_chunks_per_source=max_chunks_per_source,
        )

    items.sort(
        key=lambda x: (
            -x["score"],
            (x["dist"] if x["dist"] is not None else 999.0),
            x["idx"],
        )
    )

    # Safeguard for very specific questions: keep the strongest lexical-overlap chunks first.
    overlap_ranked = sorted(
        items,
        key=lambda x: _lexical_overlap_ratio(question, x["doc"]),
        reverse=True,
    )
    pinned = [it for it in overlap_ranked[:4] if _lexical_overlap_ratio(question, it["doc"]) >= 0.18]
    if pinned:
        pinned_ids = {it["id"] for it in pinned}
        items = pinned + [it for it in items if it["id"] not in pinned_ids]

    print("[RERANK] Top 10 after rerank (score, dist, id, source, chunk):")
    for it in items[:10]:
        print(f"         score={it['score']:.2f} dist={it['dist']} id={it['id']} source={it['meta'].get('source')} chunk={it['meta'].get('chunk')}")

    reranked_docs = [it["doc"] for it in items]
    reranked_metas = [it["meta"] for it in items]
    reranked_distances = [it["dist"] for it in items]
    reranked_ids = [it["id"] for it in items]
    
    # Apply comprehensive deduplication and diversification
    return _smart_deduplicate_and_diversify(
        reranked_docs, reranked_metas, reranked_distances, reranked_ids,
        max_chunks_per_source=max_chunks_per_source,
    )


def _lexical_overlap_ratio(question: str, doc: str) -> float:
    if not question or not doc:
        return 0.0
    # Extra safety: handle None explicitly
    if doc is None or question is None:
        return 0.0
    q_tokens = set(_tokenize_normalized(str(question)))
    d_tokens = set(_tokenize_normalized(str(doc)))
    if not q_tokens or not d_tokens:
        return 0.0
    inter = len(q_tokens & d_tokens)
    return inter / max(1, len(q_tokens))


def _is_short_query(question: str) -> bool:
    """Detect if query is short (likely needs lenient matching)."""
    word_count = len(question.strip().split())
    return word_count <= 8


def _extract_query_keywords(question: str, max_terms: int = 8) -> List[str]:
    """Extract compact keyword terms from a user question for fallback retrieval."""
    if not question:
        return []
    stopwords = {
        "the", "is", "are", "a", "an", "and", "or", "to", "of", "in", "on", "for", "with", "by",
        "from", "as", "at", "be", "this", "that", "these", "those", "it", "its", "if", "then",
        "how", "what", "when", "where", "why", "who", "can", "could", "should", "would", "do",
        "does", "did", "i", "we", "you", "they", "he", "she", "my", "our", "your", "their",
        "about", "into", "through", "using", "use", "want", "need", "like", "please",
    }
    tokens = re.findall(r"[a-zA-Z0-9_\-]+", question.lower())
    ranked = [t for t in tokens if len(t) > 2 and t not in stopwords]
    seen = set()
    deduped = []
    for term in ranked:
        if term in seen:
            continue
        seen.add(term)
        deduped.append(term)
        if len(deduped) >= max_terms:
            break
    return deduped


def _build_fallback_query_variants(question: str) -> List[str]:
    """Build generic fallback variants for broader recall across any topic."""
    if not question:
        return []
    variants: List[str] = []
    base = question.strip()
    if base:
        variants.append(base)

    normalized = _normalize_question(base)
    if normalized and normalized not in variants:
        variants.append(normalized)

    # Split compound questions into smaller clauses for better retrieval hit rate
    clauses = re.split(r"\?|\.|,|\band\b|\bthen\b|\balso\b", base, flags=re.IGNORECASE)
    for clause in clauses:
        c = clause.strip()
        if len(c.split()) >= 2 and c not in variants:
            variants.append(c)

    keywords = _extract_query_keywords(base)
    if keywords:
        keyword_query = " ".join(keywords)
        if keyword_query and keyword_query not in variants:
            variants.append(keyword_query)

    return variants[:3]


def _context_confidence_profile(
    question: str,
    docs: List[str],
    distances: Optional[List[float]],
) -> dict:
    """Compute topic-agnostic retrieval confidence from semantic and lexical signals."""
    if not docs:
        return {
            "confidence": 0.0,
            "best_distance": None,
            "max_overlap": 0.0,
            "is_short": _is_short_query(question),
        }

    best_distance = min(distances) if distances else None
    top_docs = [d for d in docs[:5] if d]
    max_overlap = max((_lexical_overlap_ratio(question, d) for d in top_docs), default=0.0)

    # Distance score: lower distance => higher confidence
    if best_distance is None:
        distance_score = 0.35
    else:
        distance_score = 1.0 - min(max(best_distance, 0.0), 0.95) / 0.95

    # Overlap score saturates once overlap is reasonably meaningful
    overlap_score = min(max_overlap / 0.22, 1.0)

    # Blend scores with slight leniency for short follow-up style questions
    is_short = _is_short_query(question)
    confidence = (0.7 * distance_score) + (0.3 * overlap_score)
    if is_short:
        confidence = max(confidence, 0.45 * distance_score + 0.55 * overlap_score)

    return {
        "confidence": round(max(0.0, min(1.0, confidence)), 4),
        "best_distance": best_distance,
        "max_overlap": round(max_overlap, 4),
        "is_short": is_short,
    }


def _normalize_history_messages(raw_messages: Optional[List[Dict[str, Any]]]) -> List[Dict[str, str]]:
    """Normalize user-provided history messages to role/content pairs."""
    if not raw_messages:
        return []
    normalized: List[Dict[str, str]] = []
    for msg in raw_messages:
        role = str((msg or {}).get("role") or "").strip().lower()
        content = str((msg or {}).get("content") or "").strip()
        if role not in {"user", "assistant"} or not content:
            continue
        normalized.append({"role": role, "content": content})
    return normalized


def _format_history_for_prompt(messages: List[Dict[str, str]], max_messages: int = 8) -> str:
    """Format recent conversation history for retrieval and prompt context.

    Only the user's prior questions are included. Prior assistant answers are
    intentionally omitted so the model cannot copy stale wording from earlier
    turns (which caused old/incorrect facts to leak into new answers even when
    retrieval was correct). The user turns are still enough for follow-up
    pronoun resolution ("what about for SQL?", "explain that more").
    """
    if not messages:
        return ""
    user_msgs = [m for m in messages if m.get("role") == "user"]
    if not user_msgs:
        return ""
    tail = user_msgs[-max_messages:]
    lines = [f"User: {m.get('content', '')}" for m in tail]
    return "\n".join(lines)


def _question_looks_like_followup(question: str) -> bool:
    candidate = str(question or "").strip().lower()
    if not candidate:
        return False

    followup_starters = (
        "and ", "also ", "what about", "how about", "why ", "then ", "so ",
        "can you elaborate", "can you explain more", "tell me more", "continue",
        "based on that", "for that", "for this", "regarding that", "regarding this",
    )
    if candidate.startswith(followup_starters):
        return True

    referential_patterns = [
        r"\bthat\b",
        r"\bthis\b",
        r"\bthese\b",
        r"\bthose\b",
        r"\bit\b",
        r"\bthey\b",
        r"\bthem\b",
        r"\bthe previous\b",
        r"\bthe last\b",
        r"\bearlier\b",
        r"\babove\b",
        r"\bbefore\b",
    ]
    short_question = len(candidate) <= 120
    if short_question and any(re.search(pattern, candidate) for pattern in referential_patterns):
        return True

    # Standalone questions with explicit nouns/topics should not inherit prior context.
    return False


def _should_use_history_for_question(
    question: str,
    stored_history: List[Dict[str, str]],
    explicit_is_followup: bool = False,
) -> bool:
    if not stored_history:
        return False

    if _question_looks_like_followup(question):
        return True

    if not explicit_is_followup:
        return False

    if _question_has_entity_fact_intent(question) or _question_requests_broad_coverage(question):
        return False

    explicit_topic_tokens = [
        token for token in _tokenize_normalized(question)
        if token not in _LEXICAL_STOPWORDS and token not in {"how", "step", "steps", "way", "ways", "method", "methods"}
    ]
    return len(explicit_topic_tokens) < 3


def _resolve_question_with_history(
    question: str,
    stored_history: List[Dict[str, str]],
    explicit_is_followup: bool = False,
) -> Dict[str, Any]:
    """
    Decide whether the current question depends on recent turns and, if needed,
    rewrite it into a standalone retrieval-friendly question.
    """
    original = str(question or "").strip()
    heuristic_followup = _should_use_history_for_question(
        original,
        stored_history,
        explicit_is_followup=explicit_is_followup,
    )
    result: Dict[str, Any] = {
        "use_history": bool(heuristic_followup),
        "relation": "related" if heuristic_followup else "standalone",
        "confidence": 0.55 if heuristic_followup else 0.7,
        "resolved_question": original,
        "reason": "heuristic",
    }

    if not original or not stored_history:
        return result

    # Keep the linkage window short so a topic shift mid-conversation drops
    # earlier topics out of the resolver's view. Pronoun-style follow-ups
    # ("is there a way from the UI?", "what about deleting it?") only ever
    # refer to the immediately preceding turn, not turns from 5+ messages ago.
    history_text = _format_history_for_prompt(stored_history, max_messages=2)
    if not history_text or not OPENAI_API_KEY:
        return result

    prompt = (
        "You are a conversation linkage resolver for a support chatbot.\n"
        "Decide if the CURRENT QUESTION should be interpreted using recent conversation history.\n"
        "Return strict JSON only with keys: relation, use_history, confidence, resolved_question, reason.\n\n"
        "Rules:\n"
        "1) relation must be one of: dependent, related, standalone.\n"
        "2) use_history=true for dependent or related; false for standalone.\n"
        "3) If dependent/related, resolved_question must rewrite the current question into a clear standalone question by adding only context explicitly present in history.\n"
        "4) If standalone, resolved_question should equal the current question.\n"
        "5) Do not invent facts or topics not present in history.\n"
        "6) confidence must be a number between 0 and 1.\n"
        "7) Keep reason short (max 20 words).\n"
        "8) IMPORTANT — incomplete-question detection: if the CURRENT QUESTION lacks a clear topic noun or object — i.e. it cannot be answered without referring back to a previous turn (examples of this pattern include but are not limited to: 'is there another way?', 'is there a way from X?', 'how about Y?', 'what about Z?', 'can I do it differently?', 'and from W?', 'why?', 'why not?') — you MUST classify it as dependent and resolve it against the MOST RECENT user turn in history (the last entry shown). Topic carries forward only from the immediately previous turn, never from earlier turns.\n"
        "9) If the question contains its own clear topic noun and is grammatically self-contained, classify it as standalone even if it is on a related theme — a fresh, fully-formed question is a topic shift, not a follow-up.\n\n"
        f"Explicit follow-up flag: {str(bool(explicit_is_followup)).lower()}\n\n"
        f"RECENT HISTORY (oldest to newest, only user turns):\n{history_text}\n\n"
        f"CURRENT QUESTION:\n{original}\n"
    )

    try:
        verdict_text = call_llm(prompt, temperature=0.0, model=OPENAI_RERANK_MODEL)
        try:
            parsed = json.loads(verdict_text)
        except json.JSONDecodeError:
            match = re.search(r"\{[\s\S]*\}", verdict_text)
            if not match:
                return result
            parsed = json.loads(match.group(0))

        relation = str(parsed.get("relation") or "").strip().lower()
        if relation not in {"dependent", "related", "standalone"}:
            relation = "related" if bool(parsed.get("use_history")) else "standalone"

        llm_use_history = bool(parsed.get("use_history"))
        if relation == "standalone":
            llm_use_history = False
        elif relation in {"dependent", "related"}:
            llm_use_history = True

        try:
            confidence = float(parsed.get("confidence", 0.0))
        except Exception:
            confidence = 0.0
        confidence = max(0.0, min(1.0, confidence))

        resolved_question = str(parsed.get("resolved_question") or "").strip() or original
        reason = str(parsed.get("reason") or "llm").strip() or "llm"

        if relation == "standalone" and heuristic_followup and confidence < 0.55:
            return result

        return {
            "use_history": llm_use_history,
            "relation": relation,
            "confidence": confidence,
            "resolved_question": resolved_question,
            "reason": reason,
        }
    except Exception as exc:
        print(f"[API][CHAT][WARN] History resolution failed; fallback to heuristic: {exc}")
        return result


_GENERIC_ANCHOR_TERMS = {
    "ticket", "issue", "problem", "error", "bug", "case", "item", "thing", "things", "one", "ones",
    "service", "system", "module", "component", "file", "document", "doc", "guide", "article",
    "log", "attachment", "image", "screenshot", "request", "record", "row", "column", "field",
    "setting", "settings", "configuration", "config", "workflow", "task", "job", "result", "answer",
}

_REFERENTIAL_ANCHOR_PATTERNS = [
    r"\bthis\b",
    r"\bthat\b",
    r"\bthese\b",
    r"\bthose\b",
    r"\bit\b",
    r"\bthey\b",
    r"\bthem\b",
    r"\bcurrent\b",
    r"\bselected\b",
    r"\bprevious\b",
    r"\blast\b",
    r"\bearlier\b",
    r"\bsame\b",
    r"\babove\b",
    r"\bbefore\b",
    r"\bwhich one\b",
    r"\bwhat about\b",
    r"\bhow about\b",
]


def _topic_bearing_tokens(question: str) -> List[str]:
    generic_fillers = _GENERIC_ANCHOR_TERMS | {
        "how", "what", "why", "which", "where", "when", "who", "tell", "show", "explain", "help",
        "solve", "fix", "use", "used", "using", "refer", "refers", "referred", "talk", "talking",
        "mean", "means", "meant", "about", "last", "previous", "current", "selected", "exact",
    }
    tokens: List[str] = []
    for token in _tokenize_normalized(question):
        if not token or token in _LEXICAL_STOPWORDS or token in generic_fillers or token.isdigit():
            continue
        tokens.append(token)
    return tokens


def _question_needs_anchor(question: str) -> bool:
    candidate = str(question or "").strip().lower()
    if not candidate:
        return False

    has_referential_language = any(re.search(pattern, candidate) for pattern in _REFERENTIAL_ANCHOR_PATTERNS)
    if not has_referential_language:
        return False

    topic_tokens = _topic_bearing_tokens(candidate)
    return len(topic_tokens) < 2


def _question_requests_previous_answer_source(question: str) -> bool:
    candidate = str(question or "").strip().lower()
    if not candidate:
        return False

    patterns = [
        r"\bwhich\s+(ticket|source|document|doc|guide|file|case|issue|item)\b.*\b(referring to|refer to|talking about|used|meant)",
        r"\bwhat\s+(ticket|source|document|doc|guide|file|case|issue|item)\b.*\b(referring to|refer to|talking about|used|meant)",
        r"\bwhat are you referring to\b",
        r"\bwhat do you mean\b",
        r"\bwhich one are you talking about\b",
        r"\bwhat was your source\b",
        r"\bwhich source did you use\b",
        r"\bwhere did that come from\b",
    ]
    return any(re.search(pattern, candidate) for pattern in patterns)


def _generate_llm_meta_response(question: str, scenario: str, facts: List[str]) -> str:
    if not OPENAI_API_KEY:
        raise RuntimeError("OPENAI_API_KEY is not configured; cannot generate chat response.")

    fact_block = "\n".join(f"- {fact}" for fact in facts if str(fact or "").strip()) or "- No additional facts provided."
    prompt = (
        "You are a grounded support assistant replying directly to the user.\n"
        "Generate the user-facing answer using ONLY the facts below.\n"
        "Do not invent details, hidden reasoning, sources, or context.\n"
        "Do not mention prompts, policies, internal state, retrieval, or chain-of-thought.\n"
        "Keep the answer concise, natural, and actionable.\n"
        "If the facts indicate missing context, clearly say what is missing and ask for the exact item needed.\n"
        "If the facts identify prior grounding sources, explain them plainly and specifically.\n\n"
        f"Scenario: {scenario}\n"
        f"User question: {question}\n\n"
        f"Facts:\n{fact_block}\n\n"
        "Return only the final answer to the user."
    )
    return (call_llm(prompt, temperature=0.1) or "").strip()


def _build_anchor_clarification_answer(question: str) -> str:
    candidate = str(question or "").lower()
    target_label = "subject"
    for token in (
        "ticket", "issue", "problem", "service", "file", "document", "guide", "log", "attachment",
        "image", "screenshot", "request", "record", "row", "column", "field", "configuration",
    ):
        if re.search(rf"\b{re.escape(token)}s?\b", candidate):
            target_label = token
            break

    facts = [
        "The current user question is referential or ambiguous.",
        "There is no active grounding anchor from a selected ticket, uploaded file, or grounded prior answer.",
    ]
    if target_label == "subject":
        facts.append("The missing anchor is a specific subject, item, document, component, or context.")
    else:
        facts.append(f"The missing anchor is the exact {target_label} the user means.")
        facts.append(f"The user can clarify by naming the exact {target_label}, selecting it in the UI, or restating the question with that subject.")
    return _generate_llm_meta_response(question, "missing-anchor", facts)


def _humanize_conversation_source(source: str) -> str:
    raw_source = str(source or "").strip()
    if not raw_source:
        return ""

    if raw_source.startswith("azure-devops:"):
        match = re.search(r"azure-devops:(\d+)", raw_source)
        if match:
            return f"selected Azure DevOps ticket {match.group(1)}"
        return "selected Azure DevOps ticket"

    learn_match = re.search(r"(?:^|::)ado:learn:(\d+)", raw_source)
    if learn_match:
        return f"knowledge-base ticket {learn_match.group(1)}"

    compact = raw_source.split(" (", 1)[0]
    if "::" in compact:
        compact = compact.split("::", 1)[1]
    return _display_source_name(compact)


def _build_previous_answer_source_answer(question: str, conversation_state: Optional[dict]) -> str:
    state = dict(conversation_state or {})
    sources = [str(item).strip() for item in list(state.get("sources") or []) if str(item).strip()]
    selected_ticket_id = state.get("selected_ticket_id")
    used_selected_ticket = bool(state.get("used_selected_ticket"))
    facts: List[str] = ["The user is asking what grounded my previous answer."]

    if used_selected_ticket and selected_ticket_id:
        facts.append(f"The selected Azure DevOps ticket {selected_ticket_id} was used as grounding context.")
        supplemental = [source for source in sources if not str(source).startswith("azure-devops:")]
        if supplemental:
            formatted = [_humanize_conversation_source(source) for source in supplemental[:3]]
            facts.append("Supplemental grounding sources were: " + ", ".join(formatted) + ".")
        return _generate_llm_meta_response(question, "previous-answer-source", facts)

    if sources:
        formatted_sources = [_humanize_conversation_source(source) for source in sources[:4]]
        if len(formatted_sources) == 1:
            facts.append(f"There was one recorded grounding source: {formatted_sources[0]}.")
        else:
            facts.append("Recorded grounding sources were: " + ", ".join(formatted_sources) + ".")
        return _generate_llm_meta_response(question, "previous-answer-source", facts)

    facts.append("There is no specific grounded source recorded for the previous answer.")
    facts.append("Ask the user to provide the exact subject, select the relevant item, or restate the question with explicit context.")
    return _generate_llm_meta_response(question, "previous-answer-source", facts)


def _store_conversation_turn(
    conversation_key: str,
    question: str,
    answer: str,
    *,
    sources: Optional[List[str]] = None,
    selected_ticket_id: Optional[int] = None,
    selected_ticket_url: Optional[str] = None,
    used_selected_ticket: bool = False,
    used_file_context: bool = False,
) -> None:
    try:
        conversation_store[conversation_key].append({"role": "user", "content": question})
        conversation_store[conversation_key].append({"role": "assistant", "content": answer})
        conversation_state_store[conversation_key] = {
            "sources": list(sources or []),
            "selected_ticket_id": selected_ticket_id,
            "selected_ticket_url": selected_ticket_url,
            "used_selected_ticket": bool(used_selected_ticket),
            "used_file_context": bool(used_file_context),
            "has_grounded_sources": bool(sources),
            "updated_at": datetime.utcnow().isoformat() + "Z",
        }
    except Exception as exc:
        print(f"[API][CHAT][WARN] Failed to persist conversation context/state: {exc}")


def _looks_like_noncode_fenced_block(language: str, content: str) -> bool:
    lang = str(language or "").strip().lower()

    lines = [line.rstrip() for line in str(content or "").strip().splitlines() if line.strip()]
    if not lines:
        return False

    # Even for "code" languages like bash/shell, detect blocks that are purely
    # UI/navigation instructions disguised as code comments (no real commands).
    if lang in {"bash", "sh", "shell", "zsh", "powershell", "ps1"}:
        _real_cmd_count = 0
        _ui_comment_count = 0
        for _ln in lines:
            _stripped = _ln.strip()
            # Lines that are actual commands (not comments)
            if _stripped and not re.match(r"^\s*#", _stripped):
                _real_cmd_count += 1
            elif re.match(r"^\s*#", _stripped):
                _body = re.sub(r"^\s*#\s*", "", _stripped)
                if re.search(
                    r"\b(click|select|open|navigate|button|tab|screen|ui|portainer|"
                    r"admin studio|field|dropdown|checkbox|page|section|restart|restarting|"
                    r"service|clearing|browser|cache|setting|menu|dialog|typically|usually|"
                    r"involves|done through|via the|this can|you can|access)\b",
                    _body, re.IGNORECASE,
                ):
                    _ui_comment_count += 1
        # If there are NO real commands and at least some UI comments, it's not code
        if _real_cmd_count == 0 and _ui_comment_count > 0:
            return True
        # If it has real commands, keep it as code
        return False

    if lang in {"python", "py", "sql", "json", "yaml", "yml", "xml", "javascript", "js", "typescript", "ts", "html", "css", "ini", "toml", "diff", "dockerfile"}:
        return False
    if not lines:
        return False

    code_like_count = sum(1 for line in lines if _looks_like_command_or_code_line(line))
    ui_like_count = sum(
        1
        for line in lines
        if re.search(r"\b(click|select|open|navigate|button|tab|screen|ui|portainer|admin studio|field|dropdown|checkbox|page|section)\b", line, re.IGNORECASE)
    )
    key_value_count = sum(1 for line in lines if re.match(r"^[A-Za-z][A-Za-z0-9_ /()-]{1,40}:\s+.+$", line))

    if code_like_count == 0 and (ui_like_count > 0 or key_value_count > 0):
        return True
    if code_like_count <= max(1, len(lines) // 4) and lang in {"", "text", "plaintext", "plain", "md", "markdown"}:
        return True
    return False


def _normalize_noncode_fenced_blocks(answer: str) -> str:
    if not answer or "```" not in answer:
        return answer

    pattern = re.compile(r"```([A-Za-z0-9_+-]*)\n([\s\S]*?)```", re.MULTILINE)

    def _replace(match: re.Match) -> str:
        language = match.group(1) or ""
        content = (match.group(2) or "").strip("\n")
        if not _looks_like_noncode_fenced_block(language, content):
            return match.group(0)

        lines = [line.strip() for line in content.splitlines() if line.strip()]
        if not lines:
            return ""
        # Strip leading # from comment-only lines (UI instructions that were in bash blocks)
        cleaned = []
        for line in lines:
            stripped = re.sub(r"^\s*#\s*", "", line).strip()
            cleaned.append(stripped if stripped else line)
        if all(re.match(r"^[A-Za-z][A-Za-z0-9_ /()-]{1,40}:\s+.+$", cl) for cl in cleaned):
            return "\n".join(f"- {cl}" for cl in cleaned)
        return "  \n".join(cleaned)

    return pattern.sub(_replace, answer)


def _get_vector_collection(collection_name: str):
    normalized_name = (collection_name or "").strip()
    if normalized_name == COLLECTION_NAME:
        return collection
    if normalized_name == MEMORY_COLLECTION_NAME:
        return memory_collection
    raise HTTPException(status_code=400, detail=f"Unknown collection: {collection_name}")


def _vector_collection_description(collection_name: str) -> str:
    if collection_name == COLLECTION_NAME:
        return "Wiki + User Guide + Old Tickets Summeries"
    if collection_name == MEMORY_COLLECTION_NAME:
        return "New User Knowledge"
    return "Vector collection"


def _safe_embedding_to_list(embedding: Any) -> List[float]:
    if embedding is None:
        return []
    if hasattr(embedding, "tolist"):
        embedding = embedding.tolist()
    return [float(value) for value in embedding]


def _display_source_name(source: str) -> str:
    source_text = str(source or "").strip()
    if not source_text:
        return "Untitled document"
    return os.path.basename(unquote(source_text)) or source_text


def _normalize_vector_input_type(value: Any) -> Optional[str]:
    normalized = str(value or "").strip().lower()
    if normalized in {"file", "upload", "uploaded_file", "uploaded-file"}:
        return "file"
    if normalized in {"free_text", "free-text", "text", "content", "manual"}:
        return "free_text"
    return None


def _looks_like_uploaded_file_content(text: str) -> bool:
    normalized_text = str(text or "").strip()
    if not normalized_text:
        return False
    if "---FILE SEPARATOR---" in normalized_text:
        return True
    if re.search(r"(?m)^\*\*File:\s*.+?\*\*(?:\s*\(.+?\))?$", normalized_text):
        return True
    if re.search(r"(?mi)^\[Image Content from .+\]$", normalized_text):
        return True
    return bool(re.search(r"(?i)\.(pdf|docx?|xlsx?|csv|png|jpe?g|gif|webp|txt|md)\b", normalized_text[:240]))


def _get_tokenizer_encoding():
    if not HAS_TIKTOKEN:
        return None

    for model_name in (OPENAI_EMBEDDING_MODEL, OPENAI_CHAT_MODEL):
        try:
            return tiktoken.encoding_for_model(model_name)
        except KeyError:
            continue

    return tiktoken.get_encoding("cl100k_base")


def _infer_vector_document_input_type(collection_name: str, source: str, items: List[dict], merged_content: str) -> Optional[str]:
    if collection_name != MEMORY_COLLECTION_NAME:
        return None

    seed_meta = dict((items[0].get("metadata") if items else {}) or {})
    is_user_knowledge = str(seed_meta.get("collection") or "").strip().lower() == "user_knowledge" or str(source or "").startswith("user:content:")
    if not is_user_knowledge:
        return None

    for item in items:
        safe_meta = item.get("metadata") or {}
        explicit_type = _normalize_vector_input_type(safe_meta.get("input_type") or safe_meta.get("source_type"))
        if explicit_type == "file":
            return explicit_type
        if explicit_type == "free_text" and not _looks_like_uploaded_file_content(merged_content):
            return explicit_type

    return "file" if _looks_like_uploaded_file_content(merged_content) else "free_text"


def _extract_document_heading(text: str) -> str:
    if not text:
        return ""
    for raw_line in text.splitlines()[:10]:
        line = raw_line.strip()
        if not line:
            continue
        file_match = re.match(r"^\*\*File:\s*(.+?)\*\*$", line)
        if file_match:
            return file_match.group(1).strip()
        if line.startswith("#"):
            return line.lstrip("#").strip()
    return ""


def _build_document_title(source: str, first_chunk_text: str) -> str:
    heading = _extract_document_heading(first_chunk_text)
    return heading or _display_source_name(source)


def _find_chunk_overlap(existing_text: str, next_chunk: str, max_overlap: int = CHUNK_OVERLAP) -> int:
    if not existing_text or not next_chunk:
        return 0
    max_len = min(max_overlap, len(existing_text), len(next_chunk))
    for overlap_size in range(max_len, 0, -1):
        if existing_text.endswith(next_chunk[:overlap_size]):
            return overlap_size
    return 0


def _merge_chunk_texts(chunks: List[str], max_chars: Optional[int] = None) -> str:
    if not chunks:
        return ""
    merged_parts: List[str] = []
    merged_length = 0
    merged_tail = ""

    for chunk in chunks:
        if not chunk:
            continue

        if not merged_parts:
            piece = chunk
        else:
            overlap_size = _find_chunk_overlap(merged_tail, chunk)
            piece = chunk[overlap_size:]

        if not piece:
            continue

        if max_chars is not None:
            remaining = max_chars - merged_length
            if remaining <= 0:
                break
            if len(piece) > remaining:
                piece = piece[:remaining]

        merged_parts.append(piece)
        merged_length += len(piece)
        merged_tail = (merged_tail + piece)[-CHUNK_OVERLAP:]

        if max_chars is not None and merged_length >= max_chars:
            break

    return "".join(merged_parts)


def _build_chunk_boundaries(rows: List[dict], preview_chars: int = 72) -> List[dict]:
    boundaries: List[dict] = []
    merged_length = 0
    merged_tail = ""

    for row in rows:
        chunk_text = row.get("content") or ""
        chunk_number = int(row.get("chunk", 0) or 0)
        if not chunk_text:
            boundaries.append(
                {
                    "chunk": chunk_number,
                    "start_position": merged_length,
                    "end_position": merged_length,
                    "visible_length": 0,
                    "original_length": 0,
                    "start_preview": "",
                    "end_preview": "",
                }
            )
            continue

        if not boundaries:
            overlap_size = 0
            visible_text = chunk_text
        else:
            overlap_size = _find_chunk_overlap(merged_tail, chunk_text)
            visible_text = chunk_text[overlap_size:]

        visible_length = len(visible_text)
        start_position = merged_length + 1 if visible_length > 0 else merged_length
        end_position = merged_length + visible_length

        start_preview = visible_text[:preview_chars].strip()
        end_preview = visible_text[-preview_chars:].strip() if visible_text else ""

        boundaries.append(
            {
                "chunk": chunk_number,
                "start_position": start_position,
                "end_position": end_position,
                "visible_length": visible_length,
                "original_length": len(chunk_text),
                "overlap_trimmed": overlap_size,
                "start_preview": start_preview,
                "end_preview": end_preview,
            }
        )

        if not visible_text:
            continue

        merged_length += visible_length
        merged_tail = (merged_tail + visible_text)[-CHUNK_OVERLAP:]

    return boundaries


def _vector_chunk_page_bounds(limit: int, offset: int) -> Tuple[int, int]:
    safe_limit = max(1, min(int(limit), VECTOR_DB_MAX_CHUNK_LIMIT))
    safe_offset = max(0, int(offset))
    return safe_limit, safe_offset


def _get_sorted_document_rows(
    collection_name: str,
    source: str,
    include_embeddings: bool = False,
) -> List[dict]:
    target_collection = _get_vector_collection(collection_name)
    include_fields = ["documents", "metadatas"]
    if include_embeddings:
        include_fields.append("embeddings")

    raw = target_collection.get(where={"source": source}, include=include_fields)
    ids = list(raw.get("ids", []) or [])
    docs = list(raw.get("documents", []) or [])
    metas = list(raw.get("metadatas", []) or [])
    embeddings = list(raw.get("embeddings", []) or []) if include_embeddings else []

    rows: List[dict] = []
    for idx, (doc_id, doc_text, meta) in enumerate(zip(ids, docs, metas)):
        safe_meta = meta or {}
        rows.append(
            {
                "id": doc_id,
                "chunk": int(safe_meta.get("chunk", 0) or 0),
                "content": doc_text or "",
                "metadata": safe_meta,
                "embedding": embeddings[idx] if include_embeddings and idx < len(embeddings) else None,
            }
        )

    rows.sort(key=lambda item: item["chunk"])
    return rows


def _build_chunk_page(rows: List[dict], limit: int, offset: int, include_embeddings: bool = False) -> Tuple[List[dict], int, int]:
    safe_limit, safe_offset = _vector_chunk_page_bounds(limit, offset)
    page_rows = rows[safe_offset:safe_offset + safe_limit]
    chunk_items: List[dict] = []
    for row in page_rows:
        chunk_items.append(
            {
                "id": row["id"],
                "chunk": row["chunk"],
                "content": row["content"],
                "metadata": row["metadata"],
                "embedding": _safe_embedding_to_list(row.get("embedding")) if include_embeddings else None,
                "embedding_loaded": bool(include_embeddings),
            }
        )
    return chunk_items, safe_limit, safe_offset


def _count_collection_documents(collection_name: str) -> int:
    target_collection = _get_vector_collection(collection_name)
    raw = target_collection.get(include=["metadatas"])
    sources = {
        str((meta or {}).get("source") or "").strip()
        for meta in list(raw.get("metadatas", []) or [])
        if str((meta or {}).get("source") or "").strip()
    }
    return len(sources)


def _count_text_tokens(text: str) -> int:
    raw_text = str(text or "")
    if not raw_text:
        return 0

    encoding = _get_tokenizer_encoding()
    if encoding is None:
        raise RuntimeError("tiktoken is required for exact token counting")

    return len(encoding.encode(raw_text, disallowed_special=()))


def _estimate_text_tokens(text: str) -> int:
    raw_text = str(text or "")
    if not raw_text:
        return 0
    try:
        return _count_text_tokens(raw_text)
    except Exception:
        return max(1, len(raw_text) // 4)


def _truncate_text_to_token_budget(text: str, max_tokens: int, label: str) -> str:
    raw_text = str(text or "")
    if not raw_text or max_tokens <= 0:
        return ""

    current_tokens = _estimate_text_tokens(raw_text)
    if current_tokens <= max_tokens:
        return raw_text

    notice = (
        f"\n\n[TRUNCATED {label}: original_tokens={current_tokens}, kept_tokens={max_tokens}]\n\n"
    )
    encoding = _get_tokenizer_encoding()
    if encoding is not None:
        encoded = encoding.encode(raw_text, disallowed_special=())
        head_count = max_tokens // 2
        tail_count = max_tokens - head_count
        if len(encoded) <= max_tokens:
            return raw_text
        head_text = encoding.decode(encoded[:head_count]).strip()
        tail_text = encoding.decode(encoded[-tail_count:]).strip() if tail_count > 0 else ""
        merged = head_text
        if tail_text:
            merged = f"{head_text}{notice}{tail_text}"
        return merged.strip()

    approx_chars = max_tokens * 4
    head_chars = approx_chars // 2
    tail_chars = approx_chars - head_chars
    head_text = raw_text[:head_chars].strip()
    tail_text = raw_text[-tail_chars:].strip() if tail_chars > 0 else ""
    merged = head_text
    if tail_text:
        merged = f"{head_text}{notice}{tail_text}"
    return merged.strip()


def _fit_prompt_to_model_budget(
    question: str,
    combined_context: str,
    foundational_instruction: str = "",
    ta9_instruction: str = "",
) -> Tuple[str, str, str]:
    """Return (system_message, user_message, trimmed_context)."""
    def _total_tokens(sys_msg: str, usr_msg: str) -> int:
        return _estimate_text_tokens(sys_msg) + _estimate_text_tokens(usr_msg)

    system_msg, user_msg = _build_system_prompt(
        question=question,
        combined_context=combined_context,
        foundational_instruction=foundational_instruction,
        ta9_instruction=ta9_instruction,
    )
    prompt_tokens = _total_tokens(system_msg, user_msg)
    if prompt_tokens <= LLM_MAX_INPUT_TOKENS:
        return system_msg, user_msg, combined_context

    context_tokens = _estimate_text_tokens(combined_context)
    overflow_tokens = prompt_tokens - LLM_MAX_INPUT_TOKENS
    target_context_tokens = max(4000, context_tokens - overflow_tokens - 1024)
    trimmed_context = _truncate_text_to_token_budget(
        combined_context,
        target_context_tokens,
        "combined context for final LLM input",
    )
    system_msg, user_msg = _build_system_prompt(
        question=question,
        combined_context=trimmed_context,
        foundational_instruction=foundational_instruction,
        ta9_instruction=ta9_instruction,
    )

    second_pass_tokens = _total_tokens(system_msg, user_msg)
    if second_pass_tokens > LLM_MAX_INPUT_TOKENS:
        second_overflow = second_pass_tokens - LLM_MAX_INPUT_TOKENS
        second_budget = max(2000, target_context_tokens - second_overflow - 1024)
        trimmed_context = _truncate_text_to_token_budget(
            trimmed_context,
            second_budget,
            "combined context for final LLM input",
        )
        system_msg, user_msg = _build_system_prompt(
            question=question,
            combined_context=trimmed_context,
            foundational_instruction=foundational_instruction,
            ta9_instruction=ta9_instruction,
        )

    final_prompt_tokens = _total_tokens(system_msg, user_msg)
    if final_prompt_tokens > LLM_MAX_INPUT_TOKENS:
        print(
            f"[API][CHAT][WARN] Prompt is still larger than the model budget after trimming: "
            f"tokens={final_prompt_tokens} budget={LLM_MAX_INPUT_TOKENS}"
        )
    else:
        print(
            f"[API][CHAT] Prompt fitted to model budget: tokens={final_prompt_tokens} budget={LLM_MAX_INPUT_TOKENS}"
        )
    return system_msg, user_msg, trimmed_context


def _estimate_rows_token_count(rows: List[dict]) -> int:
    # Fast approximation for huge documents: avoids expensive tokenizer passes.
    return sum(len(str((row or {}).get("content") or "")) // 4 for row in rows)


def _count_collection_tokens(collection_name: str) -> int:
    summaries = _group_collection_documents(collection_name)
    return sum(int(item.get("token_count", 0) or 0) for item in summaries)


def _group_collection_documents(collection_name: str) -> List[dict]:
    target_collection = _get_vector_collection(collection_name)
    raw = target_collection.get(include=["metadatas", "documents"])
    ids = list(raw.get("ids", []) or [])
    docs = list(raw.get("documents", []) or [])
    metas = list(raw.get("metadatas", []) or [])

    grouped: Dict[str, dict] = {}
    for doc_id, doc_text, meta in zip(ids, docs, metas):
        safe_meta = meta or {}
        source = str(safe_meta.get("source") or doc_id or "")
        chunk_index = int(safe_meta.get("chunk", 0) or 0)
        entry = grouped.setdefault(
            source,
            {
                "source": source,
                "items": [],
                "updated_at": "",
            },
        )
        entry["items"].append(
            {
                "id": doc_id,
                "chunk": chunk_index,
                "document": doc_text or "",
                "metadata": safe_meta,
            }
        )

        candidate_timestamp = str(
            safe_meta.get("edited_at")
            or safe_meta.get("uploaded_at")
            or safe_meta.get("created_at")
            or ""
        )
        if candidate_timestamp and candidate_timestamp > entry["updated_at"]:
            entry["updated_at"] = candidate_timestamp

    summaries: List[dict] = []
    for source, payload in grouped.items():
        items = sorted(payload["items"], key=lambda item: item["chunk"])
        merged_content = _merge_chunk_texts([item["document"] for item in items])
        preview = re.sub(r"\s+", " ", merged_content).strip()
        if len(preview) > VECTOR_DB_PREVIEW_CHARS:
            preview = preview[:VECTOR_DB_PREVIEW_CHARS].rstrip() + "..."
        title = _build_document_title(source, items[0]["document"] if items else "")
        input_type = _infer_vector_document_input_type(collection_name, source, items, merged_content)
        summaries.append(
            {
                "source": source,
                "title": title,
                "display_source": _display_source_name(source),
                "preview": preview,
                "chunk_count": len(items),
                "token_count": _count_text_tokens(merged_content),
                "input_type": input_type,
                "updated_at": payload["updated_at"] or None,
                "_search_content": merged_content,
            }
        )

    summaries.sort(key=lambda item: ((item["updated_at"] or ""), item["title"].lower()), reverse=True)
    return summaries


def _normalize_vector_search_value(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "").lower()).strip()


def _build_vector_search_preview(content: str, query: str, fallback_preview: str) -> str:
    text = re.sub(r"\s+", " ", str(content or "")).strip()
    if not text:
        return fallback_preview

    raw_query = str(query or "").strip().lower()
    normalized_terms = [term for term in _tokenize_normalized(query) if term]
    lower_text = text.lower()

    hit_index = -1
    if raw_query:
        hit_index = lower_text.find(raw_query)
    if hit_index < 0:
        for term in normalized_terms:
            hit_index = lower_text.find(term.lower())
            if hit_index >= 0:
                break

    if hit_index < 0:
        return fallback_preview or (text[:VECTOR_DB_PREVIEW_CHARS].rstrip() + "..." if len(text) > VECTOR_DB_PREVIEW_CHARS else text)

    window = max(80, VECTOR_DB_PREVIEW_CHARS // 2)
    start = max(0, hit_index - window)
    end = min(len(text), hit_index + window)
    snippet = text[start:end].strip()

    if start > 0:
        snippet = "..." + snippet
    if end < len(text):
        snippet = snippet + "..."
    return snippet


def _score_vector_document_summary(item: dict, query: str) -> Optional[dict]:
    query_text = str(query or "").strip()
    normalized_query = _normalize_vector_search_value(query_text)
    query_terms = [term for term in _tokenize_normalized(query_text) if term]
    if not normalized_query and not query_terms:
        return None

    title = str(item.get("title") or "")
    source = str(item.get("source") or "")
    display_source = str(item.get("display_source") or "")
    search_content = str(item.get("_search_content") or "")

    title_normalized = _normalize_vector_search_value(title)
    source_normalized = _normalize_vector_search_value(f"{display_source} {source}")
    content_normalized = _normalize_vector_search_value(search_content)

    title_terms = set(_tokenize_normalized(title))
    source_terms = set(_tokenize_normalized(f"{display_source} {source}"))
    content_terms = set(_tokenize_normalized(search_content))
    combined_terms = title_terms | source_terms | content_terms

    matched_terms = [term for term in query_terms if term in combined_terms]
    if normalized_query:
        exact_in_title = normalized_query in title_normalized
        exact_in_source = normalized_query in source_normalized
        exact_in_content = normalized_query in content_normalized
    else:
        exact_in_title = False
        exact_in_source = False
        exact_in_content = False

    has_any_match = exact_in_title or exact_in_source or exact_in_content
    if not has_any_match:
        return None

    all_terms_matched = bool(query_terms) and len(matched_terms) == len(query_terms)
    score = 0.0

    if exact_in_title:
        score += 220.0
    if exact_in_source:
        score += 170.0
    if exact_in_content:
        score += 120.0

    title_term_hits = sum(1 for term in query_terms if term in title_terms)
    source_term_hits = sum(1 for term in query_terms if term in source_terms)
    content_term_hits = sum(1 for term in query_terms if term in content_terms)

    score += title_term_hits * 40.0
    score += source_term_hits * 24.0
    score += content_term_hits * 12.0

    if all_terms_matched:
        score += 90.0
    elif query_terms:
        score += (len(matched_terms) / len(query_terms)) * 25.0

    preview = _build_vector_search_preview(search_content, query_text, str(item.get("preview") or ""))
    return {
        "score": score,
        "all_terms_matched": all_terms_matched,
        "preview": preview,
    }


def _build_vector_document_detail(
    collection_name: str,
    source: str,
    chunk_limit: int = VECTOR_DB_DEFAULT_CHUNK_LIMIT,
    chunk_offset: int = 0,
    include_embeddings: bool = False,
    include_full_content: bool = True,
) -> dict:
    rows = _get_sorted_document_rows(
        collection_name,
        source,
        include_embeddings=include_embeddings,
    )
    if not rows:
        raise HTTPException(status_code=404, detail="Document not found")

    chunk_items, safe_limit, safe_offset = _build_chunk_page(
        rows,
        limit=chunk_limit,
        offset=chunk_offset,
        include_embeddings=include_embeddings,
    )
    chunk_boundaries = _build_chunk_boundaries(rows)
    merged_content: Optional[str] = None
    cached_payload: Optional[dict] = None
    if include_full_content:
        cached_payload = _get_cached_vector_document_payload(collection_name, source)
        if cached_payload is not None:
            merged_content = str(cached_payload.get("full_content") or "")
        else:
            merged_content = _merge_chunk_texts([row["content"] for row in rows])
    full_content = merged_content if include_full_content else None
    if merged_content is not None:
        if cached_payload is not None:
            token_count = int(cached_payload.get("token_count") or 0)
        else:
            token_count = _count_text_tokens(merged_content)
            _store_cached_vector_document_payload(collection_name, source, merged_content, token_count)
    else:
        token_count = _estimate_rows_token_count(rows)
    title = _build_document_title(source, rows[0]["content"] if rows else "")
    input_type = _infer_vector_document_input_type(
        collection_name,
        source,
        [{"metadata": row["metadata"]} for row in rows],
        merged_content or "",
    )
    updated_at = None
    for row in rows:
        safe_meta = row["metadata"] or {}
        candidate_timestamp = safe_meta.get("edited_at") or safe_meta.get("uploaded_at") or safe_meta.get("created_at")
        if candidate_timestamp and (updated_at is None or str(candidate_timestamp) > str(updated_at)):
            updated_at = str(candidate_timestamp)

    return {
        "collection_name": collection_name,
        "collection_description": _vector_collection_description(collection_name),
        "source": source,
        "title": title,
        "display_source": _display_source_name(source),
        "chunk_count": len(rows),
        "token_count": token_count,
        "input_type": input_type,
        "updated_at": updated_at,
        "full_content": full_content,
        "chunk_boundaries": chunk_boundaries,
        "chunk_limit": safe_limit,
        "chunk_offset": safe_offset,
        "chunk_has_more": safe_offset + safe_limit < len(rows),
        "chunks": chunk_items,
    }


def _build_updated_chunk_records(source: str, content: str, existing_metas: List[dict], existing_docs: Optional[List[str]] = None, existing_embeddings: Optional[List[List[float]]] = None) -> Tuple[List[str], List[str], List[dict], List[List[float]]]:
    clean_content = (content or "").strip()
    if not clean_content:
        raise HTTPException(status_code=400, detail="content is empty")

    seed_meta = dict((existing_metas or [{}])[0] or {})
    base_meta = {
        key: value
        for key, value in seed_meta.items()
        if key not in {"chunk", "chunk_type", "table_page", "table_index", "table_row_index"}
    }
    base_meta["source"] = source
    base_meta["edited_at"] = datetime.utcnow().isoformat() + "Z"

    if str(base_meta.get("collection") or "").strip().lower() == "user_knowledge" or str(source or "").startswith("user:content:"):
        current_input_type = _normalize_vector_input_type(base_meta.get("input_type"))
        if current_input_type != "file" and _looks_like_uploaded_file_content(clean_content):
            base_meta["input_type"] = "file"
        elif current_input_type is None:
            base_meta["input_type"] = "free_text"

    chunks = (
        chunk_text_preserve_table_rows(clean_content)
        if TABLE_ROW_START_TAG in clean_content or base_meta.get("input_type") == "file"
        else chunk_text(clean_content)
    )
    if not chunks:
        raise HTTPException(status_code=400, detail="No chunks generated from content")

    # Build a reuse pool from existing chunks: content_text -> list of embeddings
    reuse_pool: Dict[str, List[List[float]]] = {}
    if existing_docs and existing_embeddings and len(existing_docs) == len(existing_embeddings):
        for doc_text, emb in zip(existing_docs, existing_embeddings):
            if doc_text and emb is not None and (isinstance(emb, list) or len(emb) > 0):
                reuse_pool.setdefault(doc_text, []).append(list(emb) if not isinstance(emb, list) else emb)

    ids_to_add: List[str] = []
    docs_to_add: List[str] = []
    metas_to_add: List[dict] = []
    embeddings_to_add: List[List[float]] = []
    chunks_needing_embedding: List[Tuple[int, str]] = []  # (index, chunk_text)

    for chunk_index, chunk_text_value in enumerate(chunks):
        chunk_meta = dict(base_meta)
        chunk_meta["chunk"] = chunk_index
        if chunk_text_value.startswith(TABLE_ROW_START_TAG):
            chunk_meta["chunk_type"] = "table_row"
            chunk_meta.update(extract_table_row_metadata(chunk_text_value))
        else:
            chunk_meta["chunk_type"] = "text"
        ids_to_add.append(str(uuid.uuid4()))
        docs_to_add.append(chunk_text_value)
        metas_to_add.append(chunk_meta)

        # Check if we can reuse an existing embedding for this exact chunk content
        pool_list = reuse_pool.get(chunk_text_value)
        if pool_list:
            embeddings_to_add.append(pool_list.pop(0))
            if not pool_list:
                del reuse_pool[chunk_text_value]
        else:
            embeddings_to_add.append([])  # placeholder
            chunks_needing_embedding.append((chunk_index, chunk_text_value))

    reused_count = len(chunks) - len(chunks_needing_embedding)
    print(f"[VECTOR_DB][UPDATE] source={source} total_chunks={len(chunks)} reused={reused_count} need_embedding={len(chunks_needing_embedding)}")

    # Batch-embed only the changed/new chunks
    if chunks_needing_embedding:
        texts_to_embed = [text for _, text in chunks_needing_embedding]
        new_embeddings = embed_texts_batch(texts_to_embed)
        for (chunk_idx, _), embedding in zip(chunks_needing_embedding, new_embeddings):
            embeddings_to_add[chunk_idx] = embedding

    return ids_to_add, docs_to_add, metas_to_add, embeddings_to_add


def _resolve_collection_name_from_meta(meta: Optional[dict]) -> str:
    label = str((meta or {}).get("collection") or "").strip().lower()
    if label == "user_knowledge":
        return MEMORY_COLLECTION_NAME
    return COLLECTION_NAME


def _build_source_context_candidates(
    question: str,
    docs: List[str],
    metas: List[dict],
    distances: Optional[List[float]] = None,
    ids: Optional[List[str]] = None,
    max_sources: int = 6,
) -> List[dict]:
    """Group retrieved chunks by source document and rank sources by relevance."""
    if not docs:
        return []

    # --- Phase 1: group chunks by (collection, source) ---
    grouped: Dict[Tuple[str, str], dict] = {}
    for idx, (doc, meta) in enumerate(zip(docs, metas)):
        safe_meta = meta or {}
        source = str(safe_meta.get("source") or "").strip()
        if not source or not doc:
            continue

        collection_name = _resolve_collection_name_from_meta(safe_meta)
        key = (collection_name, source)
        distance = None if distances is None or idx >= len(distances) else distances[idx]
        chunk_score = lexical_boost_score(question, doc, safe_meta)
        overlap_score = _lexical_overlap_ratio(question, doc) * 5.0
        position_bonus = max(0.0, 3.0 - (idx * 0.12))
        distance_bonus = 0.0 if distance is None else max(0.0, 1.5 - (float(distance) * 3.0))

        entry = grouped.setdefault(
            key,
            {
                "collection_name": collection_name,
                "source": source,
                "best_distance": distance,
                "chunk_signal_scores": [],
                "seed_meta": safe_meta,
                "matched_chunk_ids": [],
                "matched_chunk_indices": [],
                "primary_matched_chunk_index": None,
                "chunk_id_scores": {},
            },
        )
        per_chunk_score = chunk_score + overlap_score + position_bonus + distance_bonus
        entry["chunk_signal_scores"].append(per_chunk_score)
        if distance is not None and (entry["best_distance"] is None or float(distance) < float(entry["best_distance"])):
            entry["best_distance"] = distance
        if ids is not None and idx < len(ids) and ids[idx] is not None:
            chunk_id = ids[idx]
            entry["matched_chunk_ids"].append(chunk_id)
            prev = entry["chunk_id_scores"].get(chunk_id)
            if prev is None or per_chunk_score > prev:
                entry["chunk_id_scores"][chunk_id] = per_chunk_score
        chunk_idx = safe_meta.get("chunk")
        if chunk_idx is not None:
            normalized_chunk_idx = int(chunk_idx)
            entry["matched_chunk_indices"].append(normalized_chunk_idx)
            if entry["primary_matched_chunk_index"] is None:
                entry["primary_matched_chunk_index"] = normalized_chunk_idx

    if not grouped:
        return []

    # --- Phase 2: preliminary rank by chunk signals, keep top candidates ---
    prelim_ranked = sorted(
        grouped.values(),
        key=lambda item: (
            -sum(sorted(item["chunk_signal_scores"], reverse=True)[:3]),
            (item["best_distance"] if item["best_distance"] is not None else 999.0),
        ),
    )[: max_sources + 2]

    # --- Phase 3: load full content in parallel and compute source-level score ---
    def _load_source_rows(entry):
        try:
            rows = _get_sorted_document_rows(entry["collection_name"], entry["source"], include_embeddings=False)
        except Exception as exc:
            print(f"[SOURCE_RERANK][WARN] Failed to load source={entry['source']}: {exc}")
            return entry, None
        return entry, rows

    with ThreadPoolExecutor(max_workers=min(len(prelim_ranked), 6)) as pool:
        loaded = list(pool.map(_load_source_rows, prelim_ranked))

    source_candidates: List[dict] = []
    for entry, rows in loaded:
        if not rows:
            continue

        full_content = _merge_chunk_texts([row["content"] for row in rows])
        # For scoring, use only matched chunks (not full doc) so large docs aren't diluted.
        # Preserve the order from matched_chunk_ids (which reflects rerank/round-robin
        # priority) instead of doc order — this ensures the highest-ranked chunk per
        # sub-question appears FIRST in the excerpt and survives the char-budget cut.
        matched_id_order = entry.get("matched_chunk_ids") or []
        chunk_id_scores = entry.get("chunk_id_scores") or {}
        rows_by_id = {r["id"]: r for r in rows}
        seen_ids: set = set()
        unique_ordered_ids: List[str] = []
        for mid in matched_id_order:
            if mid in seen_ids:
                continue
            seen_ids.add(mid)
            unique_ordered_ids.append(mid)
        # Reorder matched ids by per-chunk score (descending) so the most-relevant
        # chunk per source appears FIRST in the excerpt and survives the char cap.
        # Stable sort preserves original ingestion order on score ties.
        unique_ordered_ids.sort(key=lambda mid: -float(chunk_id_scores.get(mid, 0.0)))
        matched_rows: List[dict] = []
        for mid in unique_ordered_ids:
            row = rows_by_id.get(mid)
            if row is not None:
                matched_rows.append(row)
        if not matched_rows:
            matched_rows = rows[:5]
        matched_text = _merge_chunk_texts([r["content"] for r in matched_rows], max_chars=28000)
        excerpt = _merge_chunk_texts([r["content"] for r in matched_rows], max_chars=8000) if matched_rows else _merge_chunk_texts([row["content"] for row in rows], max_chars=8000)
        title = _build_document_title(entry["source"], rows[0]["content"] if rows else "")

        top_chunk_signals = sorted(entry.get("chunk_signal_scores") or [0.0], reverse=True)[:3]
        chunk_signal_strength = sum(top_chunk_signals)
        # Use matched chunks for overlap — full_content dilutes relevance on large docs
        overlap_ratio = _lexical_overlap_ratio(question, matched_text)
        title_overlap = _lexical_overlap_ratio(question, title)
        source_overlap = _lexical_overlap_ratio(question, entry["source"])
        lexical_score = lexical_boost_score(question, matched_text, rows[0].get("metadata") or entry["seed_meta"])

        # Mild penalty for very large docs — capped at 2.0 so big docs aren't unfairly suppressed
        chunk_count_penalty = min(2.0, math.log(max(1, len(rows)), 2) * 0.3) if len(rows) > 20 else 0.0

        source_score = (
            chunk_signal_strength
            + lexical_score
            + (overlap_ratio * 8.0)
            + (title_overlap * 5.0)
            + (source_overlap * 4.0)
            - chunk_count_penalty
        )

        source_candidates.append({
            "collection_name": entry["collection_name"],
            "source": entry["source"],
            "title": title,
            "full_content": full_content,
            "matched_text": matched_text,
            "excerpt": excerpt,
            "chunk_count": len(rows),
            "best_distance": entry["best_distance"],
            "score": round(source_score, 4),
            "matched_chunk_ids": entry["matched_chunk_ids"],
            "matched_chunk_indices": sorted(set(entry.get("matched_chunk_indices") or [])),
            "primary_matched_chunk_index": entry.get("primary_matched_chunk_index"),
        })

    source_candidates.sort(
        key=lambda item: (
            -item["score"],
            (item["best_distance"] if item["best_distance"] is not None else 999.0),
            item["title"].lower(),
        )
    )

    print("[SOURCE_RERANK] Top source candidates:")
    for candidate in source_candidates[:10]:
        print(
            f"[SOURCE_RERANK] score={candidate['score']:.2f} dist={candidate['best_distance']} "
            f"collection={candidate['collection_name']} source={candidate['source']} chunks={candidate['chunk_count']}"
        )

    return source_candidates[:max_sources]


# ---------------------------------------------------------------------
# API Models
# ---------------------------------------------------------------------

class ChatRequest(BaseModel):
    question: str
    top_k: int = 5
    force_reingest: Optional[bool] = False
    ticket_url: Optional[str] = None
    conversation_id: Optional[str] = None
    history: Optional[List[Dict[str, str]]] = None
    teach: Optional[bool] = False
    # If a ticket is selected, set to True for follow-up messages after the first structured reply
    is_followup: Optional[bool] = False
    # File uploads with content
    files: Optional[List[Dict[str, str]]] = None  # List of {name, type, data} where data is base64
    # Enable SSE streaming for the LLM response
    stream: Optional[bool] = False


class KnowledgeAddRequest(BaseModel):
    mode: str  # "content"
    text: Optional[str] = None
    files: Optional[List[Dict[str, str]]] = None
    request_id: Optional[str] = None


class KnowledgeCancelRequest(BaseModel):
    request_id: str


class ChatResponse(BaseModel):
    answer: str
    sources: List[str]
    source_details: Optional[List[dict]] = None


class VectorDbDocumentUpdateRequest(BaseModel):
    collection_name: str
    source: str
    content: str


class VectorDbDocumentDeleteRequest(BaseModel):
    collection_name: str
    source: str


class VectorDbChunkUpdateRequest(BaseModel):
    collection_name: str
    source: str
    chunk_index: int
    content: str


class SystemPromptUpdateRequest(BaseModel):
    template: str


# ---------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------


@app.get("/vector-db/collections")
def vector_db_collections():
    items: List[dict] = []
    for collection_name in (COLLECTION_NAME, MEMORY_COLLECTION_NAME):
        target_collection = _get_vector_collection(collection_name)
        items.append(
            {
                "name": collection_name,
                "description": _vector_collection_description(collection_name),
                "document_count": _count_collection_documents(collection_name),
                "token_count": _count_collection_tokens(collection_name),
                "chunk_count": target_collection.count(),
            }
        )
    return {"items": items}


@app.get("/vector-db/documents")
def vector_db_documents(
    collection_name: str,
    search: str = "",
    limit: int = 24,
    offset: int = 0,
    sort_by: str = "updated_at",
    sort_direction: str = "desc",
):
    safe_limit = max(1, min(limit, 100))
    safe_offset = max(0, offset)
    normalized_sort_by = (sort_by or "updated_at").strip().lower()
    normalized_sort_direction = (sort_direction or "desc").strip().lower()

    if normalized_sort_by not in {"updated_at", "token_count"}:
        raise HTTPException(status_code=400, detail=f"Unsupported sort_by: {sort_by}")
    if normalized_sort_direction not in {"asc", "desc"}:
        raise HTTPException(status_code=400, detail=f"Unsupported sort_direction: {sort_direction}")

    summaries = _group_collection_documents(collection_name)
    search_value = (search or "").strip()[:40]
    if search_value:
        needle = search_value.lower()
        matched_docs: List[dict] = []
        for item in summaries:
            full_content = str(item.get("_search_content") or "").lower()
            title_text = str(item.get("title") or "").lower()
            source_text = str(item.get("source") or "").lower()
            if needle not in full_content and needle not in title_text and needle not in source_text:
                continue
            matched_item = dict(item)
            matched_item["preview"] = _build_vector_search_preview(
                item.get("_search_content", ""), search_value, str(item.get("preview") or "")
            )
            matched_docs.append(matched_item)
        summaries = matched_docs

    if normalized_sort_by == "token_count":
        summaries.sort(
            key=lambda item: (
                int(item.get("token_count") or 0),
                str(item.get("title") or "").lower(),
            ),
            reverse=(normalized_sort_direction == "desc"),
        )
    else:
        summaries.sort(
            key=lambda item: (
                str(item.get("updated_at") or ""),
                str(item.get("title") or "").lower(),
            ),
            reverse=(normalized_sort_direction == "desc"),
        )

    page_items = [
        {key: value for key, value in item.items() if not key.startswith("_")}
        for item in summaries[safe_offset:safe_offset + safe_limit]
    ]
    return {
        "collection_name": collection_name,
        "total": len(summaries),
        "limit": safe_limit,
        "offset": safe_offset,
        "sort_by": normalized_sort_by,
        "sort_direction": normalized_sort_direction,
        "items": page_items,
    }


@app.get("/vector-db/document")
def vector_db_document(
    collection_name: str,
    source: str,
    chunk_limit: int = VECTOR_DB_DEFAULT_CHUNK_LIMIT,
    chunk_offset: int = 0,
    include_embeddings: bool = False,
    include_full_content: bool = True,
):
    return _build_vector_document_detail(
        collection_name,
        source,
        chunk_limit=chunk_limit,
        chunk_offset=chunk_offset,
        include_embeddings=include_embeddings,
        include_full_content=include_full_content,
    )


@app.get("/vector-db/chunk-embedding")
def vector_db_chunk_embedding(collection_name: str, chunk_id: str):
    target_collection = _get_vector_collection(collection_name)
    raw = target_collection.get(ids=[chunk_id], include=["embeddings", "metadatas"])
    ids = list(raw.get("ids", []) or [])
    if not ids:
        raise HTTPException(status_code=404, detail="Chunk not found")

    metas = list(raw.get("metadatas", []) or [])
    embeddings = list(raw.get("embeddings", []) or [])
    safe_meta = (metas[0] if metas else {}) or {}
    embedding = _safe_embedding_to_list(embeddings[0]) if embeddings else []

    return {
        "id": ids[0],
        "chunk": int(safe_meta.get("chunk", 0) or 0),
        "metadata": safe_meta,
        "embedding": embedding,
        "vector_dimensions": len(embedding),
    }


@app.get("/vector-db/chunk")
def vector_db_get_chunk(collection_name: str, source: str, chunk_index: int):
    """Return the text of a single chunk by collection, source, and chunk index."""
    target_collection = _get_vector_collection(collection_name)
    raw = target_collection.get(
        where={"source": source},
        include=["documents", "metadatas"],
    )
    ids = list(raw.get("ids", []) or [])
    docs = list(raw.get("documents", []) or [])
    metas = list(raw.get("metadatas", []) or [])

    for idx, (doc_id, doc_text, meta) in enumerate(zip(ids, docs, metas)):
        safe_meta = meta or {}
        if int(safe_meta.get("chunk", -1)) == chunk_index:
            return {
                "id": doc_id,
                "chunk_index": chunk_index,
                "content": doc_text or "",
                "collection_name": collection_name,
                "source": source,
            }

    raise HTTPException(status_code=404, detail=f"Chunk {chunk_index} not found")


@app.put("/vector-db/chunk")
def vector_db_update_chunk(req: VectorDbChunkUpdateRequest):
    """Update a single chunk: replace its text and re-embed it in place."""
    new_content = (req.content or "").strip()
    if not new_content:
        raise HTTPException(status_code=400, detail="Chunk content cannot be empty")

    target_collection = _get_vector_collection(req.collection_name)
    raw = target_collection.get(
        where={"source": req.source},
        include=["documents", "metadatas"],
    )
    ids = list(raw.get("ids", []) or [])
    metas = list(raw.get("metadatas", []) or [])

    old_id = None
    old_meta = None
    for idx, (doc_id, meta) in enumerate(zip(ids, metas)):
        safe_meta = meta or {}
        if int(safe_meta.get("chunk", -1)) == req.chunk_index:
            old_id = doc_id
            old_meta = dict(safe_meta)
            break

    if old_id is None:
        raise HTTPException(status_code=404, detail=f"Chunk {req.chunk_index} not found")

    new_embedding = embed_text(new_content)
    new_id = str(uuid.uuid4())
    new_meta = dict(old_meta)
    new_meta["edited_at"] = datetime.utcnow().isoformat() + "Z"
    if new_content.startswith(TABLE_ROW_START_TAG):
        new_meta["chunk_type"] = "table_row"
        new_meta.update(extract_table_row_metadata(new_content))
    else:
        new_meta["chunk_type"] = "text"

    try:
        target_collection.add(
            ids=[new_id],
            documents=[new_content],
            metadatas=[new_meta],
            embeddings=[new_embedding],
        )
        target_collection.delete(ids=[old_id])
        invalidate_collection_cache(req.collection_name)
    except Exception as exc:
        try:
            target_collection.delete(ids=[new_id])
        except Exception:
            pass
        print(f"[API][VECTOR_DB][ERROR] chunk update failed: {exc}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to update chunk: {exc}")

    _invalidate_cached_vector_document_payload(req.collection_name, req.source)

    return {
        "message": "Chunk updated and re-embedded successfully.",
        "id": new_id,
        "chunk_index": req.chunk_index,
        "source": req.source,
        "collection_name": req.collection_name,
    }


@app.put("/vector-db/document")
def vector_db_update_document(req: VectorDbDocumentUpdateRequest):
    target_collection = _get_vector_collection(req.collection_name)
    existing = target_collection.get(where={"source": req.source}, include=["metadatas", "documents", "embeddings"])
    existing_ids = list(existing.get("ids") or [])
    existing_metas = list(existing.get("metadatas") or [])
    _raw_docs = existing.get("documents")
    _raw_embs = existing.get("embeddings")
    existing_docs = list(_raw_docs) if _raw_docs is not None else []
    existing_embeddings = list(_raw_embs) if _raw_embs is not None else []
    if not existing_ids:
        raise HTTPException(status_code=404, detail="Document not found")

    new_ids, new_docs, new_metas, new_embeddings = _build_updated_chunk_records(
        req.source,
        req.content,
        existing_metas,
        existing_docs,
        existing_embeddings,
    )

    try:
        CHROMA_BATCH = 500
        for i in range(0, len(new_ids), CHROMA_BATCH):
            end = min(i + CHROMA_BATCH, len(new_ids))
            target_collection.add(
                ids=new_ids[i:end],
                documents=new_docs[i:end],
                metadatas=new_metas[i:end],
                embeddings=new_embeddings[i:end],
            )
        for i in range(0, len(existing_ids), CHROMA_BATCH):
            end = min(i + CHROMA_BATCH, len(existing_ids))
            target_collection.delete(ids=existing_ids[i:end])
        invalidate_collection_cache(req.collection_name)
    except Exception as exc:
        try:
            for i in range(0, len(new_ids), CHROMA_BATCH):
                end = min(i + CHROMA_BATCH, len(new_ids))
                target_collection.delete(ids=new_ids[i:end])
        except Exception:
            pass
        print(f"[API][VECTOR_DB][ERROR] update failed: {exc}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to update document: {exc}")

    _invalidate_cached_vector_document_payload(req.collection_name, req.source)

    return {
        "message": "Document updated successfully.",
        "source": req.source,
        "chunks_added": len(new_ids),
        "detail": _build_vector_document_detail(
            req.collection_name,
            req.source,
            chunk_limit=VECTOR_DB_DEFAULT_CHUNK_LIMIT,
            chunk_offset=0,
            include_embeddings=False,
            include_full_content=False,
        ),
    }


@app.delete("/vector-db/document")
def vector_db_delete_document(req: VectorDbDocumentDeleteRequest):
    target_collection = _get_vector_collection(req.collection_name)
    existing = target_collection.get(where={"source": req.source}, include=["metadatas"])
    existing_ids = list(existing.get("ids", []) or [])
    if not existing_ids:
        raise HTTPException(status_code=404, detail="Document not found")

    try:
        target_collection.delete(ids=existing_ids)
        invalidate_collection_cache(req.collection_name)
    except Exception as exc:
        print(f"[API][VECTOR_DB][ERROR] delete failed: {exc}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {exc}")

    _invalidate_cached_vector_document_payload(req.collection_name, req.source)

    return {
        "message": "Document deleted successfully.",
        "source": req.source,
        "deleted_chunks": len(existing_ids),
    }


@app.get("/health")
def health():
    print("[API][HEALTH] /health called")
    info = {
        "status": "ok",
        "chroma_empty": collection_empty(),
        "wiki_root": WIKI_ROOT,
        "collection_name": COLLECTION_NAME,
        "memory_collection_name": MEMORY_COLLECTION_NAME,
        "embedding_backend": "openai",
        "embedding_model": OPENAI_EMBEDDING_MODEL,
    }
    print(f"[API][HEALTH] {info}")
    return info


@app.get("/azure/tickets")
def azure_tickets(tag: str = "CC"):
    print(f"[API][ADO] /azure/tickets loading merged support-query results tag_param_ignored={tag}")
    try:
        items = ado_list_tickets(tag_contains=tag)
        return {"items": items}
    except Exception as exc:
        print(f"[API][ADO][ERROR] {exc}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(exc))


@app.get("/azure/ticket-by-id")
def azure_ticket_by_id(id: int):
    print(f"[API][ADO] /azure/ticket-by-id lookup id={id}")
    if id <= 0:
        raise HTTPException(status_code=400, detail="id must be a positive integer")

    try:
        item = ado_get_ticket_picker_item(id)
        return {"item": item}
    except HTTPException:
        raise
    except Exception as exc:
        print(f"[API][ADO][ERROR] ticket-by-id lookup failed for id={id}: {exc}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(exc))


@app.get("/azure/tickets/search")
def azure_ticket_search(query: str, limit: int = 100, max_scan: int = 500):
    normalized_query = str(query or "").strip()
    if len(normalized_query) < 2:
        raise HTTPException(status_code=400, detail="query must be at least 2 characters")

    print(f"[API][ADO] /azure/tickets/search query='{normalized_query}' limit={limit} max_scan={max_scan}")
    try:
        items = ado_search_tickets_by_keywords(normalized_query, limit=limit, max_scan=max_scan)
        return {"items": items, "query": normalized_query, "count": len(items)}
    except HTTPException:
        raise
    except Exception as exc:
        print(f"[API][ADO][ERROR] ticket search failed: {exc}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(exc))


@app.get("/system-prompt")
def get_system_prompt():
    template = _read_system_prompt_template()
    return {
        "template": template,
        "path": SYSTEM_PROMPT_FILE,
        "display_path": SYSTEM_PROMPT_DISPLAY_PATH,
    }


@app.put("/system-prompt")
def update_system_prompt(req: SystemPromptUpdateRequest):
    try:
        template = _write_system_prompt_template(req.template)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        print(f"[API][PROMPT][ERROR] Failed to save system prompt: {exc}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to save system prompt: {exc}") from exc

    return {
        "message": "System prompt saved successfully.",
        "template": template,
        "path": SYSTEM_PROMPT_FILE,
        "display_path": SYSTEM_PROMPT_DISPLAY_PATH,
    }


def _knowledge_add_process(req: KnowledgeAddRequest, request_id: Optional[str]) -> dict:
    """Core knowledge ingestion logic — runs in a background thread.
    Returns a result dict with at least {approved: bool}; may include {canceled: bool}.
    """
    _ingest_log("knowledge_add start", request_id, force=True)
    content_text = ""
    source_name = ""
    collection_label = "user_knowledge"
    mode = (req.mode or "").strip().lower()

    if mode == "content":
        text_part = (req.text or "").strip()
        file_part = ""
        if req.files:
            file_part = build_file_context(
                req.files,
                describe_image_func=_describe_image_via_vision,
                cancel_check=lambda: _is_knowledge_request_cancelled(request_id),
                request_id=request_id,
            )
        content_text = "\n\n".join([p for p in [text_part, file_part] if p])
        source_name = f"user:content:{uuid.uuid4()}"
    else:
        return {"approved": False, "message": "mode must be content"}

    if not content_text.strip():
        return {"approved": False, "message": "No content to ingest"}

    if _is_knowledge_request_cancelled(request_id):
        _ingest_log("knowledge_add canceled before embedding started", request_id, force=True)
        return {"approved": False, "message": "Upload stopped by user.", "canceled": True}

    _ingest_log(f"knowledge_add content_ready len={len(content_text)}", request_id, force=True)

    # Validation policy: reject unusable or clearly irrelevant knowledge for both free text and files.
    is_valid_ta9_content, rejection_reason = _validate_ta9_knowledge_content(content_text)
    if not is_valid_ta9_content:
        return {"approved": False, "message": rejection_reason}

    # Chunk, embed, and add with the same retrieval weight as every other document.
    input_type = "file" if req.files or _looks_like_uploaded_file_content(content_text) else "free_text"
    chunks = chunk_text_preserve_table_rows(content_text) if input_type == "file" else chunk_text(content_text)
    if not chunks:
        return {"approved": False, "message": "No chunks generated from content"}

    _ingest_log(f"embedding phase start chunks_total={len(chunks)} input_type={input_type}", request_id, force=True)

    ids: List[str] = []
    docs: List[str] = []
    metas: List[dict] = []
    embeds: List[List[float]] = []
    uploaded_at = datetime.utcnow().isoformat() + "Z"

    for i, ch in enumerate(chunks):
        if _is_knowledge_request_cancelled(request_id):
            _ingest_log(
                f"cancel acknowledged during embedding at chunk={i + 1}/{len(chunks)} partial_chunks_added={len(ids)}",
                request_id,
                force=True,
            )
            return {
                "approved": False,
                "message": "Embedding stopped by user.",
                "chunks_added": len(ids),
                "source": source_name,
                "canceled": True,
            }
        try:
            chunk_kind = "text"
            if ch.startswith(TABLE_ROW_START_TAG):
                chunk_kind = "table_row"
            elif "[IMAGE_BLOCK_ANALYSIS]" in ch or "[PAGE_IMAGE_ANALYSIS_FALLBACK]" in ch:
                chunk_kind = "image_description"
            _ingest_log(
                f"embedding chunk={i + 1}/{len(chunks)} kind={chunk_kind} chars={len(ch)}",
                request_id,
            )
            emb = embed_text(ch)
        except Exception as e:
            print(f"[KNOWLEDGE][ERROR] Embedding failed chunk={i}: {e}")
            traceback.print_exc()
            return {
                "approved": False,
                "message": "Knowledge upload failed during embedding. Nothing was saved to the database.",
                "failed_chunk": i,
                "source": source_name,
            }
        metadata = {
            "source": source_name,
            "chunk": i,
            "uploaded_at": uploaded_at,
            "mode": mode,
            "input_type": input_type,
            "path": "",
            "collection": collection_label,
        }
        if ch.startswith(TABLE_ROW_START_TAG):
            metadata["chunk_type"] = "table_row"
            metadata.update(extract_table_row_metadata(ch))
        else:
            metadata["chunk_type"] = "text"
        ids.append(str(uuid.uuid4()))
        docs.append(ch)
        metas.append(metadata)
        embeds.append(emb)

    if not ids:
        return {"approved": False, "message": "Failed to embed any chunks"}

    memory_collection.add(ids=ids, documents=docs, metadatas=metas, embeddings=embeds)
    invalidate_collection_cache(MEMORY_COLLECTION_NAME)
    _ingest_log(f"knowledge_add complete chunks_added={len(ids)}", request_id, force=True)
    return {
        "approved": True,
        "message": "Knowledge added successfully.",
        "chunks_added": len(ids),
        "source": source_name,
    }


@app.post("/knowledge/add")
def knowledge_add(req: KnowledgeAddRequest):
    """Start knowledge ingestion in a background thread and return immediately.
    Poll GET /knowledge/status/{request_id} for progress and the final result.
    """
    mode = (req.mode or "").strip().lower()
    if mode != "content":
        raise HTTPException(status_code=400, detail="mode must be content")

    request_id = (req.request_id or "").strip() or str(uuid.uuid4())
    try:
        _mark_knowledge_request_started(request_id)
        _store_job_result(request_id, {"status": "processing", "started_at": datetime.utcnow().isoformat() + "Z"})
        _ingest_log("knowledge_add queued as background job", request_id, force=True)

        def _run() -> None:
            try:
                result = _knowledge_add_process(req, request_id)
                _store_job_result(request_id, {"status": "done", **result})
            except Exception as exc:
                print(f"[API][KNOWLEDGE][ERROR] background job exception: {exc}")
                traceback.print_exc()
                _store_job_result(request_id, {"status": "done", "approved": False, "message": f"Processing error: {exc}"})
            finally:
                _clear_knowledge_request(request_id)

        threading.Thread(target=_run, daemon=True).start()
        return {"status": "processing", "request_id": request_id}
    except HTTPException:
        raise
    except Exception as exc:
        print(f"[API][KNOWLEDGE][ERROR] knowledge_add setup failed: {exc}")
        traceback.print_exc()
        _clear_knowledge_request(request_id)
        raise HTTPException(status_code=500, detail=str(exc))


# Maximum characters to send to LLM for content prep (stays within 128k token context)
# ~400k chars ≈ 100k tokens, safe for any OpenAI model
_LLM_PREP_MAX_CHARS = int(os.getenv("LLM_PREP_MAX_CHARS", "400000"))


def _prepare_rag_content(raw_text: str) -> str:
    """Normalize and structure knowledge content for better RAG ingestion.
    
    For large files (>400k chars), LLM prep is skipped — the raw extracted text is already
    high quality from PyMuPDF + vision analysis and does not need LLM restructuring.
    Chunking handles large content; LLM prep is only useful for short user-typed content.
    """
    if not raw_text.strip():
        return ""
    if not OPENAI_API_KEY:
        return raw_text

    # Skip LLM prep for large files — sending 1M+ tokens to LLM would crash.
    # Large files (PDFs etc.) are already well-structured from extraction pipeline.
    if len(raw_text) > _LLM_PREP_MAX_CHARS:
        print(f"[KNOWLEDGE] Content too large for LLM prep ({len(raw_text)} chars > {_LLM_PREP_MAX_CHARS}), skipping — using raw extracted text directly.")
        return raw_text

    prep_prompt = (
        "You are a strict TA9 knowledge normalizer."
        "then rewrite it into clean, structured text for retrieval. "
        "Use ONLY information explicitly present in the input. "
        "Do NOT infer, assume, enrich, generalize, or add domain details that are not written in the source. "
        "Do NOT invent product names, categories, industries, procedures, tags, or entities.\n\n"
        "Output English plain text with these sections in order, each on its own line header:\n"
        "Title:\n"
        "Summary:\n"
        "Key Details:\n"
        "Steps or Procedures (if any):\n"
        "Entities/Fields:\n"
        "Tags:\n\n"
        "Guidelines:\n"
        "- Title: short, specific, based only on source text.\n"
        "- Summary: 2-4 sentences, source-faithful only.\n"
        "- Key Details: bullet list of facts explicitly present in source.\n"
        "- Steps or Procedures: numbered steps only if applicable.\n"
        "- Entities/Fields: list only entities/fields explicitly named in source.\n"
        "- Tags: only source-grounded terms; no invented tags.\n"
        "- If a section has no explicit source data, write: Not provided in source.\n\n"
        "Input:\n" + raw_text
    )
    try:
        prepared = call_llm(prep_prompt, temperature=0.1)
        prepared = (prepared or "").strip()
        return prepared if prepared else raw_text
    except Exception as e:
        print(f"[KNOWLEDGE][WARN] RAG prep failed: {e}")
        return raw_text


def _find_similar_knowledge_sources(text: str, limit: int = 5, max_distance: float = 0.28) -> List[str]:
    """Find existing knowledge sources that are highly similar to the provided text.
    
    The OpenAI embedding API has an 8,191 token limit per request (~32,000 chars).
    For large files, we sample the first representative chunk only.
    """
    if not text.strip():
        return []
    if not OPENAI_API_KEY:
        return []
    try:
        # Truncate to embedding API limit: 8191 tokens ≈ 32,000 characters
        _EMBED_MAX_CHARS = 32000
        sample_text = text[:_EMBED_MAX_CHARS]
        if len(text) > _EMBED_MAX_CHARS:
            print(f"[KNOWLEDGE] Truncating text for similarity check: {len(text)} -> {_EMBED_MAX_CHARS} chars")
        emb = embed_text(sample_text)
        similar_sources: List[str] = []
        for col in (memory_collection, collection):
            results = col.query(query_embeddings=[emb], n_results=limit, include=["metadatas", "distances"])
            metas = results.get("metadatas", [[]])[0] if results else []
            distances = results.get("distances", [[]])[0] if results else []
            for meta, dist in zip(metas or [], distances or []):
                if dist is not None and dist <= max_distance:
                    source = meta.get("source") if meta else None
                    if source:
                        similar_sources.append(str(source))
        seen = set()
        deduped = []
        for src in similar_sources:
            if src in seen:
                continue
            seen.add(src)
            deduped.append(src)
        return deduped
    except Exception as e:
        print(f"[KNOWLEDGE][WARN] Similarity check failed: {e}")
        return []


@app.post("/knowledge/prepare")
def knowledge_prepare(req: KnowledgeAddRequest):
    """Prepare user content for RAG and return the normalized text without validating it."""
    try:
        request_id = (req.request_id or "").strip() or None
        _mark_knowledge_request_started(request_id)
        _ingest_log("knowledge_prepare start", request_id, force=True)
        text_part = (req.text or "").strip()
        file_part = ""
        if req.files:
            file_part = build_file_context(
                req.files,
                describe_image_func=_describe_image_via_vision,
                cancel_check=lambda: _is_knowledge_request_cancelled(request_id),
                request_id=request_id,
            )

        raw_content = "\n\n".join([p for p in [text_part, file_part] if p])
        if not raw_content.strip():
            raise HTTPException(status_code=400, detail="No content to prepare")

        if _is_knowledge_request_cancelled(request_id):
            _ingest_log("knowledge_prepare canceled", request_id, force=True)
            return {
                "approved": False,
                "message": "Preparation stopped by user.",
                "canceled": True,
            }

        prepared_content = _prepare_rag_content(raw_content)
        final_content = prepared_content if prepared_content.strip() else raw_content
        return {
            "approved": True,
            "message": "Knowledge prepared for RAG.",
            "prepared_text": final_content,
        }
    except HTTPException:
        raise
    except Exception as exc:
        print(f"[API][KNOWLEDGE][ERROR] prepare failed: {exc}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(exc))
    finally:
        _clear_knowledge_request((req.request_id or "").strip() or None)


@app.post("/knowledge/cancel")
def knowledge_cancel(req: KnowledgeCancelRequest):
    request_id = (req.request_id or "").strip()
    if not request_id:
        raise HTTPException(status_code=400, detail="request_id is required")
    existed = _cancel_knowledge_request(request_id)
    _ingest_log(f"/knowledge/cancel accepted (was_active={existed})", request_id, force=True)
    return {
        "ok": True,
        "message": "Cancellation requested.",
        "request_id": request_id,
        "was_active": existed,
    }


@app.get("/knowledge/status/{request_id}")
def knowledge_job_status(request_id: str):
    """Poll the status of a background knowledge ingestion job started by POST /knowledge/add.
    Returns {status: 'processing'} while running, or {status: 'done', approved: bool, ...} when finished.
    """
    if not request_id:
        raise HTTPException(status_code=400, detail="request_id is required")
    job = _get_job_result(request_id)
    if not job:
        return {"status": "not_found", "request_id": request_id}
    return {"request_id": request_id, **job}


# ---------------------------------------------------------------------
# Debug/Inspect Endpoint
# ---------------------------------------------------------------------

@app.get("/debug/inspect_vectors")
def inspect_vectors(source_pattern: str = "ado:learn:", limit: int = 10):
    """
    Inspect stored vectors and their documents.
    Use source_pattern to filter (e.g., 'ado:learn:' for tickets, or a specific file name).
    """
    print(f"[API][DEBUG] /debug/inspect_vectors source_pattern={source_pattern} limit={limit}")
    try:
        # Get all items (up to limit) matching the pattern
        results = collection.get(
            include=["documents", "metadatas"],
            limit=min(limit * 100, 10000),  # Fetch more to filter
        )
        
        docs = results.get("documents", [])
        metas = results.get("metadatas", [])
        ids = results.get("ids", [])
        
        # Filter by source pattern
        filtered = []
        for i, (doc, meta) in enumerate(zip(docs, metas)):
            source = meta.get("source", "") if meta else ""
            if source_pattern in source:
                filtered.append({
                    "id": ids[i] if i < len(ids) else None,
                    "source": source,
                    "chunk": meta.get("chunk", 0) if meta else 0,
                    "ticket_id": meta.get("ticket_id") if meta else None,
                    "title": meta.get("title") if meta else None,
                    "type": meta.get("type") if meta else None,
                    "state": meta.get("state") if meta else None,
                    "document_preview": doc[:500] if doc else "",
                    "document_length": len(doc) if doc else 0,
                    "full_document": doc,
                })
                if len(filtered) >= limit:
                    break
        
        return {
            "total_matching": len(filtered),
            "items": filtered,
        }
    except Exception as exc:
        print(f"[API][DEBUG][ERROR] {exc}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(exc))


# ---------------------------------------------------------------------
# Chat Endpoint
# ---------------------------------------------------------------------

@app.post("/chat")
def chat(req: ChatRequest):
    _chat_t0 = time.time()
    print(f"[API][CHAT] /chat called question='{req.question[:200]}' top_k={req.top_k} force_reingest={req.force_reingest} files={len(req.files) if req.files else 0}")
    question = (req.question or "").strip()
    resolved_question = question
    total_rag_steps = 8
    _rag_log_step(1, total_rag_steps, "Request received", f"question_len={len(question)} files={len(req.files) if req.files else 0}")

    conversation_key = (req.conversation_id or req.ticket_url or "default").strip() or "default"
    incoming_history = _normalize_history_messages(req.history)
    if incoming_history:
        dq = deque(incoming_history[-MAX_CONVERSATION_MESSAGES:], maxlen=MAX_CONVERSATION_MESSAGES)
        conversation_store[conversation_key] = dq
    stored_history = list(conversation_store.get(conversation_key, deque()))
    conversation_state = dict(conversation_state_store.get(conversation_key) or {})

    # --- Run history resolution, file processing, and ticket fetching in parallel ---
    _parallel_history_result = [None]
    _parallel_file_context = [""]
    _parallel_ticket_text = [None]
    _parallel_ticket_id = [None]
    _parallel_ticket_key = [None]

    def _run_history_resolution():
        _parallel_history_result[0] = _resolve_question_with_history(
            question,
            stored_history,
            explicit_is_followup=bool(req.is_followup),
        )

    def _run_file_processing():
        if req.files:
            _rag_log_step(2, total_rag_steps, "Processing user-uploaded files", f"count={len(req.files)}")
            try:
                ctx = build_file_context(req.files, describe_image_func=_describe_image_via_vision)
                if ctx:
                    print(f"[API][CHAT] Extracted file context: {len(ctx)} characters")
                    _parallel_file_context[0] = ctx
            except Exception as e:
                print(f"[API][CHAT][WARN] File processing failed: {e}")
        else:
            _rag_log_step(2, total_rag_steps, "Processing user-uploaded files", "no files were attached")

    def _run_ticket_fetch():
        if req.ticket_url:
            _rag_log_step(3, total_rag_steps, "Loading selected ticket context", req.ticket_url)
            _parallel_ticket_key[0] = (req.ticket_url or "").strip() or None
            tid = ado_parse_id_from_url(req.ticket_url)
            if tid is not None:
                _parallel_ticket_id[0] = tid
                _parallel_ticket_key[0] = str(tid)
                print(f"[API][CHAT] Fetching Azure DevOps ticket id={tid}")
                try:
                    _parallel_ticket_text[0] = ado_fetch_ticket_text(int(tid), skip_image_analysis=True)
                except Exception as e:
                    print(f"[API][CHAT][WARN] failed to fetch ticket: {e}")
        else:
            _rag_log_step(3, total_rag_steps, "Loading selected ticket context", "no ticket selected")

    with ThreadPoolExecutor(max_workers=3) as executor:
        futures = [
            executor.submit(_run_history_resolution),
            executor.submit(_run_file_processing),
            executor.submit(_run_ticket_fetch),
        ]
        for f in futures:
            f.result()  # wait for all to complete

    print(f"[API][CHAT][TIMING] Parallel phase took {time.time() - _chat_t0:.2f}s")
    history_resolution = _parallel_history_result[0] or {}
    file_context = _parallel_file_context[0]
    selected_ticket_id = _parallel_ticket_id[0]
    selected_ticket_text = _parallel_ticket_text[0]
    ticket_key = _parallel_ticket_key[0]

    use_history_context = bool(history_resolution.get("use_history"))
    resolved_question = str(history_resolution.get("resolved_question") or question).strip() or question
    # Mirror the resolver window: only the last 2 user turns reach the answer
    # prompt. Anything older is treated as a different topic and excluded so
    # the model cannot drift back to it after a topic shift.
    history_context = _format_history_for_prompt(stored_history, max_messages=2) if use_history_context else ""
    if history_context:
        print(
            f"[API][CHAT] Using conversation history context key={conversation_key} "
            f"messages={len(stored_history)} relation={history_resolution.get('relation')} "
            f"confidence={history_resolution.get('confidence')}"
        )
    else:
        print(f"[API][CHAT] Treating question as standalone key={conversation_key}")

    retrieval_question = resolved_question if use_history_context else question
    # Safety: if the user's current message is itself a compound question
    # (multiple explicit sub-questions in one prompt), the history resolver
    # must NOT drop any of them. Always use the original message for retrieval
    # in that case so every sub-question is decomposed and retrieved.
    if _decompose_compound_question(question):
        if retrieval_question != question:
            print(
                f"[API][CHAT] Compound question detected; ignoring history-resolver rewrite "
                f"and retrieving against original message"
            )
        retrieval_question = question
    if retrieval_question != question:
        print(f"[API][CHAT] Resolved follow-up question for retrieval: '{retrieval_question[:240]}'")

    has_ticket = bool(ticket_key)

    # Enforce non-empty question
    if not question:
        print("[API][CHAT][ERROR] Empty question")
        raise HTTPException(status_code=400, detail="question is empty")

    if _question_requests_previous_answer_source(question):
        answer = _build_previous_answer_source_answer(question, conversation_state)
        _store_conversation_turn(
            conversation_key,
            question,
            answer,
            sources=list(conversation_state.get("sources") or []),
            selected_ticket_id=conversation_state.get("selected_ticket_id"),
            selected_ticket_url=conversation_state.get("selected_ticket_url"),
            used_selected_ticket=bool(conversation_state.get("used_selected_ticket")),
            used_file_context=bool(conversation_state.get("used_file_context")),
        )
        print("[API][CHAT] Answered previous-source clarification from conversation state")
        return ChatResponse(answer=answer, sources=list(conversation_state.get("sources") or []))

    has_active_anchor = bool(
        selected_ticket_text
        or file_context
        or (use_history_context and bool(conversation_state.get("has_grounded_sources")))
    )
    if _question_needs_anchor(question) and not has_active_anchor:
        answer = _build_anchor_clarification_answer(question)
        _store_conversation_turn(
            conversation_key,
            question,
            answer,
            sources=[],
            selected_ticket_id=selected_ticket_id,
            selected_ticket_url=req.ticket_url,
            used_selected_ticket=bool(selected_ticket_text),
            used_file_context=bool(file_context),
        )
        print("[API][CHAT] Blocked ambiguous referential question without an anchor")
        return ChatResponse(answer=answer, sources=[])

    if req.force_reingest:
        print("[API][CHAT] force_reingest=True → calling ingest_wiki_files(force=True)")
        ingest_wiki_files(force=True)

    ta9_mode = _is_ta9_question(retrieval_question)
    is_foundational = _is_foundational_question(retrieval_question)

    # Build base variants (no sub-question decomposition appended) for both
    # the ORIGINAL user question and the history-resolved retrieval question.
    # Sub-questions are added separately at the end so the affinity tracker
    # can rely on them being the tail of augmented_variants.
    def _build_base_variants(q: str) -> List[str]:
        v = _build_query_variants(q, ta9_mode)
        # Strip any decomposed sub-questions (they're appended at the tail);
        # keep only the head (original / normalized / keyword variants).
        decomposed = _decompose_compound_question(q)
        if decomposed:
            v = v[: max(0, len(v) - len(decomposed))]
        return v

    base_head_variants: List[str] = []
    seen_variants: set = set()

    # Highest priority: variants of the LATEST user message.
    for v in _build_base_variants(question):
        if v and v not in seen_variants:
            seen_variants.add(v)
            base_head_variants.append(v)

    # Secondary: variants of the history-resolved question (only if different).
    if retrieval_question.strip() and retrieval_question.strip() != question.strip():
        for v in _build_base_variants(retrieval_question):
            if v and v not in seen_variants:
                seen_variants.add(v)
                base_head_variants.append(v)

    # Lowest priority: history-seed (only if real history was injected).
    if history_context:
        history_seed = (
            f"Conversation context:\n{history_context}\n\n"
            f"Current question:\n{question}\n\n"
            f"Resolved question for retrieval:\n{retrieval_question}"
        )
        if history_seed not in seen_variants:
            seen_variants.add(history_seed)
            base_head_variants.append(history_seed)

    _rag_log_step(4, total_rag_steps, "Preparing retrieval queries", f"variants={len(base_head_variants)} history_used={bool(history_context)}")
    ticket_context_hint = None
    if selected_ticket_text:
        # Keep a short hint to enrich similarity search without overwhelming the question
        ticket_context_hint = (selected_ticket_text[:800] or "").strip()

    primary_emb = None
    ticket_aware_emb = None
    agg_ids: List[str] = []
    agg_distances: List[float] = []
    agg_docs: List[str] = []
    agg_metas: List[dict] = []

    # Per-sub-question chunk affinity. For each candidate chunk we remember the
    # best (smallest) vector distance it achieved against each sub-question's
    # own embedding. After rerank we use this to guarantee every sub-question
    # keeps its own top chunks even if those chunks score low against the full
    # compound question.
    sub_questions_list: List[str] = _decompose_compound_question(retrieval_question)

    # Also decompose the ORIGINAL question; if it has sub-questions of its own
    # they take precedence (latest user intent wins over a history rewrite).
    original_sub_questions = _decompose_compound_question(question)
    if original_sub_questions:
        sub_questions_list = original_sub_questions

    # PRIORITY-TO-LATEST-QUESTION: when no compound decomposition kicked in
    # but the current question differs from the retrieval question (i.e. the
    # history resolver rewrote it), treat the ORIGINAL user message as a
    # forced sub-question. This routes the latest question through the same
    # per-source affinity rerank that already protects compound sub-questions
    # from being crowded out by other variants — guaranteeing the latest
    # intent's top chunks survive dedup and source-grouping.
    if not sub_questions_list and question and question.strip() != retrieval_question.strip():
        sub_questions_list = [question.strip()]

    # Build the FINAL ordered variant list: head (priority order) + sub-questions tail.
    sub_q_set = set(sub_questions_list)
    final_head = [v for v in base_head_variants if v not in sub_q_set]
    query_variants = final_head + sub_questions_list
    # Hard cap to keep total bounded.
    query_variants = query_variants[: _MAX_SUB_QUESTIONS + 4]
    # If the cap truncated past sub_questions, recompute sub_questions_list to
    # match what's actually in query_variants tail.
    if len(query_variants) < len(final_head) + len(sub_questions_list):
        # Truncated into the head — keep sub_questions only if still in tail.
        tail_set = set(query_variants)
        sub_questions_list = [s for s in sub_questions_list if s in tail_set]
    subq_affinity: Dict[str, Dict[int, float]] = {}

    def _chunk_key(idv: Optional[str], doc: Optional[str]) -> str:
        if idv:
            return str(idv)
        return hashlib.sha1(((doc or "")[:512]).encode("utf-8", errors="ignore")).hexdigest()

    def _record_affinity(idv: Optional[str], doc: Optional[str], subq_idx: int, dist: Optional[float]) -> None:
        if subq_idx < 0 or dist is None:
            return
        key = _chunk_key(idv, doc)
        bucket = subq_affinity.setdefault(key, {})
        prev = bucket.get(subq_idx)
        if prev is None or dist < prev:
            bucket[subq_idx] = dist

    try:
        _rag_log_step(5, total_rag_steps, "Running vector retrieval", f"query_variants={len(query_variants)}")
        # Batch-embed all query variants (+ ticket-aware query if applicable) in a single API call
        augmented_variants = [augment_question(qv) for qv in query_variants]

        # Compute the slice of `augmented_variants` that corresponds to sub-questions
        # so we can track per-sub-question chunk affinity below. The base
        # `_build_query_variants` returns [original, normalized, keyword_only]
        # followed by appended sub-questions; if a history seed was prepended
        # the order shifts but sub-questions stay at the tail.
        subq_start_idx = max(0, len(augmented_variants) - len(sub_questions_list)) if sub_questions_list else len(augmented_variants)

        _ticket_aware_idx = None
        if ticket_context_hint:
            similar_query = retrieval_question + "\n\nRelated ticket context:\n" + ticket_context_hint
            if history_context:
                similar_query += "\n\nRecent conversation:\n" + history_context
            _ticket_aware_idx = len(augmented_variants)
            augmented_variants.append(similar_query)
        all_variant_embeddings = embed_texts_batch(augmented_variants)
        if _ticket_aware_idx is not None:
            ticket_aware_emb = all_variant_embeddings[_ticket_aware_idx]
            all_variant_embeddings = all_variant_embeddings[:_ticket_aware_idx]

        for idx, q_emb in enumerate(all_variant_embeddings):
            if idx == 0:
                primary_emb = q_emb
                # Log a short preview of the query embedding used for vector search
                q_preview = ",".join([f"{x:.6f}" for x in q_emb[:EMBED_PREVIEW_COUNT]])
                q_norm = math.sqrt(sum([x * x for x in q_emb]))
                q_hash = hashlib.sha256(
                    ",".join([f"{x:.6f}" for x in q_emb[:16]]).encode("utf-8")
                ).hexdigest()[:12]
                print(f"[API][CHAT] Query embedding preview=[{q_preview}] len={len(q_emb)} norm={q_norm:.6f} hash={q_hash}")
                if LOG_FULL_EMBEDDINGS:
                    print(f"[API][CHAT] Query embeddings FULL={q_emb}")

        # Per-variant retrieval depth. Sub-question variants are dedicated
        # retrievals for one atomic ask, so they need the same depth as a
        # standalone question would (otherwise relevant chunks can fall
        # outside top-K of a large collection). Base/keyword variants stay
        # at a moderate depth because they overlap heavily.
        n_variants = max(1, len(all_variant_embeddings))
        n_subq = len(sub_questions_list)
        base_per_query_k = 30 if n_subq >= 5 else 50
        # Sub-question variants need depth comparable to a single-question pool
        # so the right chunk in a 1000+ chunk doc is reachable, but not so deep
        # that downstream rerank/LLM blows the 15s budget for compound asks.
        subq_per_query_k = 30 if n_subq >= 5 else 40
        print(
            f"[API][CHAT] n_variants={n_variants} n_subq={n_subq} "
            f"base_per_query_k={base_per_query_k} subq_per_query_k={subq_per_query_k}"
        )

        # Run all (variant × collection) Chroma queries in parallel.
        def _run_query(emb, col, col_label, variant_idx, k):
            res = col.query(
                query_embeddings=[emb],
                n_results=k,
                include=["distances", "documents", "metadatas"],
            )
            return res, col_label, variant_idx

        query_jobs = []
        for v_idx, q_emb in enumerate(all_variant_embeddings):
            k = subq_per_query_k if (sub_questions_list and v_idx >= subq_start_idx) else base_per_query_k
            for col, col_label in ((memory_collection, "user_knowledge"), (collection, "wiki")):
                query_jobs.append((q_emb, col, col_label, v_idx, k))

        with ThreadPoolExecutor(max_workers=min(16, max(2, len(query_jobs)))) as ex:
            futures = [ex.submit(_run_query, *job) for job in query_jobs]
            for fut in futures:
                results, col_label, variant_idx = fut.result()
                r_ids = results.get("ids", [[]])[0]
                r_distances = results.get("distances", [[]])[0]
                r_docs = results.get("documents", [[]])[0]
                r_metas = results.get("metadatas", [[]])[0]
                normalized_metas = []
                for meta in r_metas or []:
                    meta = meta or {}
                    meta = {**meta, "collection": col_label}
                    normalized_metas.append(meta)
                agg_ids.extend(r_ids or [])
                agg_distances.extend(r_distances or [])
                agg_docs.extend(r_docs or [])
                agg_metas.extend(normalized_metas)

                # Record per-sub-question chunk affinity using the actual
                # vector distances Chroma returned for that sub-question's
                # embedding.
                if sub_questions_list and variant_idx >= subq_start_idx:
                    sq_idx = variant_idx - subq_start_idx
                    if 0 <= sq_idx < len(sub_questions_list):
                        for d_idx, doc in enumerate(r_docs or []):
                            idv = r_ids[d_idx] if d_idx < len(r_ids) else None
                            dist = r_distances[d_idx] if d_idx < len(r_distances) else None
                            _record_affinity(idv, doc, sq_idx, dist)
    except Exception as exc:
        print(f"[API][CHAT][ERROR] Embedding or vector search failed: {exc}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Embedding/vector search failed: {exc}")

    ids = agg_ids
    distances = agg_distances
    docs = agg_docs
    metas = agg_metas

    # Secondary search: ticket-aware similarity without overriding the question
    if ticket_aware_emb is not None:
        try:
            sim_emb = ticket_aware_emb
            for col, col_label in ((memory_collection, "user_knowledge"), (collection, "wiki")):
                sim_results = col.query(
                    query_embeddings=[sim_emb],
                    n_results=max(20, req.top_k * 4),
                    include=["distances", "documents", "metadatas"],
                )
                sim_ids = sim_results.get("ids", [[]])[0]
                sim_distances = sim_results.get("distances", [[]])[0]
                sim_docs = sim_results.get("documents", [[]])[0]
                sim_metas = sim_results.get("metadatas", [[]])[0]

                normalized_metas = []
                for meta in sim_metas or []:
                    meta = meta or {}
                    meta = {**meta, "collection": col_label}
                    normalized_metas.append(meta)

                if sim_docs:
                    docs = docs + sim_docs
                    metas = metas + normalized_metas
                    ids = ids + sim_ids
                    distances = distances + sim_distances
                    print(f"[API][CHAT] Added {len(sim_docs)} ticket-similar docs to candidates")
        except Exception as e:
            print(f"[API][CHAT][WARN] Similar-ticket search failed: {e}")

    print(f"[API][CHAT] Vector search returned {len(docs)} docs")
    # Log top matches briefly (id, dist, source, chunk, snippet)
    for i, (d, m) in enumerate(zip(docs, metas)):
        if i >= 10:
            break
        did = ids[i] if i < len(ids) else None
        dist = distances[i] if i < len(distances) else None
        snippet = (d[:200].replace("\n", " ") + "...") if d else ""
        m_safe = m or {}
        dist_str = f"{dist:.4f}" if dist is not None else "N/A"
        print(f"[API][CHAT] Result {i}: id={did} dist={dist_str} source={m_safe.get('source', 'N/A')} chunk={m_safe.get('chunk', 'N/A')} snippet='{snippet[:100]}'")

    # No forced root-file injection: all files treated equally

    # Also include memory snippets related to the question
    # Gather memory docs (generic + ticket-specific) but do not pin them; they will be reranked
    try:
        mem_query_emb = primary_emb or embed_text(retrieval_question)
        mem = memory_collection.query(
            query_embeddings=[mem_query_emb],
            n_results=6,
            include=["documents", "metadatas", "distances"],
        )
        mem_docs = mem.get("documents", [[]])[0]
        mem_metas = mem.get("metadatas", [[]])[0]
        mem_distances = mem.get("distances", [[]])[0]
        mem_ids = mem.get("ids", [[]])[0]
        if mem_docs:
            print(f"[API][CHAT] Collected {len(mem_docs)} memory docs for context candidates")
            for i, (d, m) in enumerate(zip(mem_docs, mem_metas)):
                meta = m or {"source": "memory", "chunk": 0}
                meta = {**meta, "collection": "user_knowledge"}
                docs.append(d)
                metas.append(meta)
                distances.append(mem_distances[i] if i < len(mem_distances) else None)
                ids.append(mem_ids[i] if i < len(mem_ids) else None)
    except Exception as e:
        print(f"[API][CHAT][WARN] memory query failed: {e}")

    # Topic-agnostic fallback retrieval when initial context is weak or empty
    initial_profile = _context_confidence_profile(retrieval_question, docs, distances)
    if not file_context and not selected_ticket_text and (not docs or initial_profile["confidence"] < 0.42):
        fallback_variants = _build_fallback_query_variants(retrieval_question)
        print(
            f"[API][CHAT] Triggering fallback retrieval: docs={len(docs)} "
            f"confidence={initial_profile['confidence']} variants={len(fallback_variants)}"
        )
        try:
            valid_variants = [v for v in fallback_variants if v.strip()]
            fallback_embeddings = embed_texts_batch(valid_variants) if valid_variants else []
            for fv_emb in fallback_embeddings:
                for col, col_label in ((memory_collection, "user_knowledge"), (collection, "wiki")):
                    fv_results = col.query(
                        query_embeddings=[fv_emb],
                        n_results=max(12, req.top_k * 3),
                        include=["distances", "documents", "metadatas"],
                    )
                    fv_ids = fv_results.get("ids", [[]])[0]
                    fv_distances = fv_results.get("distances", [[]])[0]
                    fv_docs = fv_results.get("documents", [[]])[0]
                    fv_metas = fv_results.get("metadatas", [[]])[0]

                    normalized_metas = []
                    for meta in fv_metas or []:
                        meta = meta or {}
                        meta = {**meta, "collection": col_label}
                        normalized_metas.append(meta)

                    if fv_docs:
                        docs.extend(fv_docs)
                        metas.extend(normalized_metas)
                        ids.extend(fv_ids)
                        distances.extend(fv_distances)
            print(f"[API][CHAT] Fallback retrieval completed: total docs now={len(docs)}")
        except Exception as e:
            print(f"[API][CHAT][WARN] fallback retrieval failed: {e}")

    if not docs:
        print("[API][CHAT] No docs after all retrieval attempts → returning strict no-context response")
        answer = call_llm(
            "You are a strict RAG assistant. There is no retrievable knowledge-base context for this question. "
            "Do NOT provide speculative or generic product guidance. "
            "Return a short response that clearly states the missing context and asks for the exact document/topic needed.\n\n"
            f"Question: {question}\n\n"
            "Response:"
        )
        _store_conversation_turn(
            conversation_key,
            question,
            answer,
            sources=[],
            selected_ticket_id=selected_ticket_id,
            selected_ticket_url=req.ticket_url,
            used_selected_ticket=bool(selected_ticket_text),
            used_file_context=bool(file_context),
        )
        return ChatResponse(answer=answer, sources=[])

    # ------------------------------------------------------------------
    # Sub-question-aware pre-reorder.
    # Downstream `_smart_deduplicate_and_diversify` (called inside rerank_results)
    # caps each source at max_chunks_per_source=5. When a single source document
    # contains content matching multiple sub-questions, the compound-question
    # rerank order tends to fill those 5 slots with chunks that match the
    # COMPOUND embedding best, dropping chunks each individual sub-question
    # actually needs. By front-loading docs with each sub-question's top
    # affinity chunks (round-robin), we ensure the per-source cap retains a
    # sub-question-balanced slice. Generic for any topic and any number of
    # sub-questions (capped at _MAX_SUB_QUESTIONS).
    # ------------------------------------------------------------------
    if sub_questions_list and docs:
        n_sq_pre = len(sub_questions_list)
        # How many top chunks per sub-question to guarantee survive dedup.
        # Bigger n_sq → fewer per sub to keep total bounded.
        if n_sq_pre <= 2:
            keep_per_sub_pre = 8
        elif n_sq_pre <= 4:
            keep_per_sub_pre = 6
        elif n_sq_pre <= 8:
            keep_per_sub_pre = 5
        else:
            keep_per_sub_pre = 4

        # Build per-sub sorted index list using affinity.
        per_sub_sorted_pre: List[List[int]] = []
        for sq_idx in range(n_sq_pre):
            scored: List[Tuple[float, int]] = []
            for i, doc in enumerate(docs):
                idv = ids[i] if i < len(ids) else None
                key = _chunk_key(idv, doc)
                aff = subq_affinity.get(key, {}).get(sq_idx)
                if aff is None:
                    continue
                scored.append((aff, i))
            scored.sort(key=lambda x: x[0])
            per_sub_sorted_pre.append([i for _aff, i in scored[:keep_per_sub_pre]])

        # Round-robin interleave so EVERY sub-question gets representation
        # in the head of the list before per-source caps apply.
        head_indices: List[int] = []
        head_seen: set = set()
        cursors = [0] * n_sq_pre
        progress = True
        while progress:
            progress = False
            for sq_idx in range(n_sq_pre):
                while cursors[sq_idx] < len(per_sub_sorted_pre[sq_idx]):
                    i = per_sub_sorted_pre[sq_idx][cursors[sq_idx]]
                    cursors[sq_idx] += 1
                    if i in head_seen:
                        continue
                    head_seen.add(i)
                    head_indices.append(i)
                    progress = True
                    break

        # Append the rest of docs (compound rerank order) after the head.
        tail_indices = [i for i in range(len(docs)) if i not in head_seen]
        new_order = head_indices + tail_indices

        docs = [docs[i] for i in new_order]
        metas = [metas[i] for i in new_order]
        distances = [
            (distances[i] if i < len(distances) else None) for i in new_order
        ]
        ids = [(ids[i] if i < len(ids) else None) for i in new_order]
        print(
            f"[PRE-RERANK][SUBQ] reordered docs: subs={n_sq_pre} "
            f"keep_per_sub={keep_per_sub_pre} head={len(head_indices)} total={len(docs)}"
        )

    # For compound questions, raise the per-source cap so that EACH sub-question
    # can keep its top chunks within a single large source document.
    _n_sq_for_cap = len(sub_questions_list) if sub_questions_list else 0
    if _n_sq_for_cap >= 2:
        # Roughly: 5 base + 4 per extra sub-question, capped at 25.
        _max_per_source = min(25, 5 + 4 * _n_sq_for_cap)
    else:
        _max_per_source = 5
    docs, metas, distances, ids = rerank_results(
        retrieval_question, docs, metas,
        distances=distances, ids=ids,
        max_chunks_per_source=_max_per_source,
    )
    if sub_questions_list and len(docs) > 0:
        n_sq = len(sub_questions_list)
        # Generous per-sub quota so each sub-question has enough material for the
        # COMPLETENESS_SELF_CHECK rules in the system prompt.
        if n_sq <= 2:
            per_sub_quota = 8
        elif n_sq <= 4:
            per_sub_quota = 6
        elif n_sq <= 8:
            per_sub_quota = 5
        else:
            per_sub_quota = 4
        # Reserve room for compound-question high scorers as well.
        global_keep = max(8, 24 - per_sub_quota * n_sq)

        kept_keys: set = set()
        kept_order: List[int] = []

        # Pre-compute each sub-question's affinity-sorted chunk list.
        per_sub_sorted: List[List[int]] = []
        for sq_idx in range(n_sq):
            scored_indices = []
            for i, doc in enumerate(docs):
                idv = ids[i] if i < len(ids) else None
                key = _chunk_key(idv, doc)
                aff = subq_affinity.get(key, {}).get(sq_idx)
                if aff is None:
                    continue
                scored_indices.append((aff, i))
            scored_indices.sort(key=lambda x: x[0])
            per_sub_sorted.append([i for _aff, i in scored_indices])

        # Round-robin interleave so every sub-question gets representation in the
        # top positions BEFORE downstream truncation (max_sources, source grouping).
        # Without this, sub1 fills slot 0..N, sub2 fills N..2N, sub3 fills 2N..3N
        # and trailing sub-questions get dropped when the source list is truncated.
        sub_cursors = [0] * n_sq
        sub_taken = [0] * n_sq
        while True:
            progress = False
            for sq_idx in range(n_sq):
                if sub_taken[sq_idx] >= per_sub_quota:
                    continue
                cursor = sub_cursors[sq_idx]
                while cursor < len(per_sub_sorted[sq_idx]):
                    i = per_sub_sorted[sq_idx][cursor]
                    cursor += 1
                    idv = ids[i] if i < len(ids) else None
                    key = _chunk_key(idv, docs[i])
                    if key in kept_keys:
                        continue
                    kept_keys.add(key)
                    kept_order.append(i)
                    sub_taken[sq_idx] += 1
                    progress = True
                    break
                sub_cursors[sq_idx] = cursor
            if not progress:
                break

        # Fill the rest with compound-question rerank order so single-topic
        # high scorers and cross-cutting chunks are not lost.
        for i in range(len(docs)):
            if len(kept_order) >= per_sub_quota * n_sq + global_keep:
                break
            idv = ids[i] if i < len(ids) else None
            key = _chunk_key(idv, docs[i])
            if key in kept_keys:
                continue
            kept_keys.add(key)
            kept_order.append(i)

        docs = [docs[i] for i in kept_order]
        metas = [metas[i] for i in kept_order]
        distances = [distances[i] if i < len(distances) else None for i in kept_order]
        ids = [ids[i] if i < len(ids) else None for i in kept_order]
        print(
            f"[RERANK][SUBQ] compound rerank: subs={n_sq} per_sub={per_sub_quota} "
            f"global_keep={global_keep} kept={len(docs)}"
        )

    print(f"[API][CHAT][TIMING] Retrieval+rerank done at {time.time() - _chat_t0:.2f}s")
    _rag_log_step(6, total_rag_steps, "Ranking and consolidating sources", f"candidate_docs={len(docs)}")
    question_is_procedural = _question_has_procedural_intent(retrieval_question)
    entity_fact_intent = _question_has_entity_fact_intent(retrieval_question)
    broad_coverage_intent = _question_requests_broad_coverage(retrieval_question)

    # Confidence-based behavior: avoid hard refusals and answer in best-effort mode when confidence is low
    final_profile = _context_confidence_profile(retrieval_question, docs, distances)
    low_confidence_mode = (
        not file_context
        and not selected_ticket_text
        and final_profile["confidence"] < 0.35
    )
    print(
        "[API][CHAT] Retrieval confidence "
        f"confidence={final_profile['confidence']} best_distance={final_profile['best_distance']} "
        f"max_overlap={final_profile['max_overlap']} low_confidence_mode={low_confidence_mode}"
    )

    context_blocks: List[str] = []
    source_strings: List[str] = []
    source_detail_list: List[dict] = []
    # For compound questions, allow more source documents so every sub-question's
    # best source survives truncation (each sub-question typically maps to a
    # different doc; with 3 sub-questions and max_sources=8, late sub-questions
    # can lose their source).
    _n_sq_for_sources = len(sub_questions_list) if sub_questions_list else 0
    if has_ticket:
        _max_sources = 5
    elif broad_coverage_intent:
        _max_sources = 12 if _n_sq_for_sources >= 2 else 10
    elif _n_sq_for_sources >= 2:
        _max_sources = min(20, 8 + 3 * _n_sq_for_sources)
    else:
        _max_sources = 8
    source_candidates = _build_source_context_candidates(
        retrieval_question,
        docs,
        metas,
        distances=distances,
        ids=ids,
        max_sources=_max_sources,
    )

    # More sources but cap content per source to stay within 15s response budget
    if has_ticket:
        MAX_CONTEXT_SOURCES = 5
    else:
        MAX_CONTEXT_SOURCES = 8
    print(f"[API][CHAT] Building context with MAX_CONTEXT_SOURCES={MAX_CONTEXT_SOURCES} has_ticket={has_ticket}")

    if source_candidates:
        remaining_slots = MAX_CONTEXT_SOURCES - (1 if selected_ticket_text else 0)
        for idx, candidate in enumerate(source_candidates[:remaining_slots]):
            doc_number = idx + 1
            match_label = f"DOCUMENT {doc_number} — PRIMARY MATCH" if idx == 0 else f"DOCUMENT {doc_number} — SUPPLEMENTAL MATCH"
            src = (
                f"{candidate['collection_name']}::{candidate['source']} "
                f"({candidate['chunk_count']} chunks, best_dist={candidate['best_distance']})"
            )
            source_strings.append(src)
            source_detail_list.append({
                "collection_name": candidate["collection_name"],
                "source": candidate["source"],
                "chunk_count": candidate["chunk_count"],
                "best_distance": candidate["best_distance"],
                "matched_chunks": candidate.get("matched_chunk_indices") or [],
                "primary_matched_chunk": candidate.get("primary_matched_chunk_index"),
            })
            snippet = candidate["excerpt"][:200].replace("\n", " ")
            # Cap content per source to keep total prompt small → fast LLM response
            if has_ticket:
                max_content_chars = 4000
            else:
                max_content_chars = 24000
            # Use matched_text (full matched chunks) instead of excerpt for richer context
            raw_content = str(candidate.get("matched_text") or candidate.get("excerpt") or candidate.get("full_content") or "")
            candidate_content = raw_content[:max_content_chars] if len(raw_content) > max_content_chars else raw_content
            print(f"[API][CHAT] Context source: {src} | snippet='{snippet}...'")
            context_blocks.append(
                f"=== {match_label}: {src} ===\n"
                f"Title: {candidate['title']}\n"
                f"(This is a separate document. Do NOT mix its steps with other documents.)\n"
                f"{candidate_content}\n"
                f"=== END DOCUMENT {doc_number} ==="
            )
    else:
        # Fallback to chunk-level context if source consolidation could not be built.
        remaining_slots = 10 - (1 if selected_ticket_text else 0)
        for i, (doc, meta) in enumerate(list(zip(docs, metas))[:remaining_slots]):
            id_val = ids[i] if i < len(ids) else None
            dist_val = distances[i] if i < len(distances) else None
            src = f"{meta.get('source')} (chunk {meta.get('chunk', 0)}) id={id_val} dist={dist_val}"
            source_strings.append(src)
            snippet = doc[:200].replace("\n", " ")
            print(f"[API][CHAT] Context source: {src} | snippet='{snippet}...'")
            context_blocks.append(f"Source: {src}\n{doc}")

    if selected_ticket_text:
        # Include full ticket content — _fit_prompt_to_model_budget will truncate if needed
        src = f"azure-devops:{selected_ticket_id} (selected ticket)"
        source_strings.append(src)
        snippet = selected_ticket_text[:200].replace("\n", " ")
        print(f"[API][CHAT] Context source: {src} | snippet='{snippet}...' | ticket_chars={len(selected_ticket_text)}")
        context_blocks.append(f"Source: {src}\n{selected_ticket_text}")

    context_text = "\n\n---\n\n".join(context_blocks)

    # Enhanced prompt that handles potentially redundant context intelligently
    foundational_instruction = (
        "10. FOUNDATIONAL QUESTIONS: If asked about core product concepts (data models, entities, features, Federated Search, Cases, Link Analysis etc.), "
        "provide a clear explanation using only the available context. If context is incomplete, clearly state what is missing. "
        "Do not add unsupported product details.\n\n"
        if is_foundational or ta9_mode
        else ""
    )
    
    ta9_instruction = (
        "11. If the question is about TA9/IntSight features or platform capabilities, "
        "provide a detailed answer grounded in retrieved context, and avoid examples that are not present in context. "
        "Use bullet points when helpful.\n\n"
        if ta9_mode
        else ""
    )

    # Build context sections with proper priority
    # When files are attached, they should be PRIMARY source of truth
    context_sections = []
    
    if file_context:
        # Files come FIRST and are marked as primary
        context_sections.append(
            f"=== USER-PROVIDED FILES (PRIMARY SOURCE) ===\n"
            f"The user has attached files. Answer their question using this content as the main source.\n\n"
            f"{file_context}\n"
        )
    
    if selected_ticket_text:
        # Ticket context is secondary if files exist, primary otherwise
        priority_label = "REFERENCE" if file_context else "PRIMARY"
        context_sections.append(
            f"=== SELECTED TICKET ({priority_label}) ===\n"
            f"{selected_ticket_text}\n"
        )
    
    # Include KB context as reference even with files for how-to/procedural questions.
    # Screenshot-only context is often insufficient for actionable guidance.
    should_include_rag_context = True
    if file_context:
        question_lc = (retrieval_question or "").lower()
        procedural_intent = any(term in question_lc for term in [
            "how", "step", "steps", "add", "create", "configure", "setup", "set up"
        ])
        best_distance = distances[0] if distances else 1.0
        top_overlap = _lexical_overlap_ratio(question, docs[0]) if docs else 0.0
        # More permissive threshold so attached screenshots still benefit from the manual/KB.
        should_include_rag_context = procedural_intent or ta9_mode or best_distance < 0.75 or top_overlap >= 0.10
        if not should_include_rag_context:
            print(
                "[API][CHAT] Skipping RAG context with file attachment: "
                f"best_distance={best_distance} top_overlap={top_overlap}."
            )
    
    if context_text and should_include_rag_context:
        # Wiki context is always supplemental
        priority_label = "REFERENCE" if (file_context or selected_ticket_text) else "PRIMARY"
        context_sections.append(
            f"=== KNOWLEDGE BASE ({priority_label}) ===\n"
            f"{context_text}\n"
        )

    if source_candidates and question_is_procedural:
        context_sections.append(
            "=== RETRIEVAL & COMPLETENESS PRIORITY ===\n"
            "For how-to or configuration questions, identify which retrieved document BEST answers the specific question asked.\n"
            "Use that document's procedure as your primary answer. Do NOT merge procedural steps from different documents.\n"
            "If the user's question asks about MULTIPLE distinct operations (e.g. 'entities AND relations'), "
            "provide the procedure for each from its own source document, clearly separated.\n"
            "CRITICAL: Each retrieved document describes a SPECIFIC procedure. Before using content from a document, "
            "verify that the document's procedure matches what the user is asking. A document about 'creating a relation data model' "
            "is NOT the same as 'defining link analysis in an existing data model'. Use only the document whose topic matches the question.\n"
            "If only one document matches the user's actual question, use only that document and ignore unrelated ones.\n"
        )

    if source_candidates and broad_coverage_intent:
        context_sections.append(
            "=== BREADTH COVERAGE PRIORITY ===\n"
            "The user is asking for alternatives, additional options, different methods, or broader coverage. "
            "Synthesize distinct valid approaches from all relevant matches. Do not say there is only one way unless the retrieved context clearly supports that conclusion after considering the supplemental matches too.\n"
        )

    if source_candidates and entity_fact_intent:
        context_sections.append(
            "=== ENTITY FACT PRIORITY ===\n"
            "When the question asks about company names, product names, aliases, naming rules, or how many products are listed, "
            "prefer the documents that explicitly define or enumerate those facts. "
            "If the context lists the relevant items, answer directly from that list and state the count.\n"
        )

    if file_context and should_include_rag_context:
        context_sections.append(
            "=== ANSWERING STRATEGY ===\n"
            "Use the screenshot to identify the current UI state, and use knowledge-base chunks to provide concrete step-by-step instructions.\n"
            "If the exact button/path is not visible in screenshot but appears in KB, state that clearly as KB-based guidance.\n"
        )

    if low_confidence_mode:
        context_sections.append(
            "=== RETRIEVAL CONFIDENCE NOTICE ===\n"
            "The retrieved context may only partially match the question. "
            "Provide a best-effort answer grounded in available context, state assumptions clearly, "
            "and ask one concise follow-up question to close the gap.\n"
        )

    if history_context:
        context_sections.append(
            "=== RECENT CONVERSATION CONTEXT ===\n"
            "Use this to resolve references like 'that', 'those', or follow-up clarifications.\n"
            f"{history_context}\n"
        )
    else:
        context_sections.append(
            "=== QUESTION ISOLATION RULE ===\n"
            "Treat the current user question as standalone. Do not use previous conversation topics to infer missing nouns, products, areas, or procedures unless the current question explicitly refers back to them.\n"
        )

    if retrieval_question != question:
        context_sections.append(
            "=== FOLLOW-UP INTERPRETATION ===\n"
            "The current question was interpreted with recent context for retrieval.\n"
            f"Original user wording: {question}\n"
            f"Resolved standalone retrieval question: {retrieval_question}\n"
            "Answer the user's original wording while using the resolved intent for accuracy.\n"
        )
    
    combined_context = "\n\n".join(context_sections)
    combined_context_tokens = _estimate_text_tokens(combined_context)
    _rag_log_step(
        7,
        total_rag_steps,
        "Building final prompt context",
        f"context_chars={len(combined_context)} context_tokens={combined_context_tokens}",
    )

    system_msg, user_msg, combined_context = _fit_prompt_to_model_budget(
        question=question,
        combined_context=combined_context,
        foundational_instruction=foundational_instruction,
        ta9_instruction=ta9_instruction,
    )
    prompt_chars = len(system_msg) + len(user_msg)

    try:
        print(f"[API][CHAT][TIMING] Pre-LLM at {time.time() - _chat_t0:.2f}s | prompt_chars={prompt_chars}")
        _rag_log_step(
            8,
            total_rag_steps,
            "Generating grounded answer",
            f"prompt_chars={prompt_chars} prompt_tokens={_estimate_text_tokens(system_msg) + _estimate_text_tokens(user_msg)}",
        )

        # --- Streaming path ---
        if req.stream:
            def _stream_generator():
                # Send sources first
                yield f"data: {json.dumps({'type': 'sources', 'sources': source_strings, 'source_details': source_detail_list})}\n\n"
                # Stream LLM chunks
                full_answer_parts = []
                try:
                    for chunk_text_part in call_llm_stream(user_msg, temperature=0.0, max_tokens=1000, system_message=system_msg):
                        full_answer_parts.append(chunk_text_part)
                        yield f"data: {json.dumps({'type': 'chunk', 'content': chunk_text_part})}\n\n"
                except Exception as stream_exc:
                    yield f"data: {json.dumps({'type': 'error', 'content': str(stream_exc)})}\n\n"
                    return
                # Signal done
                full_answer = "".join(full_answer_parts)
                answer = _normalize_noncode_fenced_blocks(full_answer)
                yield f"data: {json.dumps({'type': 'done'})}\n\n"
                # Post-processing (conversation store, teach) in background
                _store_conversation_turn(
                    conversation_key,
                    question,
                    answer,
                    sources=source_strings,
                    selected_ticket_id=selected_ticket_id,
                    selected_ticket_url=req.ticket_url,
                    used_selected_ticket=bool(selected_ticket_text),
                    used_file_context=bool(file_context),
                )
                print(f"[API][CHAT] Stream complete len={len(answer)} | total_time={time.time() - _chat_t0:.2f}s")
            return StreamingResponse(
                _stream_generator(),
                media_type="text/event-stream",
                headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
            )

        # --- Non-streaming path ---
        draft_answer = call_llm(user_msg, temperature=0.0, max_tokens=1000, system_message=system_msg)
        # NOTE: Removed 4 post-processing LLM calls that were degrading answers:
        # _ground_answer_against_context — was stripping valid content and saying "not found"
        # _enforce_specific_grounded_answer — redundant with system prompt rules
        # _llm_ensure_answer_completeness — redundant with system prompt rules
        # _llm_verify_answer_relevance — was nuking entire answers with "does not contain a direct answer"
        # The system prompt already handles grounding, specificity, and completeness.
        answer = _normalize_noncode_fenced_blocks(draft_answer)
    except Exception as exc:
        print(f"[API][CHAT][ERROR] LLM call failed: {exc}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"LLM call failed: {exc}")

    _store_conversation_turn(
        conversation_key,
        question,
        answer,
        sources=source_strings,
        selected_ticket_id=selected_ticket_id,
        selected_ticket_url=req.ticket_url,
        used_selected_ticket=bool(selected_ticket_text),
        used_file_context=bool(file_context),
    )

    # Optional teach: persist summarized lesson to memory collection
    if req.teach:
        try:
            lesson_prompt = (
                "Summarize the key knowledge learned from the following Q&A. "
                "Return a concise explanation (<= 180 words) followed by 6-10 bullet points of exact steps or commands. "
                "Avoid sensitive secrets. Keep it generally reusable.\n\n"
                f"Question:\n{question}\n\nAnswer:\n{answer}\n\nSummary + Steps:"
            )
            lesson = call_llm(lesson_prompt)
            emb = embed_text(lesson)
            src = "memory:generic"
            if req.ticket_url:
                tid = ado_parse_id_from_url(req.ticket_url)
                if tid:
                    src = f"memory:ticket:{tid}"
            memory_collection.add(
                ids=[str(uuid.uuid4())],
                documents=[lesson],
                embeddings=[emb],
                metadatas=[{"source": src, "created_at": datetime.utcnow().isoformat()}],
            )
            print(f"[API][CHAT] Stored lesson to memory collection source={src} len={len(lesson)}")
        except Exception as e:
            print(f"[API][CHAT][WARN] Failed to store memory: {e}")

    print(f"[API][CHAT] Returning answer len={len(answer)} with {len(source_strings)} sources | total_time={time.time() - _chat_t0:.2f}s")
    return ChatResponse(answer=answer, sources=source_strings, source_details=source_detail_list)